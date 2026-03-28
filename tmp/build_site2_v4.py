from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph

ROOT = Path(r"C:\Algonquin\Winter2026\Emerging_Tech\Project")
SRC = ROOT / "Site2_SourceBased_Technical_Report_V3.9_Revised.docx"
DST = ROOT / "Site2_SourceBased_Technical_Report_V4.0.docx"
SMB_CONF = ROOT / "c2fs_smb.conf.new"


def find_exact(doc: Document, text: str) -> Paragraph:
    for para in doc.paragraphs:
        if para.text.strip() == text:
            return para
    raise ValueError(f"Paragraph not found: {text!r}")


def next_nonempty(doc: Document, paragraph: Paragraph) -> Paragraph:
    found = False
    for para in doc.paragraphs:
        if found and para.text.strip():
            return para
        if para._p == paragraph._p:
            found = True
    raise ValueError("No next nonempty paragraph found")


def prev_nonempty(doc: Document, paragraph: Paragraph) -> Paragraph:
    previous = None
    for para in doc.paragraphs:
        if para._p == paragraph._p:
            if previous is None:
                raise ValueError("No previous nonempty paragraph found")
            return previous
        if para.text.strip():
            previous = para
    raise ValueError("No previous nonempty paragraph found")


def insert_after(paragraph: Paragraph, text: str = "", style: Optional[str] = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._p
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def style_name(doc: Document, preferred: str, fallback: str = "Normal") -> str:
    try:
        _ = doc.styles[preferred]
        return preferred
    except KeyError:
        return fallback


def add_monospace_paragraph(after: Paragraph, text: str, doc: Document) -> Paragraph:
    para = insert_after(after, style=style_name(doc, "No Spacing"))
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return para


def replace_paragraph_text(doc: Document, old: str, new: str) -> Paragraph:
    para = find_exact(doc, old)
    para.text = new
    return para


def replace_first_containing(doc: Document, needle: str, new: str) -> Paragraph:
    for para in doc.paragraphs:
        if needle in para.text:
            para.text = new
            return para
    raise ValueError(f"Paragraph containing {needle!r} not found")


def apply_cover_page_updates(doc: Document) -> None:
    next_nonempty(doc, find_exact(doc, "Document Version")).text = "4.0"
    next_nonempty(doc, find_exact(doc, "Document Date")).text = "March 27, 2026"
    contributors_value = next_nonempty(doc, find_exact(doc, "Engineering Contributors"))
    team_label = insert_after(contributors_value, "Team Name", contributors_value.style.name)
    insert_after(team_label, "Site 2 Team", contributors_value.style.name)
    date_value = next_nonempty(doc, find_exact(doc, "Document Date"))
    due_label = insert_after(date_value, "Submission Due Date", date_value.style.name)
    insert_after(due_label, "March 26, 2026", date_value.style.name)


def apply_front_matter_updates(doc: Document) -> None:
    find_exact(doc, "Executive Summary").style = doc.styles["Heading 1"]
    find_exact(doc, "List of Tables").style = doc.styles["Heading 1"]


def rewrite_executive_summary(doc: Document) -> None:
    heading = find_exact(doc, "Executive Summary")
    first_para = next_nonempty(doc, heading)
    to_remove = []
    current = first_para
    while current.text.strip() != "1. Introduction":
        to_remove.append(current)
        current = next_nonempty(doc, current)
    summaries = [
        "Site 2 was designed and deployed as a multi-tenant managed service environment supporting MSP administration together with Company 1 and Company 2 operations inside one controlled site. OPNsense provides segmentation and remote-entry control, Samba-based directory services provide identity and DNS, isolated iSCSI-backed storage supports file services, internal web applications are published under named hostnames, and Veeam provides backup with offsite protection.",
        "Validation confirmed that the main service layers are operating as intended. Administrative access reaches the approved management systems, both company namespaces resolve across the documented tenant contexts, storage is mounted and presented through structured shares, internal web delivery responds by hostname rather than by raw address, and the backup design aligns with the observed repository and copy paths. Site 2 therefore behaves as an integrated service platform rather than as a collection of unrelated virtual machines.",
        "This report is written to support formal handover, assessment, and ongoing support. It explains what was built, why the design was structured this way, which parts were directly observed, and where the remaining limits of direct validation still need to be acknowledged by the receiving team. On that basis, Site 2 can be handed over with a clear operational record of both its design intent and its current state.",
    ]
    first_para.text = summaries[0]
    first_para.style = doc.styles[style_name(doc, "Body Text")]
    anchor = first_para
    for text in summaries[1:]:
        anchor = insert_after(anchor, text, style_name(doc, "Body Text"))
    for para in to_remove[1:]:
        remove_paragraph(para)


def update_introduction_and_background(doc: Document) -> None:
    replace_paragraph_text(doc, "The purpose of this document is to describe Site 2 in a form that supports operational handover, ongoing support, and formal service review.", "The purpose of this document is to describe Site 2 in a form that supports formal assessment, operational handover, ongoing support, and later service review.")
    replace_paragraph_text(doc, "This version is grounded in the available design inputs, the current Site 2 evidence set, official vendor documentation for the deployed technologies, and current operating observations recorded on March 23, March 24, and March 25, 2026.", "This version is grounded in the available design inputs, the current Site 2 evidence set, official vendor documentation for the deployed technologies, and operating observations recorded during the March 23-25, 2026 validation window.")
    replace_paragraph_text(doc, "The report deliberately follows a formal handover structure so that it can be reviewed by technical leads, client stakeholders, and successor operators without requiring a separate explanatory walkthrough. The emphasis is on design rationale, operating state, and service relationships rather than on build notes alone.", "We approached Site 2 as a whole-site operating problem rather than a server-by-server build recap. The hardest part of the project was keeping the explanation honest: some behaviors were directly observed live, while other parts of the story had to be reconstructed from configuration evidence, inventory records, and the way the services interacted under test.")
    bg_heading = find_exact(doc, "2. Background")
    insert_after(bg_heading, "This document is written for IT staff taking over responsibility for Site 2 and for MSP technicians who may need to support or modify the environment later. Readers are expected to be comfortable with basic networking concepts, Windows Server administration, and Linux command-line operations. No prior knowledge of Site 2 is assumed; this section provides the context needed before the service discussions begin.", style_name(doc, "Body Text"))


def update_discussion(doc: Document) -> None:
    replace_paragraph_text(doc, "3.2 Site 2 Logical Service Inventory and Platform Roles", "3.2 Service Inventory and Platform Layout")
    replace_paragraph_text(doc, "3.9 Requirement-to-Implementation Traceability", "3.9 Requirements Coverage")
    replace_paragraph_text(doc, "3.10 Service Dependency, Failure Domains, and Access Model", "3.10 Dependencies, Failure Domains, and Access")
    replace_paragraph_text(doc, "3.11 Storage, Backup, and Recovery Flow", "3.11 Data Protection Flow")
    replace_paragraph_text(doc, "3.13 Troubleshooting and Fast Triage Guide", "3.13 Fast Triage Guide")
    replace_paragraph_text(doc, "Site 2 should be read as a complete service site, not as a loose grouping of servers. It combines MSP administrative control, Company 1 cross-site service delivery, Company 2 core tenant services, isolated storage transport, and backup or recovery functions inside one routed operating model. That whole-site view matters because the design only makes sense when those layers are evaluated together.", "The first thing a new operator needs to understand about Site 2 is that pulling any one service out of context makes the rest harder to explain. The network boundary only makes sense when you know which identities it protects, the storage story only makes sense when the SAN isolation underneath it is visible, and the recovery plan only makes sense when the management path is already understood.")
    replace_first_containing(doc, "The current operating observations and the configuration evidence support that whole-site reading clearly.", "Read that way, Site 2 is one operating site with three service perspectives. MSP systems define entry, policy, and recovery control; Company 1 contributes directory, file, client, web, and isolated storage roles; Company 2 contributes its own identity, file, client, web, and storage chain. Together, those components form the actual operating boundary of Site 2 — the set of systems a support engineer would need to understand before any meaningful triage could begin.")
    replace_paragraph_text(doc, "The system inventory matters because it demonstrates planned platform structure rather than ad-hoc implementation. Site 2 was built as a separated service estate with clear roles for MSP control, Company 1 delivery, Company 2 operations, storage transport, and protection services. That role-mapping view is the foundation for understanding dependency, failure domain, and support ownership.", "The inventory matters because it shows where responsibility really lives. Site 2 was not assembled as a generic pool of Windows and Linux virtual machines; it was laid out as a service estate with distinct roles for MSP control, Company 1 delivery, Company 2 operations, storage transport, and protection. Once that layout is clear, the later sections on dependency, failure domains, and support ownership become much easier to follow.")
    replace_paragraph_text(doc, "This platform layout is useful because it demonstrates role separation, shows intentional infrastructure design, and makes the environment easier to understand for an operations handoff without relying on unnecessary platform-specific background detail.", "This platform layout makes the environment easier to read because role separation is visible before any service-specific detail is introduced. It shows intentional infrastructure design without forcing the reader to wade through hypervisor detail that does not change the operating story.")
    replace_paragraph_text(doc, "Seen in that light, the inventory is not merely a host list. It is the service map of the entire site: MSP systems govern entry and recovery, Company 1 and Company 2 systems provide tenant-facing capability, and isolated SAN nodes exist only to feed the file-service layer. That is why later sections repeatedly return to the same systems when explaining dependencies, user-visible behavior, and recovery posture.", "Seen in that light, the inventory is more than a host list. It is the service map of the entire site: MSP systems govern entry and recovery, Company 1 and Company 2 systems provide tenant-facing capability, and isolated SAN nodes exist only to feed the file-service layer. The same systems reappear later because the environment depends on their relationships, not just on their presence.")
    replace_paragraph_text(doc, "The network layer is the controlling discipline of Site 2. If routing, segmentation, and exposure control are not well designed, then identity, storage, web delivery, and backup quickly become difficult to defend. The available OPNsense configuration evidence shows that Site 2 applies a disciplined bastion-first model rather than relying on broad edge exposure.", "Everything else in Site 2 depends on the network design being correct. If segmentation or exposure control is wrong, no amount of careful Samba configuration or Veeam scheduling will make the site secure or supportable. The available OPNsense evidence shows a disciplined bastion-first design instead of broad edge publication.")
    replace_paragraph_text(doc, "From a documentation standpoint, this policy model explains several otherwise separate observations in one place: why only jump systems appear at the edge, why inter-tenant reachability is bounded, why dual-hostname web access works across sites, and why Veeam can use a specific inter-site path without collapsing the rest of the segmentation model.", "This policy model explains several otherwise separate observations in one place: why only jump systems appear at the edge, why inter-tenant reachability is bounded, why dual-hostname web access works across sites, and why Veeam can use a specific inter-site path without collapsing the rest of the segmentation model.")
    replace_paragraph_text(doc, "The Company 1 section is therefore documented in the same terms as the rest of the report: service role, reachable administrative path, observed operating state, and architectural meaning. That keeps the narrative focused on how the service set behaves within the site model rather than reducing Company 1 to a single cross-site dependency.", "Company 1 is documented in the same terms as the rest of the report: service role, reachable administrative path, observed operating state, and architectural meaning. That keeps the narrative focused on how the service set behaves inside the site model rather than reducing Company 1 to a simple cross-site dependency.")
    replace_first_containing(doc, "The Company 1 service arrangement is deliberate rather than incidental.", "The Company 1 service arrangement is deliberate rather than incidental. Two domain controllers distribute authentication and naming responsibility across separate nodes, reducing single-point risk and giving the receiving support team a familiar enterprise baseline to reason from. C1DC1 and C1DC2 follow the same dual-node discipline applied to Company 2 identity, and they are expected to behave consistently across directory, DNS, and Kerberos-path functions. That parallel structure reflects a deliberate decision to apply the same resilience model to both tenant identity layers. C1FS remains separate from the domain-controller role so that file-service activity, share administration, and storage consumption do not compete directly with directory duties. The separation keeps user-data operations, identity operations, and storage operations readable as different support concerns.")
    table8 = find_exact(doc, "Table 8. Company 1 service summary")
    insert_after(table8, "Company 1 was not observed through exactly the same live CLI depth as Company 2, and that distinction needs to stay explicit. Direct observation covered cross-site reachability, Company 1 web behavior, C1UbuntuClient hostname validation, and the addressing evidence for C1SAN. C1WindowsClient placement, parts of the C1FS story, and some of the broader Company 1 service model were corroborated through inventory evidence, network-path checks, and cross-site behavior rather than by deep local inspection. That does not make Company 1 less important; it only describes the form of evidence that was available.", style_name(doc, "Body Text"))
    replace_paragraph_text(doc, "The identity layer is one of the strongest and most operationally important parts of Site 2. Current observations on C2IdM1 and C2IdM2 showed:", "The identity layer is where the Company 2 design becomes easiest to trust. Live checks on C2IdM1 and C2IdM2 showed the services that later sections depend on most heavily:")
    replace_paragraph_text(doc, "The storage layer is one of the clearest indicators that Site 2 was engineered rather than improvised. The environment uses isolated SAN connectivity, mounted block storage, structured share presentation, and synchronization behavior in a way that aligns cleanly with the project’s storage objectives.", "The storage section is where the site stops looking abstract and starts looking operational. Site 2 uses isolated SAN connectivity, mounted block storage, structured share presentation, and synchronization in a way that matches the project goals without collapsing storage traffic into the ordinary routed networks.")
    replace_paragraph_text(doc, "That sequence is operationally important: storage is presented over isolated transport, mounted on the file-service host, structured into public and private paths, and then surfaced through Samba. Writing the section in that order makes it clear where responsibility changes from storage delivery to file presentation.", "This sequence matters because responsibility changes at each step: isolated transport delivers the block device, the file-service host mounts it, Samba turns it into named shares, and only then do users encounter it. Keeping those layers separate is what makes file problems diagnosable without treating storage, permissions, and namespace issues as the same kind of fault.")
    table10 = find_exact(doc, "Table 10. Storage and isolated SAN summary")
    insert_after(table10, "We chose an iSCSI-backed model here because it keeps the file server in control of block storage without turning the tenant network into a general storage fabric. An NFS-first design would have exposed file transport higher in the stack, and a direct-attached model would have made the storage path too tightly bound to a single host. iSCSI lets the file layer own presentation while the SAN remains a small, isolated transport domain underneath it. That isolation matters just as much for security as it does for support: block traffic stays off tenant LAN and DMZ paths, and users never interact with the storage network directly.", style_name(doc, "Body Text"))
    replace_paragraph_text(doc, "This separation is also one of the design’s strongest operational characteristics. End users and routine support processes never need to interact with the SAN endpoints directly. The file-service hosts absorb the storage complexity and present stable SMB paths above it, which is the more supportable model for day-to-day operations and formal handover.", "This separation is also one of the design’s strongest operational characteristics. End users and routine support processes never need to interact with the SAN endpoints directly. The file-service hosts absorb the storage complexity and present stable SMB paths above it, which keeps day-to-day operations focused on shares, permissions, and mounted volumes instead of on raw storage transport.")
    replace_paragraph_text(doc, "Client-side observation is essential because a service environment is only meaningful when directory, web, and file outcomes are visible from real consumer systems. For Site 2, the client perspective shows that the design is not only correctly configured on servers, but also usable from both company contexts.", "Client systems are where the design either holds together or falls apart. Looking only at servers would say very little about whether users can actually resolve names, open shares, or reach the internal web applications from their own company context.")
    replace_paragraph_text(doc, "From an architectural standpoint, that distinction matters because it aligns the web layer to named service delivery rather than opportunistic address-based exposure. It is the more defensible model for an internal enterprise service environment and it reinforces the idea that DNS, routing, and web configuration were designed to operate together.", "Architecturally, this aligns the web layer to named service delivery instead of opportunistic address-based exposure. DNS, routing, certificates, and web configuration all point to the same contract: users reach the service by its intended name, not by guessing its address.")
    replace_paragraph_text(doc, "Backup and recovery must be documented at the same level of care as identity or networking because they define whether the environment is supportable after failure. In Site 2, the available design evidence and the current operating observations complement each other well.", "Backup is one of the places where the project becomes most obviously operational. A site that can authenticate users and publish internal applications but cannot recover from failure is not ready for real support, so this section treats protection with the same weight as identity and networking.")
    backup_basis = replace_paragraph_text(doc, "This protection model is intentionally broader than a simple virtual-machine backup statement. It combines workload protection, file-share protection, repository separation, and offsite copy so that failure handling does not depend on a single recovery mechanism.", "This protection model is intentionally broader than a simple virtual-machine backup statement. It combines workload protection, file-share protection, repository separation, and offsite copy so that failure handling does not depend on a single recovery mechanism.")
    insert_after(backup_basis, "The backup scope is also intentionally split. Virtual-machine jobs preserve operating system state, application state, and role configuration, while the file-share backup covers user data recovery without forcing an operator to restore an entire server. Keeping those paths separate matters in practice: a user-data incident should not require full-host recovery, and a host failure should not be confused with a routine file-restore request.", style_name(doc, "Body Text"))
    recovery_para = replace_paragraph_text(doc, "Each protection layer in Table 17 answers a different failure scenario. Synchronization supports content continuity under normal operating conditions. Local Veeam retention covers rapid recovery without inter-site dependency. The Site 1 copy addresses the scenario where Site 2 storage is itself the failure. Documenting all three is necessary because a handover audience needs to understand which mechanism to invoke before they are in the position of needing it.", "Each protection layer in Table 17 answers a different failure scenario. Synchronization supports content continuity under normal operating conditions. Local Veeam retention covers rapid recovery without inter-site dependency. The Site 1 copy addresses the scenario where Site 2 storage is itself the failure. Documenting all three is necessary because a receiving team needs to understand which mechanism to invoke before an outage forces that decision.")
    insert_after(recovery_para, "A realistic recovery sequence makes the design easier to understand. If C2FS failed, the first response would be to confirm whether the fault sits in the mounted storage path or in the host itself, check the most recent good backup chain on S2Veeam, and decide whether the correct response is a file-level restore or a full host recovery. If Site 2 storage were the failure boundary, the Site 1 copy becomes the recovery anchor instead. We kept those layers distinct in the report because operators make different decisions in the first ten minutes of a storage incident than they do in a host-loss or site-loss scenario.", style_name(doc, "Body Text"))
    replace_paragraph_text(doc, "Requirement Coverage", "Coverage Matrix")
    replace_paragraph_text(doc, "This view improves the handover quality because it explains not only where services live, but which identities are meant to operate them.", "This view matters because it explains not only where services live, but which identities are meant to operate them.")
    replace_paragraph_text(doc, "The data-protection model in Site 2 should be read as three different but related mechanisms: synchronized file content, SAN-backed storage delivery, and Veeam-based backup with offsite copy. Keeping those flows distinct is important because each solves a different operational problem and each requires different evidence.", "Site 2 uses three related data-protection mechanisms: synchronized file content, SAN-backed storage delivery, and Veeam-based backup with offsite copy. They belong together, but they do not solve the same problem and should not be treated as interchangeable.")
    replace_paragraph_text(doc, "For a handover document, that distinction is critical because it determines which mechanism would be used under different failure conditions. A sync issue, a share-permission issue, a mounted-volume issue, and a repository or offsite-copy issue all belong to different layers even if they are all loosely described as data problems.", "That distinction determines which mechanism would be used under different failure conditions. A sync issue, a share-permission issue, a mounted-volume issue, and a repository or offsite-copy issue all belong to different layers even if they are all loosely described as data problems.")


def update_conclusion_and_appendices(doc: Document) -> None:
    conclusion_heading = find_exact(doc, "4. Conclusion")
    limitations_heading = insert_after(prev_nonempty(doc, conclusion_heading), "3.15 Limitations and Outstanding Items", "Heading 2")
    p1 = insert_after(limitations_heading, "The final validation set is strong, but it is not uniform across every component. C1WindowsClient remained part of the recorded inventory and supporting evidence set, yet it was not re-inspected through the same live CLI path used for the Linux systems in the final pass. Its place in the environment is therefore supported more by inventory, backup, and network-path evidence than by fresh interactive observation.", style_name(doc, "Body Text"))
    p2 = insert_after(p1, "OPNsense management reachability was confirmed only to the point of a web response. The approved path returned HTTP 403, which is enough to show that the management plane is present and protected, but not enough to document a full authenticated GUI session in this revision.", style_name(doc, "Body Text"))
    insert_after(p2, "C1SAN is also less directly observed than C2SAN. Addressing evidence and design records confirm its isolated interface and its intended relationship to the Company 1 file layer, but no live Company 1 iSCSI session was captured during this final pass. We chose to leave those differences explicit instead of overstating certainty, because a handover record is more useful when it distinguishes direct observation from supported inference.", style_name(doc, "Body Text"))
    first_para = next_nonempty(doc, conclusion_heading)
    first_para.text = "Taken together, the design inputs, environment evidence, vendor references, and operating observations describe Site 2 as a complete operating environment supporting MSP, Company 1, and Company 2 responsibilities inside one managed site. The management boundary, segmented network, dual-tenant identity services, isolated storage paths, hostname-based web delivery, and Veeam protection layer all point to the same operating model instead of competing with one another."
    to_remove = []
    current = next_nonempty(doc, first_para)
    while current.text.strip() != "5. Appendices":
        to_remove.append(current)
        current = next_nonempty(doc, current)
    second = insert_after(first_para, "What the document has tried to show is that Site 2 can be understood in layers without losing the larger picture. MSP entry and policy control establish the management boundary, Company 1 and Company 2 provide the tenant-facing workloads, the SAN and file layers define how data is transported and presented, and the backup design extends the site beyond live service delivery into recoverability.", style_name(doc, "Body Text"))
    insert_after(second, "We learned most clearly that some problems only appear when the environment is read end to end. DNS resolver scope looked like a client detail until it affected web reachability. Synchronization looked like backup until the recovery workflow had to be explained honestly. Isolated SAN design looked straightforward until it became clear how much clearer file-service troubleshooting becomes when block transport stays off the tenant networks. Those lessons are part of the value of the project, because they turned the final report into more than an inventory of machines.", style_name(doc, "Body Text"))
    for para in to_remove:
        remove_paragraph(para)
    appendix_c = find_exact(doc, "Table C1. Service verification and assurance matrix")
    appendix_d = insert_after(appendix_c, "Appendix D. Unresolved Items and Known Gaps", "Heading 2")
    d1 = insert_after(appendix_d, "This appendix records the parts of Site 2 that were supported by indirect evidence or partial live validation at the time of writing. It is included so that a receiving team can distinguish fully rechecked behavior from components that were supported by configuration records, inventory evidence, or limited management-path responses.", style_name(doc, "Body Text"))
    d2 = insert_after(d1, "C1WindowsClient remained part of the recorded inventory and backup scope, but its interactive workflow was not revalidated in the final pass.", style_name(doc, "Normal"))
    d3 = insert_after(d2, "OPNsense management returned HTTP 403 from the approved path, confirming reachability without documenting a fully authenticated GUI session.", style_name(doc, "Normal"))
    d4 = insert_after(d3, "C1SAN was confirmed through addressing evidence and isolation design, but a live Company 1 iSCSI session was not captured in the final pass.", style_name(doc, "Normal"))
    appendix_e = insert_after(d4, "Appendix E. Sanitized SMB Configuration Excerpt", "Heading 2")
    intro = insert_after(appendix_e, "The following excerpt from the C2FS Samba configuration shows the active share model used for public and per-user private access.", style_name(doc, "Body Text"))
    last = intro
    for line in SMB_CONF.read_text(encoding="utf-8").splitlines():
        last = add_monospace_paragraph(last, line, doc)


def ensure_banned_phrases_removed(doc: Document) -> None:
    banned = ["in a handover context", "that is the more supportable model", "it is important to note that", "from a documentation standpoint", "that is why later sections", "the environment reads coherently from end to end"]
    hits = []
    for para in doc.paragraphs:
        low = para.text.lower()
        for phrase in banned:
            if phrase in low:
                hits.append(phrase)
    if hits:
        raise ValueError(f"Banned phrases still present: {sorted(set(hits))}")


def main() -> None:
    doc = Document(str(SRC))
    apply_cover_page_updates(doc)
    apply_front_matter_updates(doc)
    rewrite_executive_summary(doc)
    update_introduction_and_background(doc)
    update_discussion(doc)
    update_conclusion_and_appendices(doc)
    ensure_banned_phrases_removed(doc)
    doc.save(str(DST))
    print(DST)


if __name__ == "__main__":
    main()


