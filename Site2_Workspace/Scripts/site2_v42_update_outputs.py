from copy import copy
from pathlib import Path
import shutil

from docx import Document
from openpyxl import load_workbook

ROOT = Path(r"C:\Algonquin\Winter2026\Emerging_Tech\Project")
SRC_DOC = ROOT / "Site2_SourceBased_Technical_Report_V4.0.docx"
DST_DOC = ROOT / "Site2_SourceBased_Technical_Report_V4.2.docx"
SRC_XLSX = ROOT / "Site2_Service_Block_Demo_Checklist_V3.1.xlsx"
DST_XLSX = ROOT / "Site2_Service_Block_Demo_Checklist_V3.2.xlsx"
LIVE_DIR = sorted(ROOT.glob("site2_v42_live_2026-03-27_*"))[-1]


def clone_row_style(ws, src_row: int, dst_row: int, max_col: int = 11):
    ws.row_dimensions[dst_row].height = ws.row_dimensions[src_row].height
    for col in range(1, max_col + 1):
        src = ws.cell(src_row, col)
        dst = ws.cell(dst_row, col)
        if src.has_style:
            dst.font = copy(src.font)
            dst.fill = copy(src.fill)
            dst.border = copy(src.border)
            dst.alignment = copy(src.alignment)
            dst.protection = copy(src.protection)
            dst.number_format = src.number_format


def update_doc():
    shutil.copy2(SRC_DOC, DST_DOC)
    doc = Document(str(DST_DOC))
    p = doc.paragraphs

    replacements = {
        176: "current operating observations recorded on March 23, March 24, March 25, and March 27, 2026",
        179: "The observation approach followed a dual-bastion management model. MSPUbuntuJump remained the primary Linux bastion for targeted CLI inspection into Company 1 and Company 2 paths, while Jump64 was used as an active Windows inspection platform for Company 1 infrastructure, Company 1 client validation, and S2Veeam administration. Linux systems were inspected directly through low-impact CLI checks; Windows systems were inspected through PowerShell remoting where available and through WMI- or SMB-backed Windows management methods where WinRM was not available. Where a current service behavior had drifted from the intended hostname-based design, the final operating state included only narrowly scoped corrections followed immediately by revalidation.",
        181: "This method matters because Site 2 is not just a list of hosts. Administrative reachability leads into network control, network control enables name resolution and tenant service delivery, Windows and Linux management paths expose different parts of the operating truth, storage paths support file presentation, and the protection layer sits beyond the live service path. Keeping that chain intact is what makes the rest of the document readable as a true handover narrative.",
        237: "The management-path checks confirmed what the configuration evidence already suggested:",
        238: "Approved remote management access to both MSPUbuntuJump and Jump64 was successful",
        239: "Jump64 then served as the active Windows inspection platform for Company 1 systems and S2Veeam, which closed the earlier Windows-observation gap in the report",
        240: "OPNsense management returned HTTP 403 on port 80, TCP 53 was reachable from MSPUbuntuJump, and TCP 443 timed out from the same path",
        262: "Company 1 DNS records were visible on both Company 2 identity nodes, and Jump64 successfully established active management sessions into the Company 1 Windows estate by the methods each host actually permits",
        263: "C1DC1 and C1DC2 were actively inspected from Jump64 over WinRM, returning domain, forest, service, and replication information",
        264: "C1FS was actively inspected from Jump64 and showed a dedicated F: SharedData volume, named SMB shares, and an active iSCSI session tied to the Company 1 file-service stack",
        265: "C1WebServer was actively inspected from Jump64 with its local administrator context and showed a workgroup-hosted IIS role bound only to https://c1-webserver.c1.local on TCP 443",
        266: "C1WindowsClient was actively inspected from Jump64 through WMI and controlled remote process execution, showing domain membership, Company 1 DNS usage, successful resolution of both web hostnames, and HTTP 200 responses to both internal sites",
        267: "C1UbuntuClient successfully showed Company 1 realm membership, current Linux hardware baseline, resolver state, and successful access to both required web hostnames",
        269: "Company 1 is no longer represented only by reachability evidence. The March 27 inspection pass added direct Windows-side observation through Jump64 for the domain controllers, file server, web server, Windows client, and Veeam path where relevant. The only meaningful difference that remains is management protocol shape: C1DC1, C1DC2, C1FS, and S2Veeam accepted WinRM-based inspection; C1WebServer required local-administrator remoting because it is a workgroup web node; C1WindowsClient required WMI- and SMB-backed inspection because TCP 5985 was not open. Those differences describe how the environment is administered in practice rather than a lack of evidence.",
        271: "Company 1 is best understood as a composed service path rather than a flat host list. The domain controllers provide the naming and identity base, the file server provides structured data access on top of isolated storage transport, the Windows and Linux clients show how those services are consumed in practice, and the web server carries the internal application endpoint through its own local-administrator-managed IIS role. Each role has a distinct technical purpose, but the overall value comes from the way those roles reinforce one another.",
        272: "That composed reading also matters for support ownership. A directory issue, a share issue, a hostname issue, a Windows-client management issue, and a storage-transport issue may all appear to an end user as general service instability, yet they belong to different technical layers. Writing Company 1 in layered form keeps the eventual support response aligned to the actual architecture.",
        273: "From an end-to-end service perspective, the Company 1 path can be read in a simple sequence: domain services establish identity and namespace, clients consume those names, the web tier responds only under the intended hostname, the file tier exposes structured access from a dedicated data volume, and the SAN tier remains beneath that file tier as a transport dependency rather than a user-facing service. That sequence explains not only what exists, but why the Company 1 layer behaves predictably when approached from MSP administration, from Jump64, or from Company 1 endpoints.",
        335: "Three different client perspectives were available:",
        336: "C1UbuntuClient, identified in the environment evidence as the Company 1 Linux client role",
        337: "C1WindowsClient, the Company 1 Windows client observed from Jump64 through WMI and controlled remote process execution",
        338: "C2LinuxClient, the Company 2 Linux client",
        339: "All three client perspectives were able to resolve and consume the required internal web hostnames through name-based access. C1UbuntuClient and C2LinuxClient did so directly from their own shells. C1WindowsClient did so through a Jump64-managed client-side probe because WinRM on the client itself was not available:",
        341: "C1UbuntuClient additionally showed Company 1 realm membership and the current shell identity admin@C1UbuntuClient, C1WindowsClient showed domain membership and Company 1 DNS use from the Windows endpoint perspective, and C2LinuxClient showed valid C2.LOCAL realm membership together with valid domain-user lookups for employee1@c2.local and employee2@c2.local.",
        342: "These client observations are the clearest expression of end-to-end service success. They show that routing, naming, identity, certificate handling, and web publication converge into a predictable user-visible result across both Linux and Windows endpoints instead of remaining separate server-side claims.",
        345: "The Linux client also now presents fully formed home directories for employee1@c2.local and employee2@c2.local, which makes domain-user sessions behave like first-class endpoint sessions rather than partial identity lookups. In parallel, the Windows client path proved that Company 1 DNS and web consumption also hold together on a Windows endpoint even when the management method for that endpoint differs from the Linux workflow. Together with the hostname-based SMB access to c2fs.c2.local, that gives Site 2 a cleaner end-user story: directory users can authenticate, land in the expected home context, resolve the expected service names, browse named file services, and consume the internal web applications entirely through the documented namespace model.",
        353: "Publishing services exclusively under their intended hostnames means that misconfigured or stray clients cannot reach them by raw address alone - a deliberate choice that keeps the web layer aligned with the documented naming model.",
        354: "Architecturally, this aligns the web layer to named service delivery instead of opportunistic address-based exposure. DNS, routing, certificates, and web configuration all point to the same contract: users reach the service by its intended name, not by guessing its address. The March 27 Windows inspection also showed that C1WebServer enforces that rule inside IIS as a dedicated 443 binding on a workgroup-hosted web node, which makes the Company 1 web tier operationally distinct while still matching the same hostname-first design principle.",
        355: "The current nginx evidence on C2WebServer reinforces that interpretation from the Linux side. The host is not simply listening on 443; it is configured to answer the c2-webserver.c2.local virtual host specifically, and it returns 404 when the same service is addressed by raw IP without the expected host header. Read together with the IIS evidence on C1WebServer and the Windows-client curl probe from C1WindowsClient, the two tenant web tiers now present one consistent policy story.",
        382: "S2Veeam was reachable from the approved management path on the following ports, confirming that both Veeam control and agent communication channels are accessible from the bastion:",
        383: "TCP 445 and TCP 9392 were reachable from MSPUbuntuJump, confirming SMB repository and Veeam-console path visibility from the Linux bastion",
        384: "TCP 5985, TCP 10005, and TCP 10006 were also reachable from MSPUbuntuJump, confirming the expected Windows-management and agent-control channels",
        385: "Jump64 then established a live administrative session into S2Veeam using the local .\\Administrator context, which confirms that the backup server is not only reachable at port level but also administratively operable from the Windows bastion",
        386: "Within that live session, VeeamBackupSvc, VeeamBrokerSvc, VeeamDeploySvc, VeeamExplorersRecoverySvc, VeeamFilesysVssSvc, VeeamMountSvc, and VeeamNFSSvc were running",
        387: "Other Veeam-branded services such as cloud- or platform-specific connectors were present but appropriately stopped or disabled, which is consistent with a role-focused Veeam deployment rather than a misconfigured service set",
        390: "Documented this way, the protection layer can be explained at both the architectural and operational levels. Repository placement, workload count, file-share coverage, inter-site copy path, and the verified Windows administrative path into S2Veeam all align to the same recovery narrative, which is the main requirement of the technical report.",
        433: "The final validation set is now substantially more balanced than the earlier draft because Company 1 was re-inspected through Jump64 as an active Windows management platform rather than being described mainly by MSP-side port checks. C1DC1, C1DC2, C1FS, C1WebServer, C1WindowsClient, and S2Veeam were all revisited through live Windows-native management methods that match their actual operating context.",
        434: "The remaining observation limits are narrower and more specific. OPNsense management reachability was confirmed to the point of an HTTP 403 response on port 80 and successful TCP 53 from MSPUbuntuJump, but a full authenticated GUI walkthrough was not repeated in this pass. The Veeam GUI evidence set also still needs refreshed screenshots under the latest login context even though live administrative access and service state were confirmed from Jump64.",
        435: "C1SAN remains intentionally opaque as a directly managed host. The relevant evidence is that MSPUbuntuJump and Jump64 did not receive routed management access to its isolated address, while C1FS showed the expected iSCSI consumption path above that storage layer. For handover purposes, that is the correct distinction: isolation was confirmed, and the consuming file server exposed the dependent storage session.",
        437: "Taken together, the design inputs, environment evidence, vendor references, and operating observations describe Site 2 as a complete operating environment supporting MSP, Company 1, and Company 2 responsibilities inside one managed site. The management boundary, segmented network, dual-tenant identity services, isolated storage paths, hostname-based web delivery, Jump64-centred Windows administration, and Veeam protection layer all point to the same operating model instead of competing with one another.",
        438: "What the document has tried to show is that Site 2 can be understood in layers without losing the larger picture. MSP entry and policy control establish the management boundary, Company 1 and Company 2 provide the tenant-facing workloads, the SAN and file layers define how data is transported and presented, the Windows and Linux bastions expose the site through the right management methods, and the backup design extends the site beyond live service delivery into recoverability.",
        439: "We learned most clearly that some problems only appear when the environment is read end to end and through the management path each host actually expects. DNS resolver scope looked like a client detail until it affected web reachability. Synchronization looked like backup until the recovery workflow had to be explained honestly. Company 1 looked simpler than it really was until Jump64-based inspection showed the difference between domain controllers, a workgroup web node, a Windows client without WinRM, and a file server with isolated storage underneath it. Those lessons are part of the value of the project, because they turned the final report into more than an inventory of machines.",
    }
    for idx, text in replacements.items():
        p[idx].text = text

    # Table updates
    doc.tables[2].cell(3, 0).text = "Jump64"
    doc.tables[2].cell(3, 1).text = "To perform active Windows-side inspection of Company 1 systems and S2Veeam through the approved Windows bastion"

    t8 = doc.tables[8]
    t8.cell(1, 2).text = "Actively inspected from Jump64 over WinRM; domain, forest, DC inventory, and service state returned successfully"
    t8.cell(1, 3).text = "Directory, DNS, Kerberos, and replication roles were verified from the Windows bastion"
    t8.cell(2, 2).text = "Actively inspected from Jump64; F: SharedData volume, named SMB shares, and active iSCSI session were visible"
    t8.cell(2, 3).text = "File-service and storage-consumption paths were verified from the Windows bastion"
    t8.cell(3, 2).text = "Actively inspected from Jump64 with local administrator context; IIS returned a single HTTPS binding for c1-webserver.c1.local on TCP 443"
    t8.cell(3, 3).text = "Company 1 web tier is a workgroup web node that still follows the hostname-only publication model"
    t8.cell(4, 2).text = "Actively inspected from Jump64 through WMI and remote process execution; domain membership, DNS, and dual-web probe all returned successfully"
    t8.cell(4, 3).text = "Client was validated through Windows-native management even though WinRM was not open"
    t8.cell(5, 2).text = "Resolved both required web hostnames, showed Company 1 realm context, and returned current Linux hardware and resolver state"
    t8.cell(5, 3).text = "Linux client complements the Windows client and confirms the same named-service contract"
    t8.cell(6, 2).text = "Direct routed management stayed blocked from MSPUbuntuJump and Jump64, while C1FS showed the corresponding active iSCSI consumer session"
    t8.cell(6, 3).text = "Storage remained isolated from tenant routing while still feeding the file-service layer"

    t13 = doc.tables[13]
    t13.cell(1, 1).text = "OPNsense alias, MSP port checks, and Jump64 local-administrator session"
    t13.cell(1, 2).text = "Backup server exists at 172.30.65.180, is reachable from MSP, and is administratively operable from Jump64"
    t13.cell(2, 1).text = "Current operating observations from MSPUbuntuJump and Jump64"
    t13.cell(2, 2).text = "Management, agent, and Windows administrative paths are healthy from the approved bastion pair"

    t23 = doc.tables[23]
    t23.cell(3, 1).text = "MSP-bastion checks, Jump64 remoting or WMI inspection, hostname review, DNS visibility, and client access checks"
    t23.cell(3, 2).text = "Company 1 directory, file, web, client, and SAN roles were all revalidated, with direct Windows-side observation now covering the previously thin areas"
    t23.cell(8, 1).text = "Veeam host reachability, Jump64 administrative session, repository evidence, job inventory, copy-job inventory, route or rule review, and backup-design evidence"
    t23.cell(8, 2).text = "Protection architecture, local backup handling, file-share backup scope, offsite-copy design, and Windows administrative access to S2Veeam were all consistent with the documented operating model"

    doc.save(str(DST_DOC))


def update_workbook():
    shutil.copy2(SRC_XLSX, DST_XLSX)
    wb = load_workbook(str(DST_XLSX))
    ws = wb["Checklist"]
    ws["A3"] = "Observed Test Result and Success? reflect the latest validated state as of 2026-03-27. Suggested values for Success?: PASS / REVIEW / FAIL."

    new_rows_c1 = [
        ["SB1", "Hotseat 1", "Remote Access", "Jump64 administrative PowerShell baseline", "Jump64", "Show the Windows bastion is not only reachable but usable as the active Company 1 inspection platform.", "hostname; whoami; Get-ComputerInfo", "Administrative PowerShell context returned", "Administrative PowerShell context confirmed on Jump64", "PASS", "This closes the earlier Windows-observation gap."],
        ["SB1", "Hotseat 1", "Identity", "C1DC1 live domain-controller state from Jump64", "C1DC1", "Show the primary Company 1 domain controller can be actively inspected from the Windows bastion.", "Test-WSMan 172.30.65.2; Invoke-Command with Company 1 admin credential", "Domain, forest, DC inventory, and service state returned", "Domain, forest, service, and replication state returned successfully", "PASS", "Active inspection replaced port-only evidence."],
        ["SB1", "Hotseat 1", "Identity", "C1DC2 live domain-controller state from Jump64", "C1DC2", "Show the secondary Company 1 domain controller can be actively inspected from the Windows bastion.", "Test-WSMan 172.30.65.3; Invoke-Command with Company 1 admin credential", "Domain, forest, DC inventory, and service state returned", "Domain, forest, service, and replication state returned successfully", "PASS", "Secondary DC validated through the same Windows path."],
        ["SB1", "Hotseat 1", "File Services", "C1FS shares, data volume, and iSCSI baseline from Jump64", "C1FS", "Show the Company 1 file server presents dedicated data storage and consumes isolated SAN-backed storage.", "Invoke-Command to Get-Volume, Get-SmbShare, and Get-IscsiSession", "Dedicated data volume, named shares, and active iSCSI session visible", "F: SharedData, PublicData/Pub_S2/Priv_S2 shares, and active iSCSI session visible", "PASS", "Confirms Company 1 file-service depth from the Windows bastion."],
        ["SB1", "Hotseat 1", "HTTPS", "C1WebServer IIS binding and local-admin path from Jump64", "C1WebServer", "Show the Company 1 web node is supportable as a local-admin workgroup web server and remains hostname-bound.", "Invoke-Command with C1WEBSERVER\\administrator; Get-Website; Get-WebBinding", "W3SVC/WAS running and single HTTPS binding for c1-webserver.c1.local visible", "IIS service state and 443 hostname binding confirmed", "PASS", "Shows the web node is not domain-joined but still matches the design."],
        ["SB1", "Hotseat 1", "Client Validation", "C1WindowsClient management baseline from Jump64", "C1WindowsClient", "Show the Company 1 Windows client is domain joined and manageable through a Windows-native alternate path when WinRM is closed.", "Test-NetConnection 172.30.65.11 5985/445/3389; WMI query from Jump64", "OS, domain membership, and addressing returned", "Windows 10 Education, c1.local domain membership, and endpoint addressing returned via WMI", "PASS", "Used WMI because WinRM on 5985 was not available."],
        ["SB1", "Hotseat 1", "Client Validation", "C1WindowsClient dual-web probe from Jump64-managed execution", "C1WindowsClient", "Show the Company 1 Windows client resolves both internal hostnames and reaches both websites through name-based access.", "WMI process create; nslookup and curl.exe -k -I on the client", "Both hostnames resolve and both hostname-based web probes return 200", "Both hostnames resolved and both client-side curl probes returned HTTP 200", "PASS", "Captured from the client context without raw-IP shortcuts."],
    ]
    insert_at = 16
    ws.insert_rows(insert_at, amount=len(new_rows_c1))
    for offset, row in enumerate(new_rows_c1):
        row_num = insert_at + offset
        clone_row_style(ws, 15, row_num)
        for col, value in enumerate(row, start=1):
            ws.cell(row_num, col).value = value

    new_rows_veeam = [
        ["SB4", "Hotseat 1", "VEEAM", "S2Veeam live service baseline from Jump64", "S2Veeam", "Show the backup platform is administratively operable from the Windows bastion, not only reachable by port.", "Invoke-Command to 172.30.65.180 with .\\Administrator; Get-Service Veeam*", "Administrative session succeeds and core Veeam services are visible", "Jump64 session succeeded with .\\Administrator; core Veeam services were returned", "PASS", "Complements the MSP port checks with a real Windows administrative path."],
    ]
    # Existing VEEAM rows originally started at 56; after inserting 7 rows they start at 63.
    veeam_insert_at = 64
    ws.insert_rows(veeam_insert_at, amount=len(new_rows_veeam))
    for offset, row in enumerate(new_rows_veeam):
        row_num = veeam_insert_at + offset
        clone_row_style(ws, 63, row_num)
        for col, value in enumerate(row, start=1):
            ws.cell(row_num, col).value = value

    guide = wb["Hotseat Guide"]
    guide["A4"] = "Hotseat 1 focuses on MSP entry, Jump64 Windows administration, Company 1 service paths, DHCP/DNS control points, and Veeam."
    guide["A5"] = "Hotseat 2 focuses on Company 2 identity, hostname-based web delivery, file services, client access, and VM or service baseline."

    wb.save(str(DST_XLSX))


def write_summary():
    summary = f"""# Site 2 Live Inspection Summary - 2026-03-27

Live inspection artifacts for the current pass are stored in:
- {LIVE_DIR}

Key outcomes:
- MSPUbuntuJump and Jump64 were both used as active management vantage points.
- Jump64 provided direct Windows-side inspection into C1DC1, C1DC2, C1FS, C1WebServer, C1WindowsClient, and S2Veeam.
- C1WebServer was confirmed as a workgroup-hosted IIS node managed through its local administrator context.
- C1WindowsClient was confirmed as domain-joined and able to resolve and reach both internal web hostnames through Jump64-managed WMI or SMB-backed remote execution, even though WinRM on 5985 was closed.
- C1FS showed a dedicated F: SharedData volume, Company 1 shares, and an active iSCSI session.
- OPNsense responded on port 80 with HTTP 403 and on port 53 with TCP success from MSPUbuntuJump; port 443 timed out from the same path.
- S2Veeam was reachable from MSP and was also administratively operable from Jump64 using .\\Administrator.
"""
    (ROOT / "Site2_Live_Inspection_Summary_2026-03-27.md").write_text(summary, encoding="utf-8")


def update_word_fields():
    try:
        import win32com.client  # type: ignore

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(DST_DOC))
        doc.TablesOfContents(1).Update()
        for table in doc.TablesOfFigures:
            table.Update()
        doc.Fields.Update()
        doc.Save()
        doc.Close()
        word.Quit()
    except Exception:
        pass


if __name__ == "__main__":
    update_doc()
    update_workbook()
    write_summary()
    update_word_fields()
