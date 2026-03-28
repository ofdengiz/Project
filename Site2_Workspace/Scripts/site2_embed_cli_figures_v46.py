from pathlib import Path
import shutil

from docx import Document
from docx.shared import Inches


ROOT = Path(r"C:\Algonquin\Winter2026\Emerging_Tech\Project")
SRC = ROOT / "Site2_SourceBased_Technical_Report_V4.5.docx"
DST = ROOT / "Site2_SourceBased_Technical_Report_V4.6.docx"
IMGDIR = ROOT / "Site2_CLI_Evidence_2026-03-27_193749"


FIG_MAP = {
    "INSERT UPDATED SCREENSHOT HERE - Figure 6 / C2IdM1 AD-DNS-DHCP evidence": IMGDIR / "Figure06_C2IdM1.png",
    "INSERT UPDATED SCREENSHOT HERE - Figure 7 / C2IdM2 AD-DNS-DHCP evidence": IMGDIR / "Figure07_C2IdM2.png",
    "INSERT UPDATED SCREENSHOT HERE - Figure 8 / Shared-forest and cross-domain DNS evidence": IMGDIR / "Figure08_CrossDomainDNS.png",
    "INSERT UPDATED SCREENSHOT HERE - Figure 9 / C2FS iSCSI and mounted volume evidence": IMGDIR / "Figure09_C2FS_iSCSI_Mount.png",
    "INSERT UPDATED SCREENSHOT HERE - Figure 10 / C2FS SMB shares and sync evidence": IMGDIR / "Figure10_C2FS_Shares_Sync.png",
    "INSERT UPDATED SCREENSHOT HERE - Figure 13 / C1UbuntuClient dual-web evidence": IMGDIR / "Figure13_C1UbuntuClient_DualWeb.png",
    "INSERT UPDATED SCREENSHOT HERE - Figure 15 / S2Veeam repository, jobs, and offsite-copy evidence": IMGDIR / "Jump64_S2Veeam_CLI.png",
}


def main() -> None:
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    for para in doc.paragraphs:
        text = para.text.strip()
        if text in FIG_MAP and FIG_MAP[text].exists():
            para.clear()
            run = para.add_run()
            run.add_picture(str(FIG_MAP[text]), width=Inches(6.5))

    doc.save(str(DST))


if __name__ == "__main__":
    main()
