from pathlib import Path
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(r"C:\Algonquin\Winter2026\Emerging_Tech\Project")
OUT_DOCX = ROOT / "Site2_Final_Documentation_V5.0.docx"
OUT_MD = ROOT / "Site2_Final_Documentation_V5.0.md"
ACTIVE_DOCX = OUT_DOCX


def set_cell_text(cell, text):
    cell.text = ""
    p = cell.paragraphs[0]
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for part in str(text).split("\n"):
        if p.text:
            p.add_run().add_break()
        p.add_run(part)


def shade_cell(cell, fill="D9E2F3"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc, caption_text, headers, rows):
    cap = doc.add_paragraph(style="Caption")
    cap.add_run(caption_text)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        set_cell_text(hdr.cells[idx], header)
        shade_cell(hdr.cells[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph("")
    return table


def add_figure_placeholder(doc, label, title, description):
    cap = doc.add_paragraph(style="Caption")
    cap.add_run(f"{label}. {title}")
    p = doc.add_paragraph(style="Placeholder")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"[FIGURE PLACEHOLDER - {label}: {title}]")
    desc = doc.add_paragraph(style="Normal")
    desc.add_run("Description: ").bold = True
    desc.add_run(description)
    spacer = doc.add_paragraph(style="Placeholder")
    spacer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacer.add_run("[Insert updated figure here]")
    doc.add_paragraph("")


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph(style="CodeBlock")
        p.add_run(line)


def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def insert_paragraph_after(paragraph, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    return new_para


def ensure_style(doc, name, style_type=WD_STYLE_TYPE.PARAGRAPH, base=None):
    styles = doc.styles
    if name in styles:
        return styles[name]
    style = styles.add_style(name, style_type)
    if base:
        style.base_style = styles[base]
    return style


def configure_styles(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Cambria"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria")
    normal.font.size = Pt(11)

    title = doc.styles["Title"]
    title.font.name = "Cambria"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria")
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = ensure_style(doc, "Subtitle", base="Normal")
    subtitle.font.name = "Cambria"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria")
    subtitle.font.size = Pt(13)
    subtitle.font.italic = True
    subtitle.font.color.rgb = RGBColor(68, 68, 68)

    cover_heading = ensure_style(doc, "CoverHeading", base="Normal")
    cover_heading.font.name = "Cambria"
    cover_heading._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria")
    cover_heading.font.size = Pt(14)
    cover_heading.font.bold = True
    cover_heading.font.color.rgb = RGBColor(31, 78, 121)

    for name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = doc.styles[name]
        style.font.name = "Cambria"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria")
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(31, 78, 121) if name == "Heading 1" else RGBColor(47, 84, 150)

    caption = doc.styles["Caption"]
    caption.font.name = "Cambria"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria")
    caption.font.size = Pt(10)
    caption.font.bold = True
    caption.font.italic = False

    placeholder = ensure_style(doc, "Placeholder", base="Normal")
    placeholder.font.name = "Cambria"
    placeholder._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria")
    placeholder.font.size = Pt(10)
    placeholder.font.italic = True
    placeholder.font.color.rgb = RGBColor(96, 96, 96)

    code = ensure_style(doc, "CodeBlock", base="Normal")
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(9)

    front = ensure_style(doc, "FrontMatter", base="Normal")
    front.font.name = "Cambria"
    front._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria")
    front.font.size = Pt(11)


TABLES = {
    "Table 1. Design inputs and evidence basis for Site 2": (
        ["Input or Evidence Source", "Purpose in This Report", "How It Was Used"],
        [
            ["Live inspection from MSPUbuntuJump", "Establish Linux-side management visibility and Company 2 operating state", "Used for service-state confirmation on C2IdM1, C2IdM2, C2FS, C2LinuxClient, cross-site reachability, and bastion-path validation."],
            ["Live inspection from Jump64", "Establish Windows-side management visibility and Company 1 operating state", "Used for C1DC1, C1DC2, C1FS, C1WebServer, C1WindowsClient, and S2Veeam inspection from the approved MSP Windows bastion."],
            ["Gateway configuration evidence", "Anchor interface naming, aliases, NAT exposure, and route intent", "Used to describe OPNsense segmentation, edge publication, OpenVPN pathing, and Veeam copy-port intent."],
            ["Observed addressing and inventory records", "Fix system names, addresses, and role boundaries", "Used to populate system inventory, platform baseline, appendix addressing, and role-mapping sections."],
            ["Vendor documentation", "Explain why deployed service mechanisms are supportable", "Used to support firewall, NAT, Samba AD, SMB, iSCSI, IIS binding, and Veeam repository design choices."],
            ["Operating-state review records dated Mar. 23-27, 2026", "Preserve the latest validated behavior across both bastions", "Used to confirm hostname-only web delivery, DNS visibility, SMB behavior, and backup-path observations."],
        ],
    ),
    "Table 8. Company 1 service summary": (
        ["Company 1 System", "Observed Role", "Key State", "Operational Implication"],
        [
            ["C1DC1", "Primary Company 1 domain controller", "Domain and directory state returned successfully from Jump64.", "Company 1 authentication and naming can be inspected from the approved Windows bastion."],
            ["C1DC2", "Secondary Company 1 domain controller", "Companion domain and directory state returned successfully from Jump64.", "Company 1 identity is not concentrated on a single node."],
            ["C1FS", "Windows file server and iSCSI consumer", "Windows SMB active, F: SharedData volume present, named shares present, active iSCSI initiator session confirmed.", "Company 1 storage is layered: SAN presentation stays separate from share publication."],
            ["C1WebServer", "IIS internal web server", "Workgroup-hosted, single HTTPS binding on c1-webserver.c1.local, raw IP returned 404.", "Web access depends on the documented hostname contract, not direct IP exposure."],
            ["C1WindowsClient", "Company 1 Windows client", "Domain membership confirmed and both internal web hostnames returned HTTP 200 from Jump64-backed inspection.", "Company 1 endpoint behavior matches the intended naming model."],
            ["C1UbuntuClient", "Company 1 Linux client", "C1.LOCAL realm active and both internal web hostnames returned HTTP 200.", "Linux endpoint behavior lines up with the same naming and access expectations."],
            ["C1SAN", "Isolated Company 1 storage bridge", "Direct management was intentionally not exposed from the bastions.", "The correct health signal is the active iSCSI consumer session on C1FS, not interactive access to the SAN appliance itself."],
        ],
    ),
    "Table 9. Company 2 identity, DNS, and DHCP summary": (
        ["Company 2 System", "Observed Role", "Key State", "Operational Implication"],
        [
            ["C2IdM1", "Primary Samba AD, DNS, and DHCP node", "samba-ad-dc active, isc-dhcp-server active, both web hostnames present in DNS.", "Company 2 identity and naming begin on a Linux-hosted Microsoft-compatible stack."],
            ["C2IdM2", "Secondary Samba AD, DNS, and DHCP node", "Mirrors the same active service state and DNS visibility as C2IdM1.", "Company 2 naming and DHCP are not tied to one VM."],
            ["DHCP failover model", "Primary on C2IdM1, secondary on C2IdM2", "Roles were observed on both nodes.", "Address assignment continues to follow a known primary-secondary design."],
            ["Directory principals", "Administrator, admin, employee1, employee2, c2_file_users", "Observed on the Company 2 identity plane.", "Identity, file authorization, and Linux endpoint login use one coordinated naming authority."],
            ["Shared-forest naming view", "Zones hosted include c2.local, c1.local, and _msdcs.c2.local", "Both Company 1 and Company 2 web names were resolvable from the Company 2 identity nodes.", "Cross-domain naming is embedded in the site design, not bolted on later."],
        ],
    ),
    "Table 10. Storage and isolated SAN summary": (
        ["Storage Component", "Address or Path", "Observed State", "What It Supports"],
        [
            ["C1SAN", "172.30.65.186/29 via gateway 172.30.65.185", "Isolated Company 1 storage bridge; direct bastion administration intentionally absent.", "Block presentation to C1FS only."],
            ["C1FS storage consumer", "F: SharedData", "Dedicated data volume and active iSCSI initiator session confirmed.", "Company 1 share publication and Company 1 data access."],
            ["C2SAN", "172.30.65.194/29 via gateway 172.30.65.193", "Isolated Company 2 storage bridge with active consumer session from C2FS.", "Block presentation to C2FS only."],
            ["C2FS mounted volume", "/mnt/c2_public on /dev/sdb", "Mounted and active during the final inspection.", "Public and private share publication plus synchronized data landing zone."],
            ["C2FS public share", "/mnt/c2_public/Public", "Published as C2_Public.", "Shared Company 2 collaboration data."],
            ["C2FS private share", "/mnt/c2_public/Private/%U", "Published as C2_Private.", "Per-user Company 2 home-style storage."],
        ],
    ),
    "Table 11. Client access and identity summary": (
        ["Validation Area", "C1UbuntuClient", "C1WindowsClient", "C2LinuxClient"],
        [
            ["Identity context", "C1.LOCAL realm active in the shell context admin@C1UbuntuClient.", "Company 1 domain membership confirmed from Jump64.", "C2.LOCAL realm visible and employee1@c2.local plus employee2@c2.local resolved through getent passwd."],
            ["Resolver state", "Company 1 client consumed the published hostnames successfully from the tenant side.", "Company 1 DNS resolution supported successful access to both internal hostnames.", "DNS servers 172.30.65.66 and 172.30.65.67 with search domains c1.local and c2.local."],
            ["Client-side service consumption", "Web hostname validation for both c1-webserver.c1.local and c2-webserver.c2.local.", "Dual-hostname HTTPS validation for both internal web services from the Windows endpoint.", "Web hostname validation plus hostname-based SMB access to //c2fs.c2.local/C2_Public and //c2fs.c2.local/C2_Private."],
            ["Administrative reading", "Shows that Company 1 Linux users can consume both published names through the intended namespace.", "Shows that a Windows endpoint in Company 1 sees the same hostname-first web contract as the Linux client.", "Shows that Company 2 Linux users consume identity, web, and SMB services through names rather than addresses."],
        ],
    ),
    "Table 12. Internal web delivery summary": (
        ["Web Service", "Published Name", "Observed Result by Hostname", "Observed Result by Raw IP", "Operational Meaning"],
        [
            ["Company 1 web", "c1-webserver.c1.local", "HTTP/2 200", "HTTP/2 404 on 172.30.65.162", "IIS is deliberately bound to the hostname contract and does not serve the site as an address-based page."],
            ["Company 2 web", "c2-webserver.c2.local", "HTTP/1.1 200 OK", "HTTP/1.1 404 Not Found on 172.30.65.170", "nginx follows the same hostname-first delivery model on the Company 2 side."],
        ],
    ),
    "Table 13. Backup and offsite-protection summary": (
        ["Protection Layer", "Observed Scope", "Current State", "Operational Reading"],
        [
            ["Local backup host", "S2Veeam at 172.30.65.180", "Reachable from MSPUbuntuJump on 445, 9392, 5985, 10005, and 10006; administrative access confirmed from Jump64.", "The backup platform is visible from both approved management paths."],
            ["Local repository", "Site2Veeam on Z:\\Site2AgentBackups", "Present and active.", "Local retention exists independently of the file-service synchronization path."],
            ["Offsite copy target", "\\\\192.168.64.20\\Site2OffsiteFromSite2", "Defined as Site1OffsiteSmbShare.", "A separate failure boundary exists beyond Site 2."],
            ["Backup job families", "Ubuntu_Servers, Windows_Servers, C1_FileShare, C2_FileShare", "Observed in the Veeam configuration set.", "Protection is organized by workload class, not as one undifferentiated backup job."],
            ["Copy jobs", "Site 2 to Site 1 copy jobs", "Present.", "Recovery planning covers the case where Site 2 storage itself becomes the failed component."],
            ["Protected workload count", "10 machines", "Observed in the live Veeam state.", "The environment is being protected as a coherent service site rather than as a subset of systems."],
        ],
    ),
    "Table 14. Requirement-to-implementation traceability matrix": (
        ["Requirement Area", "Implementation", "Live or Configuration Evidence", "Operational Result"],
        [
            ["Company 1 directory services", "C1DC1 and C1DC2 observed from Jump64 through the approved Windows administration path.", "WinRM-based inspection from Jump64 returned domain and directory state successfully.", "Company 1 identity is visible, manageable, and distributed across two controllers."],
            ["Company 1 file services and isolated storage", "C1FS consumes Company 1 SAN storage and publishes named SMB shares from F: SharedData.", "Jump64 inspection confirmed Windows SMB state, the dedicated data volume, named shares, and an active iSCSI initiator session.", "Company 1 file services are online and tied to the isolated SAN design that the site intends."],
            ["Company 1 web delivery", "C1WebServer publishes HTTPS only for c1-webserver.c1.local on TCP 443.", "Jump64 inspection confirmed the IIS binding; hostname returned 200 and raw IP returned 404.", "Company 1 web access follows a hostname contract instead of open address-based serving."],
            ["Company 1 client validation", "C1WindowsClient and C1UbuntuClient were used to validate endpoint naming behavior.", "Jump64-backed Windows checks and direct Ubuntu observation both returned successful access to the two internal web names.", "Company 1 endpoints consume the site through the intended namespace."],
            ["Company 2 identity, DNS, and DHCP", "C2IdM1 and C2IdM2 host Samba AD, DNS, and DHCP with a primary-secondary split.", "Both nodes showed active samba-ad-dc and isc-dhcp-server state and both web names in DNS.", "Company 2 identity services are consistent across both nodes."],
            ["Company 2 file services and isolated storage", "C2FS consumes C2SAN storage, mounts /mnt/c2_public, and publishes C2_Public and C2_Private.", "The final inspection showed the active iSCSI session, mounted volume, share definitions, and successful synchronization.", "Company 2 data services operate on a layered path from block storage to user-facing shares."],
            ["Company 2 web delivery", "C2WebServer publishes HTTPS only for c2-webserver.c2.local.", "Hostname returned 200 and raw IP returned 404.", "Company 2 web delivery follows the same hostname-first contract as Company 1."],
            ["Company 2 client validation", "C2LinuxClient validated realm state, resolver path, SMB access, and dual-hostname web access.", "Direct client checks showed correct DNS servers, search domains, getent results, web responses, and hostname-based SMB consumption.", "Company 2 endpoint behavior matches the documented identity and file-service design."],
        ],
    ),
    "Table 2. Evidence classes used in this report": (
        ["Evidence Class", "Definition", "Typical Examples in Site 2"],
        [
            ["Direct observation", "Result obtained from a live session into the environment during the final inspection window", "Service states on C2IdM1 and C2IdM2, C2FS mount state, Jump64 inspection of Company 1 systems, and live web responses from client systems."],
            ["Configuration evidence", "State inferred from exported gateway configuration or persistent service configuration", "OPNsense aliases, NAT publication, firewall policy, and the Appendix E Samba share excerpt on C2FS."],
            ["Environment record", "Structured inventory or addressing data that fixes role names and endpoints", "System inventory, SAN addressing, jump-host roles, and observed platform baseline records."],
            ["Validated design inference", "Conclusion drawn from direct observation plus configuration evidence", "Shared-forest naming behavior, hostname-only web delivery, and the separation between synchronization, local backup, and offsite copy."],
        ],
    ),
    "Table 3. Observation vantage points": (
        ["Observation Point", "Address and Access", "Scope of Inspection", "Why This Observation Point Matters"],
        [
            ["MSPUbuntuJump", "172.30.65.179, SSH", "Primary Linux-side bastion for OPNsense edge testing, Company 2 identity and file services, Company 1 port-reachability checks, and cross-site hostname validation.", "This was the lowest-friction path into Site 2 and the most direct way to test the management plane exactly as later support staff would use it."],
            ["Jump64", "172.30.65.178, RDP", "Primary Windows-side inspection platform for C1DC1, C1DC2, C1FS, C1WebServer, C1WindowsClient, and S2Veeam.", "This is the approved Windows administration path for the tenant systems that require Microsoft-native tooling and remote management context."],
            ["Company 1 direct client observation", "C1UbuntuClient at 172.30.65.36, SSH", "Company 1 Linux client identity state, resolver behavior, and dual-hostname web access.", "This confirms that Company 1 users consume the site by service name instead of by address."],
            ["Company 2 direct client observation", "C2LinuxClient at 172.30.65.70, SSH", "Company 2 Linux client identity state, resolver path, hostname-based SMB consumption, and dual-hostname web access.", "This confirms that Company 2 clients see the intended namespace and share structure from the endpoint perspective."],
        ],
    ),
    "Table 4. Observed Site 2 systems and service roles": (
        ["Scope", "System", "Address", "Observed Role"],
        [
            ["MSP", "OPNsense", "172.30.65.177", "Site gateway, segmentation point, NAT edge, and inter-site VPN policy anchor."],
            ["MSP", "Jump64", "172.30.65.178", "Windows bastion for administrative inspection into Company 1 systems and S2Veeam."],
            ["MSP", "MSPUbuntuJump", "172.30.65.179", "Linux bastion for Company 2 and cross-site operational checks."],
            ["MSP", "S2Veeam", "172.30.65.180", "Local backup repository host, copy-job source, and management endpoint for protection workflows."],
            ["Company 1", "C1DC1", "172.30.65.2", "Primary Company 1 domain controller and DNS service node."],
            ["Company 1", "C1DC2", "172.30.65.3", "Secondary Company 1 domain controller and DNS service node."],
            ["Company 1", "C1FS", "172.30.65.4", "Windows file server consuming isolated SAN storage and publishing Company 1 data shares."],
            ["Company 1", "C1WindowsClient", "172.30.65.11", "Company 1 Windows endpoint used to validate tenant-side web consumption and endpoint state."],
            ["Company 1", "C1UbuntuClient", "172.30.65.36", "Company 1 Linux endpoint used for realm and web validation."],
            ["Company 1", "C1WebServer", "172.30.65.162", "IIS-based internal web service published only under the c1-webserver hostname."],
            ["Company 1", "C1SAN", "172.30.65.186/29", "Isolated Company 1 storage bridge presented to C1FS through iSCSI."],
            ["Company 2", "C2IdM1", "172.30.65.66", "Primary Samba AD, DNS, and DHCP node for Company 2."],
            ["Company 2", "C2IdM2", "172.30.65.67", "Secondary Samba AD, DNS, and DHCP node for Company 2."],
            ["Company 2", "C2FS", "172.30.65.68", "Company 2 Samba file server consuming isolated SAN storage and synchronization data."],
            ["Company 2", "C2LinuxClient", "172.30.65.70", "Company 2 Linux endpoint used for realm, SMB, and dual-hostname web validation."],
            ["Company 2", "C2WebServer", "172.30.65.170", "nginx-based internal web service published only under the c2-webserver hostname."],
            ["Company 2", "C2SAN", "172.30.65.194/29", "Isolated Company 2 storage bridge presented to C2FS through iSCSI."],
        ],
    ),
    "Table 5. Site 2 logical service inventory and role mapping": (
        ["Service Layer", "Systems", "Operational Role", "Primary Consumer or Dependency"],
        [
            ["Edge routing and segmentation", "OPNsense", "Owns WAN publication, tenant segmentation, DNS reachability policy, and the Site 1 VPN route.", "Every routed service path in Site 2 depends on this layer."],
            ["Approved management entry", "Jump64, MSPUbuntuJump", "Provide the two bastion paths through which the site is administered.", "All later support actions begin here."],
            ["Company 1 identity", "C1DC1, C1DC2", "Provide Company 1 directory and DNS services.", "Company 1 endpoints, C1FS, and Company 1 administrative workflows."],
            ["Company 1 application and file services", "C1FS, C1WebServer", "Publish Company 1 data and Company 1 web content.", "Company 1 endpoints and cross-site hostname validation."],
            ["Company 2 identity", "C2IdM1, C2IdM2", "Provide Samba AD, DNS, and DHCP, while hosting the shared-forest naming view.", "Company 2 endpoints and shared hostname resolution paths."],
            ["Company 2 file services", "C2FS", "Publishes public and per-user shares and receives synchronized content from Site 1.", "Company 2 users and the backup layer."],
            ["Client validation layer", "C1UbuntuClient, C1WindowsClient, C2LinuxClient", "Show how identity, naming, and web access behave from endpoints.", "Support and troubleshooting teams validating user-facing behavior."],
            ["Protection layer", "S2Veeam", "Holds local backups and copy-job paths to Site 1.", "Recovery and retention workflows."],
        ],
    ),
    "Table 5A. Observed platform baseline - all service scopes": (
        ["System", "OS", "vCPU", "Memory", "Storage", "Primary Interface"],
        [
            ["MSPUbuntuJump", "Ubuntu Linux", "Observed management endpoint", "Observed management endpoint", "Not treated as application storage", "172.30.65.179 on MSP segment"],
            ["C1UbuntuClient", "Ubuntu 25.04", "4", "7.3 GiB", "32 GB root disk plus 3.8 GiB swap", "172.30.65.36 on C1LAN"],
            ["C2IdM1", "Ubuntu 22.04.5 LTS", "4", "7.8 GiB", "32 GB system disk, 15 GB root LV", "ens18"],
            ["C2IdM2", "Ubuntu 22.04.5 LTS", "4", "7.8 GiB", "32 GB system disk, 15 GB root LV", "ens18"],
            ["C2FS", "Ubuntu 22.04.5 LTS", "4", "7.8 GiB", "16 GB system disk and 160 GB mounted data disk at /mnt/c2_public", "ens19 service, ens18 storage"],
            ["C2LinuxClient", "Ubuntu 25.04", "4", "7.3 GiB", "32 GB root disk", "ens18"],
            ["C2WebServer", "Ubuntu 22.04.5 LTS", "4", "7.8 GiB", "32 GB system disk, 30 GB root LV", "ens18"],
        ],
    ),
    "Table 6. Site 2 network segments and gateways": (
        ["Segment", "Address Range", "Gateway or Interface", "Purpose"],
        [
            ["WAN", "172.20.64.1/16", "OPNsense WAN", "External connectivity and edge publication source."],
            ["MSP", "172.30.65.177/29", "OPNsense MSP interface", "Management enclave for Jump64, MSPUbuntuJump, S2Veeam, and the gateway itself."],
            ["C1LAN", "172.30.65.1/26", "OPNsense C1LAN", "Company 1 routed LAN for controllers, file services, and clients."],
            ["C1DMZ", "172.30.65.161/29", "OPNsense C1DMZ", "Company 1 web publication segment."],
            ["C2LAN", "172.30.65.65/26", "OPNsense C2LAN", "Company 2 routed LAN for identity nodes, file services, and clients."],
            ["C2DMZ", "172.30.65.169/29", "OPNsense C2DMZ", "Company 2 web publication segment."],
            ["SITE1_OVPN", "Inter-site tunnel", "OpenVPN interface", "Cross-site web access and Veeam copy routing."],
            ["Company 1 storage bridge", "172.30.65.186/29", "Gateway 172.30.65.185", "Direct iSCSI-only storage path between C1SAN and C1FS."],
            ["Company 2 storage bridge", "172.30.65.194/29", "Gateway 172.30.65.193", "Direct iSCSI-only storage path between C2SAN and C2FS."],
        ],
    ),
    "Table 7. OPNsense exposure, routing, and firewall policy summary": (
        ["Control Area", "Observed or Documented State", "Operational Reading"],
        [
            ["WAN NAT publication", "33464 to 172.30.65.178:3389 and 33564 to 172.30.65.179:22", "Only the two bastions are published at the edge, so all administration remains funneled through controlled entry points."],
            ["Tenant egress policy", "C1LAN permitted to C1_GLOBAL, ALL_WEBS, ALL_DNS and blocked to C2_GLOBAL; C2LAN mirrors the same logic in reverse", "Tenant users receive the services they need without flattening the site into one trust zone."],
            ["Cross-site web rules", "C1_REMOTE to 172.30.65.170/32 on HTTP/HTTPS; C2_REMOTE to 172.30.65.162/32 on HTTP/HTTPS", "Cross-site web publication is intentional and hostname-oriented, not broad mutual exposure."],
            ["Veeam copy path", "SITE1_VEEAM to S2_VEEAM on VEEAM_COPY_PORTS with static route 192.168.64.20/32 via the Site 1 OpenVPN gateway", "Offsite protection uses an explicitly bounded path instead of an open-ended inter-site allowance."],
            ["Observed management-plane state", "Port 80 returned HTTP 403, TCP 53 was reachable, TCP 443 timed out", "The gateway control plane is present and filtering access, but the final pass did not include an authenticated GUI session."],
            ["Alias design", "C1_Nets, C2_Nets, C1_REMOTE, C2_REMOTE, ALL_WEBS, ALL_DNS, C1_DCs, C2_DCs, S2_VEEAM, SITE1_VEEAM, VEEAM_COPY_PORTS", "Alias-driven policy reduces rule sprawl and keeps the segmentation story readable."],
        ],
    ),
    "Table 15. Service dependency and failure-domain view": (
        ["Service or Workflow", "Primary Dependencies", "Failure Domain", "First Checks When It Breaks"],
        [
            ["Bastion access", "WAN NAT, OPNsense, Jump64, MSPUbuntuJump", "MSP edge and management segment", "Confirm NAT publication, reachability to 172.30.65.178:3389 and 172.30.65.179:22, and jump-host local service health."],
            ["Company 1 identity", "C1DC1, C1DC2, C1LAN reachability, DNS visibility", "Company 1 controller pair", "Check controller service state, DNS response, and Jump64 management reachability."],
            ["Company 2 identity", "C2IdM1, C2IdM2, DNS and DHCP services", "Company 2 identity pair", "Check samba-ad-dc, isc-dhcp-server, and cross-domain DNS visibility."],
            ["Company 1 file services", "C1FS, C1SAN, iSCSI session, SMB service", "Company 1 file server and isolated SAN path", "Confirm F: SharedData availability, SMB service state, and active iSCSI consumer session."],
            ["Company 2 file services", "C2FS, C2SAN, /mnt/c2_public, smbd", "Company 2 file server and isolated SAN path", "Check iSCSI session, mount point, smbd, share definitions, and sync status."],
            ["Internal web delivery", "C1WebServer or C2WebServer, DNS records, host bindings", "Per-web-server application boundary", "Check name resolution first, then binding or nginx/IIS behavior, then raw-IP fallback behavior."],
            ["Backup and offsite copy", "S2Veeam, local repository, OPNsense VPN path, Site 1 SMB target", "Protection platform and inter-site path", "Confirm Veeam service state, repository reachability, and the defined copy path to Site 1."],
        ],
    ),
    "Table 16. Authentication and authorization model": (
        ["Access Context", "Identity Source", "Authorization Control", "Why It Matters"],
        [
            ["MSP administrators", "Local or approved administrative identities on Jump64, MSPUbuntuJump, and S2Veeam", "Administrative entry is constrained to the published jump systems and the approved backup host.", "Support begins from fixed entry points, so credential scope and path discipline stay readable."],
            ["Company 1 domain users", "c1.local", "Directory access, endpoint logon, and Company 1 service consumption remain within the Company 1 tenant context.", "Company 1 keeps a distinct user context even while cross-tenant hostname visibility exists."],
            ["Company 1 service operators", "C1 local administrator on C1WebServer and Company 1 administrative context on C1 infrastructure", "IIS management is separate from Company 1 domain membership because C1WebServer is workgroup-hosted.", "The web server can be managed without implying that every Company 1 service shares one identity model."],
            ["Company 2 directory users", "c2.local", "Samba AD drives Company 2 logon, DNS-linked identity, and file-share authorization.", "Company 2 identity is authoritative for Company 2 endpoints and file services."],
            ["Company 2 file-share access", "c2_file_users and per-user private-share identity", "C2_Public is group-based and C2_Private is user-bound.", "Authorization lines up with the intended public-versus-private share split."],
            ["Cross-tenant name consumption", "Shared-forest naming visibility without merged tenant administration", "OPNsense segmentation still constrains traffic even though names are visible in both company contexts.", "Name resolution convenience does not flatten tenant security boundaries."],
        ],
    ),
    "Table 17. Storage, backup, and recovery data-flow summary": (
        ["Protection or Data Path", "Source", "Destination", "Operational Purpose"],
        [
            ["Company 1 SAN flow", "C1SAN", "C1FS via iSCSI", "Presents Company 1 block storage to the Windows file server for later share publication."],
            ["Company 2 SAN flow", "C2SAN", "C2FS via iSCSI", "Presents Company 2 block storage to the Linux file server for later SMB publication."],
            ["Company 2 synchronization flow", "Site 1 content source", "/mnt/c2_public on C2FS", "Maintains content continuity from Site 1 into Site 2."],
            ["Local backup flow", "Protected Site 2 workloads", "Site2Veeam local repository", "Provides short-path recovery without immediately depending on the inter-site path."],
            ["Offsite backup-copy flow", "S2Veeam", "Site1OffsiteSmbShare at \\\\192.168.64.20\\Site2OffsiteFromSite2", "Extends recovery beyond the Site 2 failure boundary."],
        ],
    ),
    "Table 18. Operational maintenance checks": (
        ["Routine Check", "Systems", "Expected Healthy State", "Why It Matters"],
        [
            ["Validate bastion reachability", "Jump64 and MSPUbuntuJump", "Both management paths reachable through their approved interfaces.", "Bastion access is the prerequisite for every other check in this table; if either jump host is unreachable, the scope of all subsequent assessments narrows significantly."],
            ["Confirm directory-service health", "C1DC1, C1DC2, C2IdM1, C2IdM2", "Controllers respond and directory services are active.", "Identity failure often shows up first as broken name resolution, logon, or file-share access."],
            ["Confirm hostname resolution", "c1-webserver.c1.local and c2-webserver.c2.local", "Both names resolve from Company 1 and Company 2 clients.", "The web tier and cross-tenant naming story depend on this behavior."],
            ["Confirm file-service state", "C1FS and C2FS", "Shares remain published and storage remains mounted or attached.", "Users notice file-service failure faster than most back-end defects."],
            ["Check iSCSI consumer sessions", "C1FS and C2FS", "Active consumer sessions remain present.", "A missing block session can look like a generic share outage unless checked early."],
            ["Check Veeam services and repositories", "S2Veeam", "Core Veeam services active and local repository reachable.", "Backup confidence depends on active services, not just configured jobs."],
        ],
    ),
    "Table 19. Troubleshooting and fast triage guide": (
        ["Reported Symptom", "First System to Check", "Immediate Follow-Up", "Likely Fault Domain"],
        [
            ["No administrative access into Site 2", "OPNsense edge and bastions", "Confirm NAT publication, jump-host service state, and MSP segment reachability.", "MSP edge, WAN publication, or bastion host."],
            ["Company 1 user cannot reach C1 web", "C1WebServer hostname resolution and IIS binding", "Check c1-webserver.c1.local resolution and verify the hostname binding still owns TCP 443.", "DNS, binding, or Company 1 web host."],
            ["Company 2 user cannot reach C2 web", "C2LinuxClient or C2WebServer naming path", "Confirm c2-webserver.c2.local resolution, then verify nginx still serves only the named host.", "DNS, nginx host rule, or Company 2 web host."],
            ["Company 2 share unavailable", "C2FS", "Check iSCSI session, mounted volume, smbd, and share definitions in that order.", "C2SAN path, C2FS mount, or Samba layer."],
            ["Company 1 file path unavailable", "C1FS", "Check F: SharedData, SMB service, and active iSCSI session.", "C1SAN path or Windows file-service layer."],
            ["Backup copy failing", "S2Veeam and inter-site copy path", "Confirm Veeam services, local repository state, and the Site 1 offsite SMB target.", "Veeam platform, inter-site route, or offsite endpoint."],
        ],
    ),
    "Table 20. Integrated design summary": (
        ["Design Domain", "What Was Implemented", "Why It Holds Together Operationally"],
        [
            ["Management", "Two bastions and a tightly bounded edge", "The environment can be supported without publishing tenant systems directly."],
            ["Company 1", "Dual controllers, dedicated file server, isolated SAN, workgroup-hosted IIS, and two client perspectives", "Company 1 services form a complete tenant stack with identity, storage, web, and endpoint validation."],
            ["Company 2", "Dual Samba AD nodes, dual DNS and DHCP role split, isolated SAN, Linux file server, nginx web, and Linux client validation", "Company 2 services deliver the same completeness through a Linux-centric stack."],
            ["Naming model", "Shared-forest visibility across c1.local and c2.local plus hostname-only web publishing", "Clients reach services through stable names even when the underlying addresses remain an implementation detail."],
            ["Protection", "Local Veeam retention plus Site 1 copy path", "Recovery planning is layered and does not depend on a single storage mechanism."],
        ],
    ),
    "Table A1. Observed addressing, gateways, and endpoints": (
        ["System or Segment", "Address", "Gateway or Path", "Notes"],
        [
            ["OPNsense MSP", "172.30.65.177/29", "Gateway interface", "Management-plane anchor for Site 2."],
            ["Jump64", "172.30.65.178", "RDP via NAT 33464", "Windows bastion."],
            ["MSPUbuntuJump", "172.30.65.179", "SSH via NAT 33564", "Linux bastion."],
            ["S2Veeam", "172.30.65.180", "MSP segment", "Backup and copy platform."],
            ["C1DC1", "172.30.65.2", "C1LAN", "Company 1 controller."],
            ["C1DC2", "172.30.65.3", "C1LAN", "Company 1 controller."],
            ["C1FS", "172.30.65.4", "C1LAN", "Company 1 file server."],
            ["C1WindowsClient", "172.30.65.11", "C1LAN", "Company 1 Windows endpoint."],
            ["C1UbuntuClient", "172.30.65.36", "C1LAN", "Company 1 Linux endpoint."],
            ["C1WebServer", "172.30.65.162", "C1DMZ", "Company 1 IIS host."],
            ["C1SAN", "172.30.65.186/29", "Gateway 172.30.65.185", "Company 1 isolated storage bridge."],
            ["C2IdM1", "172.30.65.66", "C2LAN", "Company 2 identity node."],
            ["C2IdM2", "172.30.65.67", "C2LAN", "Company 2 identity node."],
            ["C2FS", "172.30.65.68 and 172.30.65.195", "C2LAN and storage bridge", "Company 2 file server with dual network roles."],
            ["C2LinuxClient", "172.30.65.70", "C2LAN", "Company 2 Linux endpoint."],
            ["C2WebServer", "172.30.65.170", "C2DMZ", "Company 2 nginx host."],
            ["C2SAN", "172.30.65.194/29", "Gateway 172.30.65.193", "Company 2 isolated storage bridge."],
        ],
    ),
    "Table B1. Evidence and reference traceability": (
        ["Topic", "Primary Evidence Class", "Evidence Anchor", "Why It Was Needed"],
        [
            ["Network segmentation and edge exposure", "Configuration evidence", "Gateway configuration record and observed MSP management-plane behavior", "Explains the intended control boundaries before tenant service details are discussed."],
            ["Company 1 directory and file services", "Direct observation", "Jump64 inspection of C1DC1, C1DC2, C1FS, C1WebServer, and C1WindowsClient", "Provides Windows-side validation for Company 1 services."],
            ["Company 2 identity and file services", "Direct observation", "MSPUbuntuJump inspection of C2IdM1, C2IdM2, C2FS, and C2LinuxClient", "Provides Linux-side validation for Company 2 services."],
            ["Shared namespace behavior", "Validated design inference", "DNS visibility on Company 2 identity nodes plus client-side hostname tests", "Shows that the shared-forest design is operational at the endpoint level."],
            ["Backup and offsite protection", "Direct observation plus configuration evidence", "MSPUbuntuJump port checks, Jump64 service inspection, repository and copy-path records", "Separates live platform reachability from repository and copy-job design evidence."],
            ["Technology justification", "Vendor reference", "OPNsense, Samba, Ubuntu, IIS, and Veeam references", "Anchors the deployed mechanisms in authoritative implementation guidance."],
        ],
    ),
    "Table C1. Service verification and assurance matrix": (
        ["Service Scope", "Validation Method", "Result", "Assurance Reading"],
        [
            ["MSP bastion entry", "Observed from WAN NAT and internal management segment", "Pass", "Both administrative entry points are published and reachable through the intended path."],
            ["Company 1 directory", "Jump64 inspection of C1DC1 and C1DC2", "Pass", "Company 1 identity remains distributed across two controllers."],
            ["Company 1 file services", "Jump64 inspection of C1FS", "Pass", "Company 1 shares and SAN-backed storage were present at delivery."],
            ["Company 1 web delivery", "Hostname and raw-IP checks on C1WebServer", "Pass", "Company 1 web access remains hostname-bound."],
            ["Company 1 client validation", "C1UbuntuClient and C1WindowsClient web checks", "Pass", "Both Company 1 client perspectives saw the expected hostname behavior."],
            ["Company 2 identity", "MSPUbuntuJump checks on C2IdM1 and C2IdM2", "Pass", "Samba AD, DNS, and DHCP were active on both nodes."],
            ["Company 2 file services", "C2FS iSCSI, mount, share, and sync inspection", "Pass", "Company 2 file delivery operated end to end."],
            ["Company 2 client validation", "C2LinuxClient realm, SMB, and web tests", "Pass", "Company 2 client behavior matched the intended identity and share model."],
            ["Backup and offsite copy", "Live service inspection plus repository and copy-path evidence", "Pass with evidence-depth note", "Administrative reachability and core backup state were confirmed; refreshed GUI screenshots remain outstanding."],
            ["Gateway management-plane depth", "Port-level observation from MSPUbuntuJump", "Pass with evidence-depth note", "The gateway plane was present and filtering access, but the final pass did not include an authenticated GUI walkthrough."],
        ],
    ),
}


FIGURES = [
    ("Figure 1", "Site 2 topology and service-role alignment diagram", "Shows all Site 2 service domains - MSP management, Company 1 services, Company 2 services, isolated SAN bridges, inter-site VPN path, and S2Veeam - in one operational map."),
    ("Figure 2", "Site 2 logical service inventory and platform role map", "Maps every Site 2 system to its service role across MSP, Company 1, and Company 2 scopes and visually separates those scopes."),
    ("Figure 3", "OPNsense interfaces, aliases, and limited edge exposure", "Shows the OPNsense interface layout, selected aliases, and the WAN NAT publication of only Jump64 and MSPUbuntuJump."),
    ("Figure 4", "OPNsense OpenVPN and inter-site rule mapping", "Shows the inter-site rule set supporting cross-site web access and Veeam copy transport between Site 1 and Site 2."),
    ("Figure 5", "Company 1 services from the MSP management path", "Shows MSPUbuntuJump reachability checks for C1DC1, C1DC2, C1FS, C1UbuntuClient, and C1WebServer."),
    ("Figure 5A", "Jump64 Windows bastion baseline", "Shows the Windows bastion state and its Site 2 management address, confirming that Jump64 was the active Windows inspection platform."),
    ("Figure 5B", "C1DC1 service-state evidence from Jump64", "Shows C1DC1 service-state output from Jump64 and confirms the active Company 1 directory stack."),
    ("Figure 5C", "C1DC2 service-state evidence from Jump64", "Shows C1DC2 service-state output from Jump64 and reinforces the dual-controller model."),
    ("Figure 5D", "C1FS storage, shares, and iSCSI evidence from Jump64", "Shows the F: SharedData volume, named SMB shares, and active iSCSI session as observed from Jump64."),
    ("Figure 5E", "C1WebServer IIS binding evidence from Jump64", "Shows the workgroup-hosted C1WebServer state and the IIS binding restricted to c1-webserver.c1.local on TCP 443."),
    ("Figure 5F", "C1WindowsClient endpoint and dual-web evidence", "Shows Company 1 endpoint state and successful access to both internal web hostnames from the Windows client perspective."),
    ("Figure 6", "C2IdM1 Active Directory, DNS, and DHCP evidence", "Shows samba-ad-dc active, DHCP active, and DNS query output for both web hostnames on C2IdM1."),
    ("Figure 7", "C2IdM2 Active Directory, DNS, and DHCP evidence", "Shows the same identity, DNS, and DHCP pattern on the secondary Company 2 node."),
    ("Figure 8", "Shared-forest and cross-domain DNS evidence", "Shows Company 1 and Company 2 web namespaces visible within the Company 2 identity plane, confirming that the shared-forest design is operational."),
    ("Figure 9", "C2FS iSCSI-backed storage and mounted volume evidence", "Shows the active iSCSI session to 172.30.65.194:3260 and the mounted /mnt/c2_public volume on C2FS."),
    ("Figure 10", "C2FS SMB share definitions and synchronization evidence", "Shows the C2_Public and C2_Private share definitions together with the successful sync result."),
    ("Figure 11", "C1SAN isolated storage interface evidence", "Shows the C1SAN interface configuration that confirms the isolated Company 1 storage segment address and gateway."),
    ("Figure 12", "C2SAN isolated storage interface evidence", "Shows the C2SAN interface configuration that confirms the isolated Company 2 storage segment address and gateway."),
    ("Figure 13", "C1UbuntuClient Company 1 client - domain context and dual-web evidence", "Shows admin@C1UbuntuClient, C1.LOCAL realm visibility, resolver context, and successful HTTP 200 responses to both web hostnames."),
    ("Figure 14", "C2LinuxClient domain identity and dual-web evidence", "Shows C2.LOCAL realm state, employee1 and employee2 resolution, resolver configuration, and successful HTTP responses to both web hostnames."),
    ("Figure 15", "S2Veeam repository, backup jobs, and offsite-copy evidence", "Shows repository configuration, backup job families, and the copy-job path toward Site 1."),
]


REFS = [
    '[1] OPNsense Documentation, "Firewall Rules," https://docs.opnsense.org/manual/firewall.html, accessed Mar. 24, 2026.',
    '[2] OPNsense Documentation, "Network Address Translation," https://docs.opnsense.org/manual/nat.html, accessed Mar. 24, 2026.',
    '[3] OPNsense Documentation, "Setup SSL VPN Road Warrior," https://docs.opnsense.org/manual/how-tos/sslvpn_client.html, accessed Mar. 24, 2026.',
    '[4] SambaWiki, "Setting up Samba as an Active Directory Domain Controller," https://wiki.samba.org/index.php/Setting_up_Samba_as_an_Active_Directory_Domain_Controller, accessed Mar. 24, 2026.',
    '[5] Ubuntu Server Documentation, "Set up Samba as a file server," https://documentation.ubuntu.com/server/how-to/samba/file-server/, accessed Mar. 24, 2026.',
    '[6] Ubuntu Server Documentation, "iSCSI initiator (or client)," https://documentation.ubuntu.com/server/how-to/storage/iscsi-initiator-or-client/, accessed Mar. 24, 2026.',
    '[7] Microsoft Learn, "binding Element for bindings for site [IIS Settings Schema]," https://learn.microsoft.com/en-us/previous-versions/iis/settings-schema/ms691267(v=vs.90), accessed Mar. 24, 2026.',
    '[8] Veeam Help Center, "Configuring Backup Repositories," https://helpcenter.veeam.com/docs/vbr/userguide/sch_configure_repository.html, accessed Mar. 24, 2026.',
    '[9] Site 2 gateway configuration record, Mar. 23, 2026.',
    '[10] Site 2 environment inventory, SAN addressing record, and namespace design record, Mar. 24, 2026.',
    '[11] Site 2 operating-state review record, Mar. 23-27, 2026.',
]

CONTENTS_ROWS = [
    (0, "Contents", 2),
    (0, "List of Figures", 3),
    (0, "List of Tables", 4),
    (0, "Executive Summary", 5),
    (0, "1. Introduction", 5),
    (0, "2. Background", 5),
    (1, "2.1 Intended Audience and Support Scope", 5),
    (1, "2.2 Design Context and Operating Model", 6),
    (1, "2.3 Evidence Base, Observation Method, and Evidence Classes", 7),
    (0, "3. Discussion", 8),
    (1, "3.1 Environment Overview and Service Boundaries", 8),
    (1, "3.2 Service Inventory and Platform Layout", 10),
    (1, "3.3 MSP Entry, Network Segmentation, Remote Access, and Security", 12),
    (1, "3.4 Company 1 Directory Services, File Services, Web Delivery, and Client Access", 14),
    (2, "3.4.1 Service Overview", 14),
    (2, "3.4.2 Architectural Rationale", 16),
    (2, "3.4.3 Observed Operating State", 16),
    (2, "3.4.4 Service Composition and Operational Reading", 18),
    (1, "3.5 Company 2 Identity Services, DNS, DHCP, and Shared Forest Design", 18),
    (2, "3.5.1 Service Overview", 18),
    (2, "3.5.2 Architectural Rationale", 19),
    (2, "3.5.3 Observed Operating State", 19),
    (2, "3.5.4 Service Composition and Operational Reading", 20),
    (1, "3.6 Storage, File Services, and Isolated SAN Design", 20),
    (2, "3.6.1 Company 2 File-Service State (C2FS)", 20),
    (2, "3.6.2 Company 1 File-Service State (C1FS)", 20),
    (2, "3.6.3 SAN Isolation Model", 22),
    (2, "3.6.4 Share Presentation Model", 22),
    (1, "3.7 Client Access, Identity Validation, and Dual-Hostname Web Delivery", 22),
    (2, "3.7.1 Client Validation Perspectives", 23),
    (2, "3.7.2 Hostname-Based Web Publishing Behavior", 24),
    (1, "3.8 Backup, Recovery, and Offsite Protection", 24),
    (2, "3.8.1 Current Operational State", 25),
    (2, "3.8.2 Recovery Role in the Overall Site Design", 26),
    (1, "3.9 Requirement-to-Implementation Traceability", 26),
    (1, "3.10 Service Dependencies, Failure Domains, and Access Model", 27),
    (1, "3.11 Data Protection Flow", 29),
    (1, "3.12 Maintenance and Routine Checks", 30),
    (1, "3.13 Troubleshooting and Fast Triage Guide", 31),
    (1, "3.14 Integrated Design Summary", 32),
    (1, "3.15 Limitations and Outstanding Items", 33),
    (0, "4. Conclusion", 33),
    (0, "5. Appendices", 34),
    (1, "Appendix A: Observed Addressing, Gateways, and Endpoints", 34),
    (1, "Appendix B: Evidence and Reference Traceability", 34),
    (1, "Appendix C: Service Verification Matrix", 35),
    (1, "Appendix D: Unresolved Items and Known Gaps", 36),
    (1, "Appendix E: Sanitized SMB Configuration Excerpt (C2FS)", 37),
    (0, "6. References", 38),
]

FIGURE_ROWS = [
    ("Figure 1. Site 2 topology and service-role alignment diagram", 8),
    ("Figure 2. Site 2 logical service inventory and platform role map", 12),
    ("Figure 3. OPNsense interfaces, aliases, and limited edge exposure", 14),
    ("Figure 4. OPNsense OpenVPN and inter-site rule mapping", 14),
    ("Figure 5. Company 1 services from the MSP management path", 16),
    ("Figure 5A. Jump64 Windows bastion baseline", 16),
    ("Figure 5B. C1DC1 service-state evidence from Jump64", 16),
    ("Figure 5C. C1DC2 service-state evidence from Jump64", 17),
    ("Figure 5D. C1FS storage, shares, and iSCSI evidence from Jump64", 17),
    ("Figure 5E. C1WebServer IIS binding evidence from Jump64", 17),
    ("Figure 5F. C1WindowsClient endpoint and dual-web evidence", 17),
    ("Figure 6. C2IdM1 Active Directory, DNS, and DHCP evidence", 19),
    ("Figure 7. C2IdM2 Active Directory, DNS, and DHCP evidence", 19),
    ("Figure 8. Shared-forest and cross-domain DNS evidence", 20),
    ("Figure 9. C2FS iSCSI-backed storage and mounted volume evidence", 21),
    ("Figure 10. C2FS SMB share definitions and synchronization evidence", 21),
    ("Figure 11. C1SAN isolated storage interface evidence", 22),
    ("Figure 12. C2SAN isolated storage interface evidence", 22),
    ("Figure 13. C1UbuntuClient Company 1 client - domain context and dual-web evidence", 24),
    ("Figure 14. C2LinuxClient domain identity and dual-web evidence", 24),
    ("Figure 15. S2Veeam repository, backup jobs, and offsite-copy evidence", 25),
]

TABLE_ROWS = [
    ("Table 1. Design inputs and evidence basis for Site 2", 6),
    ("Table 2. Evidence classes used in this report", 7),
    ("Table 3. Observation vantage points", 7),
    ("Table 4. Observed Site 2 systems and service roles", 9),
    ("Table 5. Site 2 logical service inventory and role mapping", 10),
    ("Table 5A. Observed platform baseline - all service scopes", 11),
    ("Table 6. Site 2 network segments and gateways", 12),
    ("Table 7. OPNsense exposure, routing, and firewall policy summary", 13),
    ("Table 8. Company 1 service summary", 15),
    ("Table 9. Company 2 identity, DNS, and DHCP summary", 18),
    ("Table 10. Storage and isolated SAN summary", 21),
    ("Table 11. Client access and identity summary", 22),
    ("Table 12. Internal web delivery summary", 23),
    ("Table 13. Backup and offsite-protection summary", 25),
    ("Table 14. Requirement-to-implementation traceability matrix", 27),
    ("Table 15. Service dependency and failure-domain view", 28),
    ("Table 16. Authentication and authorization model", 29),
    ("Table 17. Storage, backup, and recovery data-flow summary", 30),
    ("Table 18. Operational maintenance checks", 31),
    ("Table 19. Troubleshooting and fast triage guide", 32),
    ("Table 20. Integrated design summary", 33),
    ("Table A1. Observed addressing, gateways, and endpoints", 34),
    ("Table B1. Evidence and reference traceability", 35),
    ("Table C1. Service verification and assurance matrix", 36),
]


def add_cover(doc):
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Site 2 Infrastructure Deployment")

    p = doc.add_paragraph(style="CoverHeading")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Integrated Technical Design, Validation, and Handover Report")

    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Design and Implementation of a Multi-Tenant Service Environment using OPNsense, Samba AD, Isolated SAN Storage, nginx and IIS Web Delivery, and Veeam Backup")

    doc.add_paragraph("")

    cover_rows = [
        ("Document Type", "Formal Technical Design, Validation, and Handover Report"),
        ("Service Scope", "MSP, Company 1, and Company 2 operations"),
        ("Document Version", "5.0"),
        ("Document Date", "March 27, 2026"),
        ("Submission Due Date", "March 26, 2026"),
        ("Intended Audience", "Client IT staff, MSP support teams, and successor operations staff"),
        ("Engineering Contributors", "Bailey Kulla, Elyazid Sidelkheir, Ru Wang, Justin Rosseleve, Yiqin Huang, Omer Deniz"),
        ("Team Name", "Site 2 Team"),
        ("Report Intent", "This document is the formal Site 2 technical handover package. It explains service design, observed operating state, support assumptions, maintenance expectations, and validation references so that routine administration and first-line troubleshooting can continue without separate verbal knowledge transfer."),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for left, right in cover_rows:
        row = table.add_row().cells
        set_cell_text(row[0], left)
        shade_cell(row[0], "EAF2F8")
        set_cell_text(row[1], right)


def add_front_matter_placeholders(doc):
    add_page_break(doc)
    doc.add_paragraph("Contents", style="Heading 1")
    doc.add_paragraph("[[CONTENTS_PLACEHOLDER]]", style="FrontMatter")
    add_page_break(doc)
    doc.add_paragraph("List of Figures", style="Heading 1")
    doc.add_paragraph("[[FIGURES_PLACEHOLDER]]", style="FrontMatter")
    add_page_break(doc)
    doc.add_paragraph("List of Tables", style="Heading 1")
    doc.add_paragraph("[[TABLES_PLACEHOLDER]]", style="FrontMatter")
    add_page_break(doc)


def add_section_intro(doc):
    doc.add_paragraph("Executive Summary", style="Heading 1")
    doc.add_paragraph(
        "Site 2 operates as an integrated service environment that combines MSP administration, Company 1 services, Company 2 services, isolated storage, and both local backup and offsite recovery capability within a single managed design. Administrative access starts at two tightly bounded bastions, then fans into tenant-specific identity, file, web, and backup services through controlled internal paths."
    )
    doc.add_paragraph(
        "The environment uses OPNsense to separate MSP, Company 1, Company 2, DMZ, storage-bridge, and inter-site traffic concerns. Company 1 relies on a Windows-led stack with dual controllers, a SAN-backed file server, a workgroup-hosted IIS web server, and both Windows and Linux client validation. Company 2 relies on a Linux-led stack with dual Samba AD nodes, dual DNS and DHCP roles, a SAN-backed Samba file server, nginx-based web delivery, and Linux client validation."
    )
    doc.add_paragraph(
        "What makes the site supportable is not just that the services exist, but that their boundaries are readable. Identity is distributed across paired nodes. Storage remains isolated from routed user networks. Web publication is tied to documented hostnames. Backup, synchronization, and offsite copy are separate mechanisms with separate operational meanings. A receiving support team can therefore reason about the site as a set of deliberate service contracts instead of as a loose cluster of virtual machines."
    )

    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph(
        "This report records the final Site 2 design, the observed operating state at handover, and the support assumptions that follow from that design. It is written as a technical handover document, not as a build diary, so the emphasis stays on architecture, operational reading, and the checks that matter once the site moves into routine administration."
    )
    doc.add_paragraph(
        "The document covers the MSP management path, Company 1 services, Company 2 services, isolated storage, internal web delivery, backup, and the cross-site protection path to Site 1. It also records the observed endpoint behavior that support staff will use to validate the environment from a user-facing perspective."
    )
    doc.add_paragraph(
        "Credentials, screenshots, and interactive console sessions are treated as separate operational artifacts and are not reproduced here in full. The report instead preserves the service model, the observed state, and the technical reasoning required to operate or troubleshoot the site later."
    )

    doc.add_paragraph("2. Background", style="Heading 1")
    doc.add_paragraph("2.1 Intended Audience and Support Scope", style="Heading 2")
    doc.add_paragraph(
        "This document is written for three readers at once: client administrators inheriting the day-to-day environment, MSP staff responsible for support or future changes, and academic reviewers testing whether the site is coherent as an engineered system. It assumes familiarity with routing, Windows and Linux administration, Active Directory concepts, Samba AD, iSCSI, and routine backup operations. It does not assume prior participation in the build."
    )
    doc.add_paragraph("2.2 Design Context and Operating Model", style="Heading 2")
    doc.add_paragraph(
        "Site 2 is best understood as a managed service site with three tightly related scopes: MSP administration, a Company 1 tenant stack, and a Company 2 tenant stack. The same network perimeter, the same bastion entry model, and the same backup platform support all three scopes, yet the tenant layers remain distinct in their identity services, storage paths, and application roles."
    )
    doc.add_paragraph(
        "That operating model drives the rest of the report. OPNsense owns segmentation and edge exposure. Company 1 and Company 2 each own their service identities and endpoint behavior. The storage bridges stay out of the routed tenant paths. The backup platform sits in MSP space so that recovery remains possible even when one tenant service is impaired."
    )
    add_table(doc, "Table 1. Design inputs and evidence basis for Site 2", *TABLES["Table 1. Design inputs and evidence basis for Site 2"])
    doc.add_paragraph("2.3 Evidence Base, Observation Method, and Evidence Classes", style="Heading 2")
    doc.add_paragraph(
        "No single source explains Site 2 completely. Direct inspection shows what was alive and reachable during the final pass. Gateway and service configuration evidence explain why those results should look the way they do. Environment records fix the names, addresses, and role boundaries so that the report does not drift into approximation."
    )
    doc.add_paragraph(
        "The report therefore distinguishes between direct observation, configuration evidence, environment records, and validated inference. That distinction keeps the handover honest and makes it clearer which findings can be checked live and which ones are architectural facts that support later interpretation."
    )
    add_table(doc, "Table 2. Evidence classes used in this report", *TABLES["Table 2. Evidence classes used in this report"])
    add_table(doc, "Table 3. Observation vantage points", *TABLES["Table 3. Observation vantage points"])


def add_discussion(doc):
    doc.add_paragraph("3. Discussion", style="Heading 1")
    doc.add_paragraph(
        "The discussion is organized from the outside in. It begins with site boundaries and inventory, moves through the MSP control layer, then opens Company 1, Company 2, storage, client behavior, protection, and operational support. Read in that order, the site resolves into one managed environment instead of a disconnected list of machines."
    )

    doc.add_paragraph("3.1 Environment Overview and Service Boundaries", style="Heading 2")
    doc.add_paragraph(
        "Site 2 has a clear physical and logical boundary. OPNsense defines the routed edge. Jump64 and MSPUbuntuJump define the approved administrative entry. Company 1 and Company 2 each bring their own service stack into that boundary, while S2Veeam occupies the MSP layer as the recovery platform that spans both tenants. The storage bridges remain intentionally outside the routed tenant paths even though their effects are visible through the file servers that consume them."
    )
    doc.add_paragraph(
        "This matters in practice because support staff do not troubleshoot abstract diagrams; they troubleshoot where responsibility changes hands. The routed edge, the bastions, the tenant identity systems, the file-service consumers, and the backup platform each represent a distinct decision point. Once those boundaries are clear, later sections can focus on how those layers interact."
    )
    add_figure_placeholder(doc, *FIGURES[0])
    add_table(doc, "Table 4. Observed Site 2 systems and service roles", *TABLES["Table 4. Observed Site 2 systems and service roles"])

    doc.add_paragraph("3.2 Service Inventory and Platform Layout", style="Heading 2")
    doc.add_paragraph(
        "The site inventory is not just a list of hostnames. It is the quickest way to see which systems live in MSP space, which systems belong to Company 1, which belong to Company 2, and which ones bridge between storage, naming, and user-facing access. The inventory is also where the difference between Linux-hosted and Windows-hosted services becomes readable."
    )
    doc.add_paragraph(
        "Table 5 maps the logical service roles. Table 5A then gives the Linux-side platform baseline that underpins those roles. Giving the platform baseline its own table matters because later support decisions depend on remembering which Linux nodes are small identity appliances, which ones carry data disks, and which ones are just endpoints."
    )
    add_table(doc, "Table 5. Site 2 logical service inventory and role mapping", *TABLES["Table 5. Site 2 logical service inventory and role mapping"])
    add_table(doc, "Table 5A. Observed platform baseline - all service scopes", *TABLES["Table 5A. Observed platform baseline - all service scopes"])
    add_figure_placeholder(doc, *FIGURES[1])

    doc.add_paragraph("3.3 MSP Entry, Network Segmentation, Remote Access, and Security", style="Heading 2")
    doc.add_paragraph(
        "The MSP layer is where Site 2 becomes supportable. The gateway publishes exactly two entry points from WAN into the site: RDP to Jump64 and SSH to MSPUbuntuJump. That decision keeps administrative access explicit and makes it possible to reason about every later inspection path from the same starting point."
    )
    doc.add_paragraph(
        "Segmentation then continues inside the site. Company 1 and Company 2 each receive their own LAN and DMZ ranges, and the storage bridges remain off the routed user path. Cross-site exceptions are narrow and legible: one set of rules for cross-site web access and another set of copy ports for Veeam traffic to Site 1. The firewall and alias design reinforces the same segmentation logic at the policy layer."
    )
    add_table(doc, "Table 6. Site 2 network segments and gateways", *TABLES["Table 6. Site 2 network segments and gateways"])
    add_table(doc, "Table 7. OPNsense exposure, routing, and firewall policy summary", *TABLES["Table 7. OPNsense exposure, routing, and firewall policy summary"])
    add_figure_placeholder(doc, *FIGURES[2])
    add_figure_placeholder(doc, *FIGURES[3])
    doc.add_paragraph(
        "With the MSP edge and segmentation model in place, the report can move into the tenant stacks themselves. Company 1 comes first because it mixes Windows controllers, a Windows file server, a workgroup IIS host, and two distinct client perspectives. That combination makes it the clearest transition from network policy into actual service delivery."
    )

    doc.add_paragraph("3.4 Company 1 Directory Services, File Services, Web Delivery, and Client Access", style="Heading 2")
    doc.add_paragraph("3.4.1 Service Overview", style="Heading 3")
    doc.add_paragraph(
        "Company 1 in Site 2 is a complete tenant stack, not a partial extension of the other company. It has two domain controllers, a dedicated Windows file server, a distinct internal web server, a Windows client, a Linux client, and an isolated SAN path that only its file server consumes. Taken together, those systems provide identity, naming, storage-backed data access, internal web delivery, and endpoint behavior from both Microsoft and Linux client perspectives."
    )
    doc.add_paragraph(
        "The service mix is also intentionally mixed-platform. Company 1 identity and file services stay on Windows, but web delivery is separated onto a workgroup-hosted IIS server, and endpoint validation includes both a Windows client and a Linux client. That split makes the tenant easier to support because directory, file, web, and client concerns do not collapse into one host class."
    )
    add_table(doc, "Table 8. Company 1 service summary", *TABLES["Table 8. Company 1 service summary"])
    doc.add_paragraph("3.4.2 Architectural Rationale", style="Heading 3")
    doc.add_paragraph(
        "Two domain controllers distribute authentication and naming responsibility across separate nodes, reducing single-point risk and giving the receiving support team a familiar enterprise baseline to reason from. C1DC1 and C1DC2 follow the same dual-node discipline applied to Company 2 identity, and they are expected to behave consistently across directory, DNS, and Kerberos-path functions. That parallel structure is deliberate because the site treats both tenant identity layers as first-class services."
    )
    doc.add_paragraph(
        "C1FS exists as a dedicated file-service node so that storage consumption, share publication, and user access can be managed without folding file operations into either a controller or the IIS host. C1WebServer remains separate and workgroup-hosted because its job is narrowly defined: publish the Company 1 internal site under a documented hostname and nothing else. The endpoint pair then closes the design loop by showing how users actually consume those services after they leave the server tier."
    )
    doc.add_paragraph("3.4.3 Observed Operating State", style="Heading 3")
    doc.add_paragraph(
        "Inspection from both bastions showed that the Company 1 stack is reachable and internally consistent. MSPUbuntuJump confirmed the reachable service surface from the MSP management path, while Jump64 was used to inspect C1DC1, C1DC2, C1FS, C1WebServer, and C1WindowsClient through the approved Windows management route. C1UbuntuClient then provided a second client-side view from the tenant LAN."
    )
    add_figure_placeholder(doc, *FIGURES[4])
    add_figure_placeholder(doc, *FIGURES[5])
    add_figure_placeholder(doc, *FIGURES[6])
    add_figure_placeholder(doc, *FIGURES[7])
    add_figure_placeholder(doc, *FIGURES[8])
    add_figure_placeholder(doc, *FIGURES[9])
    add_figure_placeholder(doc, *FIGURES[10])
    doc.add_paragraph("3.4.4 Service Composition and Operational Reading", style="Heading 3")
    doc.add_paragraph(
        "Company 1 only reads cleanly when its services are understood as a chain. Directory and DNS start on C1DC1 and C1DC2. File access depends on C1FS, which in turn depends on the isolated SAN path. Web access depends on C1WebServer and the documented hostname contract. Client behavior then shows whether all of those upstream decisions are visible in a form that users can actually consume."
    )
    doc.add_paragraph(
        "Company 1 DNS visibility on the Company 2 identity nodes confirms that the naming relationship between the two tenants is operational, not incidental. That visibility allows both C1UbuntuClient and C2LinuxClient to consume Company 1 hosted names without extra client-side rewriting, and it is one of the clearest signs that the shared-forest design is functioning across both company contexts. The Company 1 section therefore flows naturally into Company 2 identity, because that shared naming view becomes much clearer once the Company 2 controllers are examined directly."
    )

    doc.add_paragraph("3.5 Company 2 Identity Services, DNS, DHCP, and Shared Forest Design", style="Heading 2")
    doc.add_paragraph("3.5.1 Service Overview", style="Heading 3")
    doc.add_paragraph(
        "Company 2 supplies the Linux-hosted identity side of Site 2. C2IdM1 and C2IdM2 jointly provide Samba AD, DNS, and DHCP, and they anchor the shared-forest naming view that lets both company contexts consume the published internal web names. Company 2 also owns the identity group that later authorizes access to the Company 2 file-service layer."
    )
    doc.add_paragraph(
        "This part of the site has a narrower service surface than Company 1, but its role is just as important. When Company 2 identity is healthy, clients log in cleanly, DHCP remains predictable, both namespaces resolve, and later file-service authorization continues to work. When it is unhealthy, problems appear across naming, address assignment, and user-facing access at the same time."
    )
    add_table(doc, "Table 9. Company 2 identity, DNS, and DHCP summary", *TABLES["Table 9. Company 2 identity, DNS, and DHCP summary"])
    doc.add_paragraph("3.5.2 Architectural Rationale", style="Heading 3")
    doc.add_paragraph(
        "The Company 2 identity design follows the same paired-node discipline used on the Company 1 side. C2IdM1 and C2IdM2 divide risk across two systems so that authentication, naming, and DHCP do not converge on a single VM. That choice makes the Linux tenant stack easier to defend operationally because every future support engineer can immediately see where redundancy begins and where it ends."
    )
    doc.add_paragraph(
        "The shared-forest naming model is the second major design choice here. Rather than keeping c1.local and c2.local in totally isolated name spaces, the environment uses Company 2 identity nodes to host visibility for both zones while OPNsense continues to enforce the tenant traffic boundaries. This gives the site cross-tenant name resolution without turning cross-tenant access into an unrestricted flat network."
    )
    doc.add_paragraph("3.5.3 Observed Operating State", style="Heading 3")
    doc.add_paragraph(
        "Direct inspection on both Company 2 identity nodes showed the same operational shape: samba-ad-dc active, isc-dhcp-server active, and the published Company 1 and Company 2 web names present in DNS. The paired nodes therefore looked like companions in one service rather than unrelated machines with similar software installed."
    )
    add_figure_placeholder(doc, *FIGURES[11])
    add_figure_placeholder(doc, *FIGURES[12])
    add_figure_placeholder(doc, *FIGURES[13])
    doc.add_paragraph("3.5.4 Service Composition and Operational Reading", style="Heading 3")
    doc.add_paragraph(
        "Company 2 identity has to be read at two levels. At the service level, it is the usual mix of directory, DNS, and DHCP. At the site level, it is also where the cross-tenant naming story becomes concrete, because those nodes can see both the Company 1 and Company 2 web namespaces. That is what lets later client sections show both hostnames working without special per-client hacks."
    )
    doc.add_paragraph(
        "From a Company 1 perspective, the shared forest means that directory principals, Kerberos tickets, and DNS delegation all operate within one naming boundary while OPNsense still constrains traffic between the tenant networks. The result is a design feature, not a coincidence. Once the identity and naming layers are clear, the next question is how those names map to real file-service and storage paths, which is the focus of the next section."
    )
    doc.add_paragraph("3.6 Storage, File Services, and Isolated SAN Design", style="Heading 2")
    doc.add_paragraph(
        "Storage is where the site stops looking like a simple VM inventory and starts looking like an engineered service environment. Both tenants consume isolated SAN-backed storage, but the consumption models are intentionally different at the file-service layer. Company 1 uses a Windows file server and data volume model. Company 2 uses a Linux file server and path-based Samba publication model. The isolation story only holds if those storage paths stay outside the routed tenant LAN and DMZ networks."
    )
    doc.add_paragraph("3.6.1 Company 2 File-Service State (C2FS)", style="Heading 3")
    for bullet in [
        "smbd was active during the final inspection window.",
        "/dev/sdb was mounted at /mnt/c2_public.",
        "An active iSCSI session was present to 172.30.65.194:3260 using target iqn.2024-03.org.clearroots:c2san.",
        "C2_Public mapped to /mnt/c2_public/Public and C2_Private mapped to /mnt/c2_public/Private/%U.",
        "The synchronization result from Site 1 to Site 2 was successful.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")
    doc.add_paragraph("3.6.2 Company 1 File-Service State (C1FS)", style="Heading 3")
    for bullet in [
        "Windows SMB service state on C1FS was active during the Jump64 inspection.",
        "A dedicated F: data volume labeled SharedData was present on the file server.",
        "Named SMB shares were present on the F: SharedData volume.",
        "An active iSCSI initiator session tied C1FS to the Company 1 SAN path.",
        "The Windows file-service layer therefore showed both its data path and its user-facing publication path at the same time.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")
    add_table(doc, "Table 10. Storage and isolated SAN summary", *TABLES["Table 10. Storage and isolated SAN summary"])
    add_figure_placeholder(doc, *FIGURES[14])
    add_figure_placeholder(doc, *FIGURES[15])
    doc.add_paragraph("3.6.3 SAN Isolation Model", style="Heading 3")
    doc.add_paragraph(
        "The storage bridges are isolated on purpose. Bridging each SAN server only to its corresponding file host keeps block traffic entirely off the tenant LAN and DMZ segments. The iSCSI requirement is satisfied without exposing a storage network to users or introducing block transport into routed paths where it does not belong."
    )
    add_figure_placeholder(doc, *FIGURES[16])
    add_figure_placeholder(doc, *FIGURES[17])
    doc.add_paragraph("3.6.4 Share Presentation Model", style="Heading 3")
    doc.add_paragraph(
        "Keeping transport, mount, and share publication as distinct layers bounds troubleshooting effectively. A storage transport failure affects the iSCSI session and the mounted device. A share-permission failure stays within Samba or Windows SMB configuration. Operational staff responding to a reported file-access issue can therefore work through the stack in order instead of treating the entire path as a single undifferentiated problem."
    )
    doc.add_paragraph(
        "That separation also explains why C1SAN and C2SAN do not need broad administrative exposure from the bastions. Their job is to present storage to their respective consumers. The consumers - C1FS and C2FS - are the correct operational vantage points because that is where block storage becomes something users and backup systems actually depend on."
    )

    doc.add_paragraph("3.7 Client Access, Identity Validation, and Dual-Hostname Web Delivery", style="Heading 2")
    doc.add_paragraph(
        "Client observation is where the infrastructure either becomes credible or falls apart. Controllers can look healthy and file servers can publish shares, but if endpoints do not resolve names, obtain identity context, and consume the services through the intended contracts, the design remains theoretical. Site 2 therefore uses three client perspectives: C1UbuntuClient, C1WindowsClient, and C2LinuxClient."
    )
    add_table(doc, "Table 11. Client access and identity summary", *TABLES["Table 11. Client access and identity summary"])
    add_table(doc, "Table 12. Internal web delivery summary", *TABLES["Table 12. Internal web delivery summary"])
    doc.add_paragraph("3.7.1 Client Validation Perspectives", style="Heading 3")
    doc.add_paragraph(
        "C1UbuntuClient showed that Company 1 Linux users can remain inside their own realm context while still consuming both published internal web names. C1WindowsClient confirmed that the same user-facing web pattern is visible from a Windows endpoint inside Company 1. C2LinuxClient extended that validation further by combining identity checks, dual-hostname web access, and hostname-based SMB share consumption from the Company 2 side."
    )
    add_figure_placeholder(doc, *FIGURES[18])
    add_figure_placeholder(doc, *FIGURES[19])
    doc.add_paragraph("3.7.2 Hostname-Based Web Publishing Behavior", style="Heading 3")
    doc.add_paragraph(
        "Publishing services exclusively under their intended hostnames means that misconfigured or stray clients cannot reach them by raw address alone. That choice keeps the web layer aligned with the documented naming model and makes DNS part of the application contract instead of an optional convenience."
    )
    doc.add_paragraph(
        "Consistency at the naming layer also simplifies future maintenance. Certificate renewals, address reassignments, or back-end changes can be handled without altering what users and administrators type to reach the services - the documented hostnames absorb those changes beneath the surface. Once client behavior is clear, the final major service domain is recovery: how the site protects itself and how it behaves if storage or workload failure occurs."
    )

    doc.add_paragraph("3.8 Backup, Recovery, and Offsite Protection", style="Heading 2")
    doc.add_paragraph(
        "Backup deserves its own reading because it sits above both tenant stacks while still depending on them. S2Veeam lives in MSP space, so it does not belong to either tenant. At the same time, it protects tenant workloads, holds a local repository, and uses a bounded path toward Site 1 for offsite copy. In operational terms, it is the one system that has to understand the whole site even though it does not own the day-to-day identity of either tenant."
    )
    add_table(doc, "Table 13. Backup and offsite-protection summary", *TABLES["Table 13. Backup and offsite-protection summary"])
    doc.add_paragraph("3.8.1 Current Operational State", style="Heading 3")
    doc.add_paragraph(
        "S2Veeam was reachable from the approved management path on the ports required for repository, control, and agent communication, confirming that the backup platform remains accessible from the bastion. Jump64 inspection also showed the expected Veeam service set active under local administrative control. This combination is important because MSPUbuntuJump and Jump64 see the platform from different operational angles: one as a network and service endpoint, the other as a Windows-managed application host."
    )
    add_figure_placeholder(doc, *FIGURES[20])
    doc.add_paragraph("3.8.2 Recovery Role in the Overall Site Design", style="Heading 3")
    doc.add_paragraph(
        "Each protection layer in Table 13 answers a different failure scenario. Synchronization supports content continuity under normal operating conditions. Local Veeam retention covers rapid recovery without inter-site dependency. The Site 1 copy addresses the case where Site 2 storage is itself the failure. A support team needs all three mechanisms documented because recovery becomes slower and riskier if staff must rediscover which layer solves which problem during an outage."
    )

    doc.add_paragraph("3.9 Requirement-to-Implementation Traceability", style="Heading 2")
    doc.add_paragraph(
        "Tracing each requirement through to implementation and then to live evidence is what gives the report operational credibility. The matrix below is not there to re-describe the build. It is there to show that each required outcome lands on a named system, a visible mechanism, and an observation that later support staff can repeat."
    )
    add_table(doc, "Table 14. Requirement-to-implementation traceability matrix", *TABLES["Table 14. Requirement-to-implementation traceability matrix"])
    doc.add_paragraph(
        "The matrix also exposes where the site is balanced. Company 1 and Company 2 appear at the same granularity, and the MSP layer stays visible as the control plane that makes both tenant stacks supportable. That balance matters because a site like this should be handed over as a complete service model, not as one strong tenant plus one lightly sketched tenant."
    )
    doc.add_paragraph("3.10 Service Dependencies, Failure Domains, and Access Model", style="Heading 2")
    doc.add_paragraph(
        "Dependencies matter most when something fails. A controller outage, a broken mount, or a missing hostname response can look like a generic site-wide incident until the support team knows which boundary to inspect first. This section therefore separates the dependency view from the authentication view: one table shows where services depend on each other, and the other shows which identities are expected to operate them."
    )
    add_table(doc, "Table 15. Service dependency and failure-domain view", *TABLES["Table 15. Service dependency and failure-domain view"])
    add_table(doc, "Table 16. Authentication and authorization model", *TABLES["Table 16. Authentication and authorization model"])
    doc.add_paragraph(
        "Read together, the two tables show why the bastions matter so much. Services live across multiple tenant scopes, but administrative meaning begins at Jump64 and MSPUbuntuJump. If the bastion layer is unavailable, even healthy tenant systems become harder to assess because the approved observation paths disappear first."
    )

    doc.add_paragraph("3.11 Data Protection Flow", style="Heading 2")
    doc.add_paragraph(
        "Data in Site 2 does not move through one protection mechanism. It moves through several, and each one has a different purpose. The storage bridges present block devices. File servers turn that block layer into user-facing share structures. Synchronization moves content into C2FS. Veeam then captures protected backup data into a local repository and extends a copy toward Site 1."
    )
    add_table(doc, "Table 17. Storage, backup, and recovery data-flow summary", *TABLES["Table 17. Storage, backup, and recovery data-flow summary"])
    doc.add_paragraph(
        "Synchronization and backup serve different purposes, and both need to appear in the handover record. If these mechanisms are collapsed into one story, support staff can misread a healthy sync path as proof of recoverable backup or treat a healthy repository as proof that user-facing file services are intact. Site 2 avoids that confusion by documenting the flows separately."
    )

    doc.add_paragraph("3.12 Maintenance and Routine Checks", style="Heading 2")
    doc.add_paragraph(
        "Routine maintenance is where the site either remains understandable or gradually becomes tribal knowledge. The checks below are written in the order a practical operator would use them. Start with access. Confirm identity and naming. Move into file services and storage consumers. Finish with protection. That sequence mirrors how real incidents are usually encountered: a symptom appears at the edge, and only then does the deeper stack get inspected."
    )
    add_table(doc, "Table 18. Operational maintenance checks", *TABLES["Table 18. Operational maintenance checks"])
    doc.add_paragraph(
        "This sequence is also what keeps Site 2 from becoming boring infrastructure prose. The site is not interesting because it contains many services. It is interesting because each routine check passes responsibility from one layer to the next in a way that a real operator can follow without guesswork."
    )

    doc.add_paragraph("3.13 Troubleshooting and Fast Triage Guide", style="Heading 2")
    doc.add_paragraph(
        "The fastest triage starts with the reported symptom, not with the technology someone happens to like most. A user who cannot reach a share does not care whether the root cause lives in DNS, iSCSI, Samba, or a data volume. The triage guide below is organized to turn that user-facing symptom back into the smallest set of systems that actually need to be checked first."
    )
    add_table(doc, "Table 19. Troubleshooting and fast triage guide", *TABLES["Table 19. Troubleshooting and fast triage guide"])
    doc.add_paragraph(
        "That symptom-first approach is especially important in Site 2 because the environment mixes Windows and Linux service stacks. The same outward problem can sit on two very different platforms. Good triage therefore depends less on memorizing commands and more on knowing which layer owns the symptom before the keyboard work starts."
    )

    doc.add_paragraph("3.14 Integrated Design Summary", style="Heading 2")
    doc.add_paragraph(
        "The site makes the most sense when the separate service sections are collapsed back into one picture. Network, identity, storage, file services, web delivery, client behavior, and backup all point to the same operating model: narrow edge exposure, explicit tenant boundaries, hostname-based service delivery, isolated block storage, and layered protection."
    )
    add_table(doc, "Table 20. Integrated design summary", *TABLES["Table 20. Integrated design summary"])
    doc.add_paragraph(
        "This summary does not repeat the detail above. Its purpose is to confirm that the individual sections - network, identity, storage, and backup - describe one coherent site rather than separate unrelated workloads. That coherence is what makes Site 2 handoff-ready instead of merely assembled."
    )

    doc.add_paragraph("3.15 Limitations and Outstanding Items", style="Heading 2")
    doc.add_paragraph(
        "Two evidence-depth gaps remained at the final pass. The gateway management plane was confirmed at the port level, but the last inspection window did not include an authenticated OPNsense GUI walkthrough. The backup platform was administratively reachable and its service state was confirmed, but refreshed Veeam GUI screenshots had not yet been captured in the final revision window."
    )
    doc.add_paragraph(
        "These gaps do not change the design reading of the site. They do, however, affect how polished the final evidence package appears to a reader who expects every major platform to be shown through both command-line and GUI evidence. A third operational boundary also remains intentional: direct bastion access to C1SAN is not part of the design, so its correct evidence source is the active consumer session on C1FS."
    )


def add_conclusion_and_appendices(doc):
    doc.add_paragraph("4. Conclusion", style="Heading 1")
    doc.add_paragraph(
        "Site 2 is operationally coherent because each scope contributes a complete part of the service story. The MSP layer supplies the only approved entry paths, the segmentation point, and the backup platform that protects the rest of the site. Company 1 contributes a full service stack covering directory, file, web, client, and isolated SAN roles, each confirmed through either direct observation or environment evidence. Company 2 contributes an equally complete stack through Samba AD, DNS, DHCP, Linux file services, nginx web delivery, and Linux client validation."
    )
    doc.add_paragraph(
        "What was learned from the final inspection is less about any one technology than about service composition. Mixed Windows and Linux platforms can still read as one site when naming stays consistent, storage stays isolated, and management paths stay explicit. Site 2 reaches that point. It can now be handed over as a supportable environment instead of a set of one-off build notes."
    )

    doc.add_paragraph("5. Appendices", style="Heading 1")
    doc.add_paragraph("Appendix A: Observed Addressing, Gateways, and Endpoints", style="Heading 2")
    doc.add_paragraph(
        "Appendix A gathers the addressing view in one place so that later troubleshooting does not require jumping back through the discussion chapters to find a single IP, gateway, or service endpoint."
    )
    add_table(doc, "Table A1. Observed addressing, gateways, and endpoints", *TABLES["Table A1. Observed addressing, gateways, and endpoints"])

    doc.add_paragraph("Appendix B: Evidence and Reference Traceability", style="Heading 2")
    doc.add_paragraph(
        "This appendix ties the main report topics to the evidence classes and vendor references that support them. It is intended as a quick map for a reviewer who wants to know where a claim came from and what kind of support stands behind it."
    )
    add_table(doc, "Table B1. Evidence and reference traceability", *TABLES["Table B1. Evidence and reference traceability"])

    doc.add_paragraph("Appendix C: Service Verification Matrix", style="Heading 2")
    doc.add_paragraph(
        "Appendix C is the shortest operational summary in the report. It reduces the site to service scope, validation method, result, and assurance reading so that a later operator can quickly see which areas were directly validated and which ones still carry an evidence-depth note."
    )
    add_table(doc, "Table C1. Service verification and assurance matrix", *TABLES["Table C1. Service verification and assurance matrix"])

    doc.add_paragraph("Appendix D: Unresolved Items and Known Gaps", style="Heading 2")
    doc.add_paragraph(
        "1. The final report pass did not include an authenticated OPNsense GUI walkthrough. The management plane was present, port 80 returned HTTP 403, port 53 responded, and port 443 timed out from MSPUbuntuJump, so the gateway was visible but not interactively documented through the GUI."
    )
    doc.add_paragraph(
        "2. Updated Veeam GUI screenshots were still pending in the last revision window. Live administrative reachability and core service state were confirmed, so this is a presentation gap rather than a platform-state gap."
    )
    doc.add_paragraph(
        "3. C1SAN remains intentionally outside the approved bastion management surface. The correct health signal for Company 1 storage is the active iSCSI consumer session on C1FS, not a direct login path to the SAN bridge."
    )

    doc.add_paragraph("Appendix E: Sanitized SMB Configuration Excerpt (C2FS)", style="Heading 2")
    add_code_block(
        doc,
        [
            "[global]",
            "workgroup = C2",
            "realm = C2.LOCAL",
            "security = ADS",
            "server role = member server",
            "kerberos method = secrets and keytab",
            "winbind use default domain = yes",
            "winbind refresh tickets = yes",
            "idmap config * : backend = tdb",
            "idmap config * : range = 3000-7999",
            "template shell = /bin/bash",
            "template homedir = /home/%U",
            "obey pam restrictions = no",
            "log file = /var/log/samba/log.%m",
            "max log size = 1000",
            "logging = file",
            "",
            "[C2_Public]",
            "path = /mnt/c2_public/Public",
            "browseable = yes",
            "read only = no",
            "valid users = @c2_file_users",
            "force group = c2_file_users",
            "create mask = 0770",
            "directory mask = 0770",
            "",
            "[C2_Private]",
            "path = /mnt/c2_public/Private/%U",
            "browseable = no",
            "read only = no",
            "valid users = %U",
            "force group = c2_file_users",
            "create mask = 0700",
            "directory mask = 0700",
        ],
    )

    doc.add_paragraph("6. References", style="Heading 1")
    for ref in REFS:
        doc.add_paragraph(ref, style="Normal")


def write_markdown():
    lines = []
    lines.append("# Site 2 Infrastructure Deployment")
    lines.append("")
    lines.append("## Integrated Technical Design, Validation, and Handover Report")
    lines.append("")
    lines.append("Version 5.0")
    lines.append("")
    lines.append("This markdown source accompanies the DOCX handover report and preserves the same section order, table captions, and figure placeholders.")
    lines.append("")
    for heading in [
        "Executive Summary",
        "1. Introduction",
        "2. Background",
        "2.1 Intended Audience and Support Scope",
        "2.2 Design Context and Operating Model",
        "2.3 Evidence Base, Observation Method, and Evidence Classes",
        "3. Discussion",
        "3.1 Environment Overview and Service Boundaries",
        "3.2 Service Inventory and Platform Layout",
        "3.3 MSP Entry, Network Segmentation, Remote Access, and Security",
        "3.4 Company 1 Directory Services, File Services, Web Delivery, and Client Access",
        "3.5 Company 2 Identity Services, DNS, DHCP, and Shared Forest Design",
        "3.6 Storage, File Services, and Isolated SAN Design",
        "3.7 Client Access, Identity Validation, and Dual-Hostname Web Delivery",
        "3.8 Backup, Recovery, and Offsite Protection",
        "3.9 Requirement-to-Implementation Traceability",
        "3.10 Service Dependencies, Failure Domains, and Access Model",
        "3.11 Data Protection Flow",
        "3.12 Maintenance and Routine Checks",
        "3.13 Troubleshooting and Fast Triage Guide",
        "3.14 Integrated Design Summary",
        "3.15 Limitations and Outstanding Items",
        "4. Conclusion",
        "5. Appendices",
        "6. References",
    ]:
        lines.append(f"- {heading}")
    OUT_MD.write_text("\n".join(lines), encoding="utf-8")


def build_doc():
    from datetime import datetime
    global ACTIVE_DOCX
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_front_matter_placeholders(doc)
    add_section_intro(doc)
    add_discussion(doc)
    add_conclusion_and_appendices(doc)
    candidates = [
        OUT_DOCX,
        ROOT / "Site2_Final_Documentation_V5.0_Rebuilt.docx",
        ROOT / f"Site2_Final_Documentation_V5.0_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
    ]
    last_error = None
    for candidate in candidates:
        try:
            ACTIVE_DOCX = candidate
            doc.save(ACTIVE_DOCX)
            return
        except PermissionError as exc:
            last_error = exc
    raise last_error


def update_front_matter_static():
    doc = Document(ACTIVE_DOCX)

    def add_tabbed_line(paragraph, text, page, indent_level=0):
        paragraph.paragraph_format.left_indent = Inches(0.2 * indent_level)
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.0))
        paragraph.add_run(f"{text}\t{page}")

    markers = {
        "[[CONTENTS_PLACEHOLDER]]": CONTENTS_ROWS,
        "[[FIGURES_PLACEHOLDER]]": FIGURE_ROWS,
        "[[TABLES_PLACEHOLDER]]": TABLE_ROWS,
    }

    idx = 0
    while idx < len(doc.paragraphs):
        para = doc.paragraphs[idx]
        key = para.text.strip()
        if key not in markers:
            idx += 1
            continue
        para.text = ""
        rows = markers[key]
        current = para
        for offset, row in enumerate(rows):
            if key == "[[CONTENTS_PLACEHOLDER]]":
                level, text, page = row
            else:
                level, text, page = 0, row[0], row[1]
            target = current if offset == 0 else insert_paragraph_after(current, "FrontMatter")
            target.style = "FrontMatter"
            add_tabbed_line(target, text, page, level)
            current = target
        idx += len(rows)

    doc.save(ACTIVE_DOCX)


if __name__ == "__main__":
    write_markdown()
    build_doc()
    update_front_matter_static()
