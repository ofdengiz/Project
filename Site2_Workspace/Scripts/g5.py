from pathlib import Path
import re
from typing import List, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"C:\Algonquin\Winter2026\Emerging_Tech\Project")
OUTDIR = ROOT / "output" / "doc"
DOCX_PATH = OUTDIR / "Site2_Technical_Report_V5.0_Revised.docx"
PDF_PATH = OUTDIR / "Site2_Technical_Report_V5.0_Revised.pdf"
TEXT_PATH = OUTDIR / "Site2_Technical_Report_V5.0_Revised_extracted.txt"


TITLE = "Site 2 Infrastructure Deployment - Integrated Technical Design, Validation, and Handover Report"
SUBTITLE = (
    "Design and Implementation of a Multi-Tenant Service Environment using OPNsense, "
    "Samba AD, Isolated SAN Storage, Nginx and IIS Web Delivery, and Veeam Backup"
)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True

    if "Caption" in doc.styles:
        cap = doc.styles["Caption"]
        cap.font.name = "Times New Roman"
        cap.font.size = Pt(10)
        cap.font.italic = True


def add_para(doc: Document, text: str = "", style: str = "Normal", align=None, bold=False):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    if align is not None:
        p.alignment = align
    return p


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_toc(paragraph) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    run = OxmlElement("w:r")
    txt = OxmlElement("w:t")
    txt.text = "Update field in Word if entries do not render automatically."
    run.append(txt)
    fld.append(run)
    paragraph._p.append(fld)


def add_static_toc(doc: Document) -> None:
    entries = [
        ("Executive Summary", 0),
        ("1. Introduction", 0),
        ("2. Background", 0),
        ("2.1 Intended Audience and Support Scope", 1),
        ("2.2 Design Context and Operating Model", 1),
        ("2.3 Evidence Base, Observation Method, and Evidence Classes", 1),
        ("3. Discussion", 0),
        ("3.1 Environment Overview and Service Boundaries", 1),
        ("3.2 Service Inventory and Platform Layout", 1),
        ("3.3 MSP Entry, Network Segmentation, Remote Access, and Security", 1),
        ("3.4 Company 1 Directory Services, File Services, Web Delivery, and Client Access", 1),
        ("Service Overview", 2),
        ("Architectural Rationale", 2),
        ("Observed Operating State", 2),
        ("Service Composition and Operational Reading", 2),
        ("3.5 Company 2 Identity Services, DNS, DHCP, and Shared Forest Design", 1),
        ("Service Overview", 2),
        ("Architectural Rationale", 2),
        ("Observed Operating State", 2),
        ("Service Composition and Operational Reading", 2),
        ("3.6 Storage, File Services, and Isolated SAN Design", 1),
        ("Company 2 File-Service State (C2FS)", 2),
        ("Company 1 File-Service State (C1FS)", 2),
        ("SAN Isolation Model", 2),
        ("Share Presentation Model", 2),
        ("3.7 Client Access, Identity Validation, and Dual-Hostname Web Delivery", 1),
        ("3.8 Backup, Recovery, and Offsite Protection", 1),
        ("3.9 Requirement-to-Implementation Traceability", 1),
        ("3.10 Service Dependencies, Failure Domains, and Access Model", 1),
        ("3.11 Data Protection Flow", 1),
        ("3.12 Maintenance and Routine Checks", 1),
        ("3.13 Troubleshooting and Fast Triage Guide", 1),
        ("3.14 Integrated Design Summary", 1),
        ("3.15 Limitations and Outstanding Items", 1),
        ("4. Conclusion", 0),
        ("5. Appendices", 0),
        ("Appendix A: Observed Addressing, Gateways, and Endpoints", 1),
        ("Appendix B: Evidence and Reference Traceability", 1),
        ("Appendix C: Service Verification Matrix", 1),
        ("Appendix D: Unresolved Items and Known Gaps", 1),
        ("Appendix E: Sanitized SMB Configuration Excerpt (C2FS)", 1),
        ("6. References", 0),
    ]
    for text, level in entries:
        p = doc.add_paragraph(style="Normal")
        left = 0.0 if level == 0 else (0.3 if level == 1 else 0.6)
        p.paragraph_format.left_indent = Inches(left)
        p.add_run(text)


def add_table(doc: Document, caption: str, headers, rows) -> None:
    add_para(doc, caption, style="Caption" if "Caption" in doc.styles else "Normal")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    doc.add_paragraph()


def add_bullets(doc: Document, items) -> None:
    for item in items:
        add_para(doc, item, style="List Bullet")


def add_figure(doc: Document, number: str, title: str, description: str) -> None:
    add_para(doc, f"[FIGURE PLACEHOLDER - Figure {number}: {title}]", bold=True)
    add_para(doc, f"Description: {description}")


def sentence_count(texts) -> int:
    text = " ".join(texts)
    parts = [p.strip() for p in re.split(r"(?<=[.!?])\s+", text) if p.strip()]
    return len(parts)


def export_pdf_if_possible() -> str:
    try:
        import win32com.client  # type: ignore

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(DOCX_PATH))
        doc.Fields.Update()
        for toc in doc.TablesOfContents:
            toc.Update()
        doc.ExportAsFixedFormat(str(PDF_PATH), 17)
        doc.Save()
        doc.Close(False)
        word.Quit()
        return "pdf-exported"
    except Exception as exc:
        return f"pdf-export-skipped: {exc}"


FIGURES = [
    ("1", "Site 2 topology and service-role alignment diagram", "Shows all Site 2 service domains - MSP management plane, Company 1 services, Company 2 services, isolated SAN bridges, inter-site VPN path, and S2Veeam. The diagram must make all three scopes visually distinct."),
    ("2", "Site 2 logical service inventory and platform role map", "Maps every Site 2 system to its service role across MSP, Company 1, and Company 2 scopes. Should visually distinguish the three scopes."),
    ("3", "OPNsense interfaces, aliases, and limited edge exposure", "Shows the OPNsense interface layout, selected aliases, and WAN NAT publication of only the two jump-host entry points."),
    ("4", "OPNsense OpenVPN and inter-site rule mapping", "Shows the inter-site rule set supporting cross-site web access and backup-copy transport between Site 1 and Site 2."),
    ("5", "Company 1 services from the MSP management path (MSPUbuntuJump port checks)", "Shows MSPUbuntuJump port-reachability results for C1DC1, C1DC2, C1FS, C1UbuntuClient, and C1WebServer."),
    ("5A", "Jump64 Windows bastion baseline", "Shows Jump64 platform state and its internal Site 2 management address, confirming that the Windows bastion was the active inspection platform."),
    ("5B", "C1DC1 service-state evidence from Jump64", "Shows C1DC1 service-state output from Jump64 WinRM, confirming the active Company 1 directory stack."),
    ("5C", "C1DC2 service-state evidence from Jump64", "Shows C1DC2 service-state output from Jump64, reinforcing the dual-controller model."),
    ("5D", "C1FS storage, shares, and iSCSI evidence from Jump64", "Shows the F: SharedData volume, named SMB shares, and active iSCSI session as observed from Jump64 WinRM."),
    ("5E", "C1WebServer IIS binding evidence from Jump64", "Shows the workgroup-hosted C1WebServer state and IIS binding restricted to c1-webserver.c1.local on TCP 443."),
    ("5F", "C1WindowsClient endpoint and dual-web evidence", "Shows domain membership confirmation and successful access to both internal web hostnames via a WMI-backed probe from Jump64."),
    ("6", "C2IdM1 Active Directory, DNS, and DHCP evidence", "Shows samba-ad-dc active, DHCP active, and DNS query output for both web hostnames on C2IdM1."),
    ("7", "C2IdM2 Active Directory, DNS, and DHCP evidence", "Same structure as Figure 6 but for the secondary identity node, confirming dual-node consistency."),
    ("8", "Shared-forest and cross-domain DNS evidence", "Shows both Company 1 and Company 2 web namespaces visible within the Company 2 identity plane, confirming the shared-forest design is operational."),
    ("9", "C2FS iSCSI-backed storage and mounted volume evidence", "Shows the active iSCSI session to 172.30.65.194:3260 and the /mnt/c2_public mounted volume on C2FS."),
    ("10", "C2FS SMB share definitions and synchronization evidence", "Shows the C2_Public and C2_Private share definitions from testparm output together with the successful sync log result."),
    ("11", "C1SAN isolated storage interface evidence", "Shows the C1SAN interface configuration confirming the isolated Company 1 storage segment address and gateway."),
    ("12", "C2SAN isolated storage interface evidence", "Shows the C2SAN interface configuration confirming the isolated Company 2 storage segment address and gateway."),
    ("13", "C1UbuntuClient Company 1 client - domain context and dual-web evidence", "Shows C1UbuntuClient shell context (admin@C1UbuntuClient), C1.LOCAL realm visibility, resolver state, and successful HTTP 200 responses to both web hostnames."),
    ("14", "C2LinuxClient domain identity and dual-web evidence", "Shows C2.LOCAL realm state, employee1@c2.local and employee2@c2.local via getent passwd, resolver configuration, and HTTP 200 responses to both web hostnames."),
    ("15", "S2Veeam repository, backup jobs, and offsite-copy evidence", "Shows Veeam repository configuration, backup job families (Ubuntu_Servers, Windows_Servers, C1_FileShare, C2_FileShare), and copy-job configuration toward Site 1."),
]


TABLE_TITLES = [
    "Table 1. Design inputs and evidence basis for Site 2",
    "Table 2. Evidence classes used in this report",
    "Table 3. Observation vantage points",
    "Table 4. Observed Site 2 systems and service roles",
    "Table 5. Site 2 logical service inventory and role mapping",
    "Table 5A. Observed platform baseline - all service scopes",
    "Table 6. Site 2 network segments and gateways",
    "Table 7. OPNsense exposure, routing, and firewall policy summary",
    "Table 8. Company 1 service summary",
    "Table 9. Company 2 identity, DNS, and DHCP summary",
    "Table 10. Storage and isolated SAN summary",
    "Table 11. Client access and identity summary",
    "Table 12. Internal web delivery summary",
    "Table 13. Backup and offsite-protection summary",
    "Table 14. Requirement-to-implementation traceability matrix",
    "Table 15. Service dependency and failure-domain view",
    "Table 16. Authentication and authorization model",
    "Table 17. Storage, backup, and recovery data-flow summary",
    "Table 18. Operational maintenance checks",
    "Table 19. Troubleshooting and fast triage guide",
    "Table 20. Integrated design summary",
    "Table A1. Observed addressing, gateways, and endpoints",
    "Table B1. Evidence and reference traceability",
    "Table C1. Service verification and assurance matrix",
]


APPENDIX_E_BLOCK = [
    "[global]",
    "workgroup = C2",
    "realm = C2.LOCAL",
    "security = ADS",
    "server role = member server",
    "kerberos method = secrets and keytab",
    "winbind use default domain = yes",
    "winbind refresh tickets = yes",
    "idmap config * : backend = tdb",
    "idmap config * : range = 3000-7999",
    "template shell = /bin/bash",
    "template homedir = /home/%U",
    "obey pam restrictions = no",
    "log file = /var/log/samba/log.%m",
    "max log size = 1000",
    "logging = file",
    "",
    "[C2_Public]",
    "path = /mnt/c2_public/Public",
    "browseable = yes",
    "read only = no",
    "valid users = @c2_file_users",
    "force group = c2_file_users",
    "create mask = 0770",
    "directory mask = 0770",
    "",
    "[C2_Private]",
    "path = /mnt/c2_public/Private/%U",
    "browseable = no",
    "read only = no",
    "valid users = %U",
    "force group = c2_file_users",
    "create mask = 0700",
    "directory mask = 0700",
]


def build_document() -> Tuple[Document, List[str], List[str]]:
    doc = Document()
    configure_styles(doc)

    add_para(doc, TITLE, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_para(doc, SUBTITLE, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Formal Technical Design, Validation, and Handover Report", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_para(doc, "Service Scope: MSP, Company 1, and Company 2 operations", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Document Version: 5.0", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Document Date: March 27, 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Submission Due Date: March 26, 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Intended Audience: Client IT staff, MSP support teams, and successor operations staff", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Engineering Contributors: Bailey Kulla, Elyazid Sidelkheir, Ru Wang, Justin Rosseleve, Yiqin Huang, Omer Deniz", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Team Name: Site 2 Team", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Report Intent: This document is the formal Site 2 technical handover package. It explains service design, observed operating state, support assumptions, maintenance expectations, and validation references so that routine administration and first-line troubleshooting can continue without separate verbal knowledge transfer.", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_page_break(doc)

    add_heading(doc, "Table of Contents", 1)
    add_static_toc(doc)
    add_page_break(doc)

    add_heading(doc, "List of Figures", 1)
    for number, title, _desc in FIGURES:
        add_para(doc, f"Figure {number}. {title}")
    add_page_break(doc)

    add_heading(doc, "List of Tables", 1)
    for title in TABLE_TITLES:
        add_para(doc, title)
    add_page_break(doc)

    add_heading(doc, "Executive Summary", 1)
    add_para(doc, "Site 2 operates as a managed service site with three equally important scopes: MSP administration, Company 1 tenant services, and Company 2 tenant services. OPNsense at 172.30.65.177/29 enforces the routed boundary, Jump64 at 172.30.65.178 and MSPUbuntuJump at 172.30.65.179 provide controlled entry, and S2Veeam at 172.30.65.180 anchors backup and offsite-copy protection. Inside that management frame, Company 1 and Company 2 each expose their own identity, client, web, and file-service paths without collapsing into one flat network model.")
    add_para(doc, "Inspection showed that the site is operational at the levels that matter for handover. Company 1 controllers, file services, web delivery, and both client perspectives were all observable from the approved bastion paths. Company 2 identity, DNS, DHCP, file services, Linux client behavior, and nginx-based web delivery were directly validated from the Linux and Windows inspection points. The observed outputs consistently matched the documented network segments, hostnames, ports, aliases, and storage paths recorded for the site [9]-[11].")
    add_para(doc, "The report therefore treats Site 2 as a supportable operating environment rather than as a build log. Each section explains what the service layer does, why it was structured that way, what was directly observed, and how a receiving team should read faults when they occur. The final sections keep backup separate from synchronization, keep block storage separate from share presentation, and keep evidence limits explicit where the final pass did not include an authenticated OPNsense GUI walkthrough or refreshed Veeam screenshots.")

    add_heading(doc, "1. Introduction", 1)
    add_para(doc, "This report documents Site 2 as the formal design, validation, and handover package for a multi-tenant managed service environment. It is written for client IT staff who will administer the site, MSP personnel who may troubleshoot or modify it later, and academic assessors who need to see whether the environment was built coherently and defended with honest evidence.")
    add_para(doc, "The discussion is organized around operating layers instead of build chronology. First the report fixes the evidence basis, then it explains the site layout, then it follows the service chains through identity, file services, web publication, clients, backup, and recovery. That sequence lets the reader move from site context to tenant detail without losing the whole-site view that makes later troubleshooting practical.")

    add_heading(doc, "2. Background", 1)
    add_heading(doc, "2.1 Intended Audience and Support Scope", 2)
    add_para(doc, "The primary audience is the receiving operations team for Site 2. They need enough detail to understand approved management entry, tenant service boundaries, storage isolation, backup responsibilities, and the first checks that should be performed when routine faults appear. The secondary audience is the MSP support team, which needs a concise but defensible record of how Jump64, MSPUbuntuJump, OPNsense, and S2Veeam fit together across both company scopes. The third audience is academic assessment, which needs to see not only what was built but also how live observation was separated from configuration evidence and design records.")

    add_heading(doc, "2.2 Design Context and Operating Model", 2)
    add_para(doc, "Site 2 was designed as a controlled service site, not as a loose set of isolated virtual machines. The MSP scope defines entry, policy, and recovery control. Company 1 contributes Windows directory services, Windows file services, a workgroup IIS server, a Windows client, a Linux client, and an isolated SAN bridge. Company 2 contributes Samba AD, DNS, DHCP, a Samba file server with isolated storage, a Linux client, and an nginx web host. Inter-site access is limited to documented web and backup paths over the OpenVPN relationship to Site 1 [1]-[3], [9]-[11].")
    add_table(
        doc,
        "Table 1. Design inputs and evidence basis for Site 2",
        ["Input source", "Observed content", "How it is used in this report"],
        [
            ["Gateway and routing record", "OPNsense interface plan, NAT exposure, aliases, static route to 192.168.64.20/32", "Anchors the network, edge-exposure, and inter-site access discussion [9]."],
            ["Environment inventory", "Authoritative IP addresses, hostnames, interfaces, storage bridges, and roles", "Defines the host-level facts that must not be altered or approximated [10]."],
            ["Operating-state review record", "March 23-27, 2026 service checks from MSPUbuntuJump and Jump64", "Provides the live observation base for service validation [11]."],
            ["Vendor documentation", "OPNsense, Samba, Ubuntu Server, IIS, nginx, Veeam, WMI, and SMB references", "Supports architectural rationale without replacing the site-specific evidence [1]-[8], [12]-[14]."],
        ],
    )

    add_heading(doc, "2.3 Evidence Base, Observation Method, and Evidence Classes", 2)
    add_para(doc, "The report uses a mixed evidence model because not every part of the site was validated in the same way. Some components were observed directly through service-state checks, port probes, WinRM inspection, WMI-based inspection, or client-side access validation. Other components were confirmed through authoritative records or configuration artifacts when direct administration was intentionally restricted. The table set below makes those evidence classes explicit before the service discussion begins.")
    add_table(
        doc,
        "Table 2. Evidence classes used in this report",
        ["Evidence class", "Definition", "Examples in this report"],
        [
            ["Live observation", "Direct result returned during the March 23-27, 2026 inspection window", "Service state on C2IdM1, C2IdM2, C2FS, C1FS, C1DC1, C1DC2, C1WebServer, C1WindowsClient, and S2Veeam [11]."],
            ["Client validation", "Result seen from an endpoint consuming the service", "HTTP 200 from named web hosts on C1UbuntuClient, C1WindowsClient, and C2LinuxClient; hostname-based SMB access on C2LinuxClient [11]."],
            ["Configuration evidence", "State inferred from active configuration artifacts that were directly reviewed", "DNS records, aliases, zones, SMB share definitions, and routed segment layout [9]-[11]."],
            ["Design evidence", "Authoritative site record that describes intended topology or support boundary", "Shared forest interpretation, SAN bridge addressing, and intended offsite-copy path [9]-[11]."],
            ["Corroborated inference", "Operational reading supported by more than one evidence class but not presented as a fresh reconfiguration event", "The shared-forest behavior inferred from cross-domain DNS visibility and client resolution results [10]-[11]."],
        ],
    )
    add_table(
        doc,
        "Table 3. Observation vantage points",
        ["Vantage point", "Address", "Inspection role", "Primary systems or paths covered"],
        [
            ["MSPUbuntuJump", "172.30.65.179", "Primary Linux bastion for CLI inspection and path validation", "Jump64, OPNsense reachability, C1DC1, C1DC2, C1FS, C1UbuntuClient, C1WebServer, C2IdM1, C2IdM2, C2FS, and C2LinuxClient."],
            ["Jump64", "172.30.65.178", "Primary Windows-side inspection platform", "C1DC1, C1DC2, C1FS, C1WebServer, C1WindowsClient, and S2Veeam."],
            ["C2IdM1", "172.30.65.66", "Identity-node observation point", "samba-ad-dc, isc-dhcp-server, DNS records, hosted zones, and principal inventory."],
            ["C2IdM2", "172.30.65.67", "Secondary identity-node observation point", "Dual-node consistency for AD, DNS, DHCP, and shared namespace visibility."],
            ["C2FS", "172.30.65.68", "File-service observation point", "smbd state, /dev/sdb mount at /mnt/c2_public, iSCSI session to 172.30.65.194:3260, and share definitions."],
            ["C2LinuxClient", "172.30.65.70", "Client-side identity and service consumer", "Resolver state, employee principal resolution, hostname-based web access, and hostname-based SMB access."],
        ],
    )
    add_para(doc, "With the evidence model fixed, Section 3 can explain the site itself without implying that every statement comes from the same depth of observation.")

    add_heading(doc, "3. Discussion", 1)
    add_heading(doc, "3.1 Environment Overview and Service Boundaries", 2)
    add_para(doc, "Site 2 is easiest to understand as one operating site with three service scopes. The MSP scope defines how administrators enter, how policy is enforced, and how recovery tooling is reached. Company 1 and Company 2 then consume that boundary to deliver their own tenant services, while the isolated SAN bridges and Site 1 OpenVPN path extend the design below and beyond the ordinary tenant LANs.")
    add_figure(doc, *FIGURES[0])
    add_para(doc, "That boundary model matters because service behavior is intentionally not uniform across every system. Some hosts answer only by hostname, some storage paths are intentionally hidden behind file-service hosts, and some systems are meant to be touched only from one bastion. Reading the site in scopes prevents the receiving team from confusing deliberate isolation with missing functionality.")
    add_table(
        doc,
        "Table 4. Observed Site 2 systems and service roles",
        ["Scope", "System", "Address", "Role", "Observed operating reading"],
        [
            ["MSP", "OPNsense", "172.30.65.177", "Gateway and segmentation control", "Authoritative routed boundary for MSP, C1LAN, C1DMZ, C2LAN, C2DMZ, and SITE1_OVPN."],
            ["MSP", "Jump64", "172.30.65.178", "Windows bastion", "Primary Windows inspection path for Company 1 services and S2Veeam."],
            ["MSP", "MSPUbuntuJump", "172.30.65.179", "Linux bastion", "Primary Linux inspection path into Site 2 service networks."],
            ["MSP", "S2Veeam", "172.30.65.180", "Backup and recovery host", "Observed repository, job families, copy jobs, and active Veeam services."],
            ["Company 1", "C1DC1", "172.30.65.2", "Primary domain controller", "Observed from Jump64 with active directory-service state returned successfully."],
            ["Company 1", "C1DC2", "172.30.65.3", "Secondary domain controller", "Observed from Jump64 with matching controller and directory context."],
            ["Company 1", "C1FS", "172.30.65.4", "Windows file server", "Observed from Jump64 with SMB active, F: SharedData present, and iSCSI session confirmed."],
            ["Company 1", "C1WindowsClient", "172.30.65.11", "Windows client endpoint", "Observed from Jump64 as a Company 1 consumer of directory and web services."],
            ["Company 1", "C1UbuntuClient", "172.30.65.36", "Linux client endpoint", "Observed from MSPUbuntuJump with C1.LOCAL realm and dual-web HTTP 200 results."],
            ["Company 1", "C1WebServer", "172.30.65.162", "IIS internal web server", "Observed with hostname-only HTTPS publication and HTTP 404 on raw IP."],
            ["Company 1", "C1SAN", "172.30.65.186/29", "Isolated Company 1 storage bridge", "Consumed indirectly through the live iSCSI session on C1FS."],
            ["Company 2", "C2IdM1", "172.30.65.66", "Primary Samba AD, DNS, DHCP", "Observed with samba-ad-dc and isc-dhcp-server active."],
            ["Company 2", "C2IdM2", "172.30.65.67", "Secondary Samba AD, DNS, DHCP", "Observed with matching service state and shared-zone visibility."],
            ["Company 2", "C2FS", "172.30.65.68", "Samba file server", "Observed with mounted iSCSI-backed storage and active shares."],
            ["Company 2", "C2LinuxClient", "172.30.65.70", "Linux client endpoint", "Observed with resolver state, employee principals, and hostname-based service access."],
            ["Company 2", "C2WebServer", "172.30.65.170", "nginx HTTPS web server", "Observed with hostname success and raw-IP HTTP 404 behavior."],
            ["Company 2", "C2SAN", "172.30.65.194/29", "Isolated Company 2 storage bridge", "Consumed directly by C2FS over the storage NIC."],
        ],
    )
    add_para(doc, "The next section turns that site map into a role map so the relationship between hosts, platforms, and support ownership is explicit before the network policy details are introduced.")

    add_heading(doc, "3.2 Service Inventory and Platform Layout", 2)
    add_para(doc, "The inventory view shows where each service responsibility lives and which platform was chosen to carry it. That matters operationally because support work depends on choosing the right inspection path first. A Windows-side role that was validated from Jump64 should not be treated like a Linux node that was validated from MSPUbuntuJump, and a storage consumer should not be treated like the isolated SAN endpoint underneath it.")
    add_table(
        doc,
        "Table 5. Site 2 logical service inventory and role mapping",
        ["Service scope", "Systems", "Primary function", "Administrative path", "Operational boundary"],
        [
            ["MSP control plane", "OPNsense, Jump64, MSPUbuntuJump, S2Veeam", "Entry, segmentation, inspection, and recovery control", "WAN NAT to the two bastions and internal administrative access to S2Veeam", "Defines who can reach the rest of the site and how offsite protection is carried."],
            ["Company 1 identity", "C1DC1, C1DC2", "Directory, DNS, and authentication", "Jump64 WinRM plus MSP path reachability", "Authoritative base for Company 1 naming and logon behavior."],
            ["Company 1 service layer", "C1FS, C1WebServer", "File presentation and internal web publication", "Jump64 for Windows inspection", "Turns storage and IIS configuration into user-visible services."],
            ["Company 1 client layer", "C1WindowsClient, C1UbuntuClient", "Endpoint validation from Windows and Linux perspectives", "Jump64 for Windows, MSPUbuntuJump for Linux", "Confirms that named services are consumable from endpoints, not only observable on servers."],
            ["Company 2 identity", "C2IdM1, C2IdM2", "Samba AD, DNS, DHCP, and shared-forest visibility", "MSPUbuntuJump and direct node inspection", "Control plane for Company 2 logon, name resolution, and lease continuity."],
            ["Company 2 service layer", "C2FS, C2WebServer", "File presentation and internal HTTPS publication", "MSPUbuntuJump and node-side checks", "Turns isolated storage and hostname bindings into user-visible service endpoints."],
            ["Company 2 client layer", "C2LinuxClient", "Endpoint-side identity and service validation", "MSPUbuntuJump and direct client checks", "Confirms that the Company 2 control plane is consumable by an actual tenant client."],
            ["Isolated storage", "C1SAN, C2SAN", "Block transport to file-service hosts", "Indirect through the consuming hosts", "Remains off tenant LAN and DMZ paths."],
        ],
    )
    add_table(
        doc,
        "Table 5A. Observed platform baseline - all service scopes",
        ["System", "Platform", "vCPU", "RAM", "Disk or volume state", "Interface detail", "Operational note"],
        [
            ["MSPUbuntuJump", "Ubuntu Linux bastion", "Not directly observed", "Not directly observed", "No tenant data role", "172.30.65.179 on MSP segment", "Served as the main Linux inspection host during the final pass."],
            ["C1DC1", "Windows Server domain controller", "Not directly observed", "Not directly observed", "Not directly observed", "172.30.65.2 on C1LAN", "Service-state evidence was captured from Jump64, but the final evidence set did not include a numeric hardware inventory."],
            ["C1DC2", "Windows Server domain controller", "Not directly observed", "Not directly observed", "Not directly observed", "172.30.65.3 on C1LAN", "Service-state evidence was captured from Jump64, but the final evidence set did not include a numeric hardware inventory."],
            ["C1FS", "Windows Server file server", "Not directly observed", "Not directly observed", "System disk not directly observed; dedicated F: SharedData volume confirmed", "172.30.65.4 on C1LAN", "File-service state and dedicated data volume were observed directly, but full numeric hardware values were not captured in the final evidence set."],
            ["C1WindowsClient", "Windows 10/11 client", "Not directly observed", "Not directly observed", "Not directly observed", "172.30.65.11 on C1LAN", "Endpoint role, domain state, and dual-web access were confirmed from Jump64; numeric hardware inventory was not captured."],
            ["C1UbuntuClient", "Ubuntu 25.04", "4", "7.3 GiB", "32 GB root disk plus 3.8 GiB swap", "172.30.65.36/26 on C1LAN", "Used as the Company 1 Linux client validation point."],
            ["C1WebServer", "Windows Server IIS host, workgroup-hosted", "Not directly observed", "Not directly observed", "Not directly observed", "172.30.65.162 on C1DMZ", "IIS binding and hostname-only publication were observed from Jump64, but numeric hardware inventory was not captured in the final evidence set."],
            ["C2IdM1", "Ubuntu 22.04.5 LTS", "4", "7.8 GiB", "32 GB system disk, 15 GB root LV", "ens18", "Primary Samba AD, DNS, DHCP node."],
            ["C2IdM2", "Ubuntu 22.04.5 LTS", "4", "7.8 GiB", "32 GB system disk, 15 GB root LV", "ens18", "Secondary Samba AD, DNS, DHCP node."],
            ["C2FS", "Ubuntu 22.04.5 LTS", "4", "7.8 GiB", "16 GB system disk, 160 GB mounted data disk at /mnt/c2_public", "ens19 172.30.65.68/26, ens18 172.30.65.195/29", "Dual-NIC file-service host with isolated storage transport."],
            ["C2LinuxClient", "Ubuntu 25.04", "4", "7.3 GiB", "32 GB disk", "ens18", "Company 2 Linux client validation point."],
            ["C2WebServer", "Ubuntu 22.04.5 LTS", "4", "7.8 GiB", "32 GB system disk, 30 GB root LV", "ens18", "nginx HTTPS host in C2DMZ."],
        ],
    )
    add_para(doc, "Table 5A deliberately separates platform facts that were measured from platform facts that were only service-validated. C1UbuntuClient had a direct hardware baseline in the March evidence set, while the Company 1 Windows systems were validated for service behavior, role placement, and access path without a full numeric hardware inventory. Keeping that distinction visible is more honest than implying that every C1 host was hardware-profiled to the same depth as the Linux nodes.")
    add_para(doc, "This platform baseline explains why the next section concentrates on OPNsense and routed separation. Once the host roles are fixed, the important question is how policy controls the paths between them.")
    add_figure(doc, *FIGURES[1])

    add_heading(doc, "3.3 MSP Entry, Network Segmentation, Remote Access, and Security", 2)
    add_para(doc, "The Site 2 network design is intentionally strict. Only Jump64 and MSPUbuntuJump are published at the edge, tenant LAN and DMZ paths are separated, and the OpenVPN path to Site 1 is limited to documented web and backup-copy use. That design matches the OPNsense operating model described in the vendor guidance [1]-[3] and the site records [9]-[11].")
    add_table(
        doc,
        "Table 6. Site 2 network segments and gateways",
        ["Segment", "Addressing or interface", "Gateway", "Purpose", "Key members"],
        [
            ["WAN", "172.20.64.1/16", "Upstream", "External network for controlled NAT publication", "OPNsense WAN edge."],
            ["MSP segment", "172.30.65.177/29", "OPNsense MSP interface", "Administrative entry and control plane", "OPNsense 172.30.65.177, Jump64 172.30.65.178, MSPUbuntuJump 172.30.65.179, S2Veeam 172.30.65.180."],
            ["C1LAN", "172.30.65.1/26", "OPNsense C1LAN interface", "Company 1 routed tenant LAN", "C1DC1 172.30.65.2, C1DC2 172.30.65.3, C1FS 172.30.65.4, C1WindowsClient 172.30.65.11, C1UbuntuClient 172.30.65.36."],
            ["C1DMZ", "172.30.65.161/29", "OPNsense C1DMZ interface", "Company 1 web segment", "C1WebServer 172.30.65.162."],
            ["C2LAN", "172.30.65.65/26", "OPNsense C2LAN interface", "Company 2 routed tenant LAN", "C2IdM1 172.30.65.66, C2IdM2 172.30.65.67, C2FS service NIC 172.30.65.68, C2LinuxClient 172.30.65.70."],
            ["C2DMZ", "172.30.65.169/29", "OPNsense C2DMZ interface", "Company 2 web segment", "C2WebServer 172.30.65.170."],
            ["SITE1_OVPN", "OpenVPN inter-site interface", "Site 1 OpenVPN gateway", "Inter-site web access and backup-copy path", "Cross-site rules and the static route to Site 1 Veeam."],
            ["C1SAN", "172.30.65.186/29", "172.30.65.185", "Company 1 isolated storage bridge", "Consumed by C1FS through iSCSI; not routed through OPNsense."],
            ["C2SAN", "172.30.65.194/29", "172.30.65.193", "Company 2 isolated storage bridge", "Consumed by C2FS through iSCSI; not routed through OPNsense."],
        ],
    )
    add_figure(doc, *FIGURES[2])
    add_table(
        doc,
        "Table 7. OPNsense exposure, routing, and firewall policy summary",
        ["Policy item", "Exact value", "Observed meaning"],
        [
            ["WAN NAT", "33464 -> 172.30.65.178:3389", "Publishes only Jump64 RDP to the edge."],
            ["WAN NAT", "33564 -> 172.30.65.179:22", "Publishes only MSPUbuntuJump SSH to the edge."],
            ["Alias set", "C1_Nets, C2_Nets, C1_REMOTE, C2_REMOTE, ALL_WEBS, ALL_DNS, C1_DCs, C2_DCs", "Abstracts tenant and service boundaries into reusable policy objects."],
            ["Backup aliases", "S2_VEEAM = 172.30.65.180; SITE1_VEEAM = 192.168.64.20", "Pins the backup-copy path to explicit endpoints."],
            ["Backup ports", "VEEAM_COPY_PORTS = 135, 445, 6160, 6162, 2500-3000, 10005, 10006", "Limits the copy path to documented Veeam transport ports."],
            ["LAN policy", "C1LAN allowed to C1_GLOBAL, ALL_WEBS, ALL_DNS; blocked to C2_GLOBAL", "Keeps Company 1 functional without allowing arbitrary Company 2 access."],
            ["LAN policy", "C2LAN allowed to C2_GLOBAL, ALL_WEBS, ALL_DNS; blocked to C1_GLOBAL", "Keeps Company 2 functional without allowing arbitrary Company 1 access."],
            ["Cross-site web rule", "C1_REMOTE -> 172.30.65.170/32 on HTTP/HTTPS", "Allows Site 1 users to consume the Company 2 web host by rule, not by flat reachability."],
            ["Cross-site web rule", "C2_REMOTE -> 172.30.65.162/32 on HTTP/HTTPS", "Allows reciprocal Site 2 to Company 1 web access by rule."],
            ["Static route", "192.168.64.20/32 via Site 1 OpenVPN gateway", "Directs offsite backup traffic to Site 1 Veeam over the VPN path."],
            ["MSP validation result", "OPNsense returned HTTP 403 on port 80, TCP 53 reachable, TCP 443 timed out", "Shows the management plane exists and is restricted, while the final pass did not include authenticated GUI review."],
        ],
    )
    add_figure(doc, *FIGURES[3])
    add_heading(doc, "Operational Interpretation", 3)
    add_para(doc, "The OPNsense results matter because they show both restraint and function at the same time. Only the two jump hosts were published at the edge, cross-tenant reachability was bounded by alias-driven policy, and the inter-site rules were narrow enough to support web consumption and Veeam copy traffic without flattening the site into one trust zone.")
    add_para(doc, "That policy reading also explains the observed management-plane behavior. HTTP 403 on port 80 showed that the gateway was present and not anonymously exposed, TCP 53 reachability showed that core network service remained reachable from the MSP side, and the timed-out HTTPS response simply marked the evidence limit of the final pass rather than a tenant-facing outage. With the boundary model clear, the tenant sections can be read in service terms instead of network terms, and Company 1 comes first because it shows how the Windows service chain fits inside the same boundary.")

    add_heading(doc, "3.4 Company 1 Directory Services, File Services, Web Delivery, and Client Access", 2)
    add_heading(doc, "Service Overview", 3)
    section_34 = [
        "Company 1 contributes the Windows directory pair C1DC1 and C1DC2, the file-service host C1FS, the workgroup IIS host C1WebServer, the Windows endpoint C1WindowsClient, the Linux endpoint C1UbuntuClient, and the isolated C1SAN bridge.",
        "Together they provide identity, internal file presentation, hostname-based web delivery, and client-facing validation inside the Site 2 operating boundary.",
        "The service set matters operationally because Company 1 is not a single remote dependency; it is a complete tenant stack with its own controller, client, storage, and web paths.",
        "The MSP bastions therefore had to validate Company 1 through both Linux-side and Windows-side inspection routes, not through one generic reachability check.",
    ]
    for text in section_34:
        add_para(doc, text)
    add_heading(doc, "Architectural Rationale", 3)
    more_34 = [
        "Microsoft directory services stay on C1DC1 and C1DC2 so authentication and DNS responsibilities remain on dedicated hosts.",
        "C1FS stays separate from the domain controllers so share administration and storage consumption do not compete with directory workloads.",
        "C1WebServer remains in the DMZ, workgroup hosted, and bound only to c1-webserver.c1.local on TCP 443, which preserves a named-service contract without adding unnecessary domain membership [7].",
        "C1SAN stays off the routed tenant networks, which keeps block transport hidden behind C1FS and lets support teams diagnose storage, share, and client issues as separate layers.",
    ]
    section_34.extend(more_34)
    for text in more_34:
        add_para(doc, text)
    add_heading(doc, "Observed Operating State", 3)
    obs_34 = [
        "Inspection showed that C1DC1 at 172.30.65.2 and C1DC2 at 172.30.65.3 returned domain, forest, controller inventory, and directory-service state successfully from Jump64 over WinRM.",
        "Jump64 also showed C1FS at 172.30.65.4 with the Windows SMB service active, a dedicated F: volume labeled SharedData, named shares on that volume, and an active iSCSI initiator session tied to Company 1 storage.",
        "C1WebServer at 172.30.65.162 showed a single HTTPS IIS binding for c1-webserver.c1.local on TCP 443, while raw IP access returned HTTP 404 as designed.",
        "C1WindowsClient at 172.30.65.11 was inspected from Jump64 through WMI because TCP 5985 was not open on that host, and the checks confirmed c1.local domain membership, Company 1 DNS use, and successful HTTP 200 access to both internal web hostnames [13].",
        "MSPUbuntuJump also showed C1UbuntuClient at 172.30.65.36 with admin@C1UbuntuClient shell context, active C1.LOCAL realm membership, and HTTP 200 results for both web hostnames, while C1SAN remained intentionally indirect and was confirmed through the live C1FS iSCSI consumer session.",
    ]
    section_34.extend(obs_34)
    for text in obs_34:
        add_para(doc, text)
    add_table(
        doc,
        "Table 8. Company 1 service summary",
        ["System", "Address", "Service role", "Observed state", "Operational reading"],
        [
            ["C1DC1", "172.30.65.2", "Primary domain controller", "WinRM inspection returned domain, forest, and directory-service state successfully", "Authoritative Company 1 identity source."],
            ["C1DC2", "172.30.65.3", "Secondary domain controller", "WinRM inspection returned matching controller and service state", "Provides Company 1 controller parity and resilience."],
            ["C1FS", "172.30.65.4", "Windows file server", "SMB active, F: SharedData present, named shares visible, active iSCSI session confirmed", "Converts isolated storage into user-visible SMB service."],
            ["C1WebServer", "172.30.65.162", "IIS internal web server", "Workgroup-hosted, HTTPS binding only for c1-webserver.c1.local, raw IP returns HTTP 404", "Publishes the site by name, not by generic address."],
            ["C1WindowsClient", "172.30.65.11", "Windows endpoint", "Domain membership, Company 1 DNS use, and dual-web HTTP 200 responses confirmed from Jump64", "Validates the Company 1 service contract from a Windows endpoint."],
            ["C1UbuntuClient", "172.30.65.36", "Linux endpoint", "C1.LOCAL realm active, dual-web HTTP 200 responses confirmed", "Validates the same contract from a Linux endpoint."],
            ["C1SAN", "172.30.65.186/29 via gateway 172.30.65.185", "Isolated Company 1 storage bridge", "Not managed directly from bastions; storage consumer confirmed on C1FS", "The supportable signal is the live iSCSI consumer on the file-service host."],
        ],
    )
    add_figure(doc, *FIGURES[4])
    add_figure(doc, *FIGURES[5])
    add_figure(doc, *FIGURES[6])
    add_figure(doc, *FIGURES[7])
    add_figure(doc, *FIGURES[8])
    add_figure(doc, *FIGURES[9])
    add_figure(doc, *FIGURES[10])
    add_heading(doc, "Service Composition and Operational Reading", 3)
    tail_34 = [
        "These observations read as one service chain rather than unrelated host checks.",
        "C1DC1 and C1DC2 establish the naming and authentication base, C1FS turns isolated block storage into user-visible SMB paths, and C1WebServer publishes the internal application under the documented hostname.",
        "The two client perspectives then show that the Company 1 service contract is consumable from both Windows and Linux endpoints.",
        "The Company 1 storage story is also supportable because direct management of C1SAN is not required for normal operations; the authoritative operational signal is the active storage consumer on C1FS.",
        "That arrangement gives the receiving team a clear order for troubleshooting: identity first, file or storage next, web binding next, and client experience last.",
        "The Windows emphasis in this tenant section also matters for support realism. Jump64 was not used as a symbolic bastion but as the active Windows-side inspection platform that could interrogate controllers, the file server, the IIS host, the endpoint, and S2Veeam through the same administrative context. That continuity makes the Company 1 evidence easier to trust because the same management path observed both infrastructure state and endpoint-side consumption.",
        "Company 1 also shows why isolation does not mean fragmentation. The web server remains workgroup hosted, the storage bridge remains hidden behind C1FS, and the two clients were validated through different methods, yet the service chain still reads coherently because every layer points back to the same named services and tenant boundary. In handover terms, that gives the receiving team a tenant model they can reason about without needing an undocumented verbal explanation.",
    ]
    section_34.extend(tail_34)
    for text in tail_34:
        add_para(doc, text)

    add_heading(doc, "3.5 Company 2 Identity Services, DNS, DHCP, and Shared Forest Design", 2)
    add_heading(doc, "Service Overview", 3)
    section_35 = [
        "Company 2 centers on C2IdM1 and C2IdM2 as the shared-forest identity, DNS, and DHCP pair for c2.local and the related namespace visibility into c1.local.",
        "That pair supports the rest of the Company 2 stack, including C2FS, C2LinuxClient, and C2WebServer, by keeping authentication, naming, and lease services consistent across the tenant.",
        "The section focuses on identity services because those are the control-plane systems that make later file, web, and client behaviors understandable.",
        "It also explains the shared-forest reading because cross-domain name visibility is part of the design evidence, not an accidental side effect.",
    ]
    for text in section_35:
        add_para(doc, text)
    add_heading(doc, "Architectural Rationale", 3)
    more_35 = [
        "Using two Samba AD nodes keeps directory, DNS, and DHCP roles readable and resilient inside the Company 2 scope [4]-[6].",
        "C2IdM1 holds the primary DHCP failover role and C2IdM2 holds the secondary role, which avoids hiding lease continuity inside one host.",
        "Hosting c2.local, c1.local, and _msdcs.c2.local on both identity nodes keeps namespace availability aligned with the dual-node design.",
        "The shared forest model also fits the wider site because it supports cross-domain visibility without collapsing tenant service boundaries.",
    ]
    section_35.extend(more_35)
    for text in more_35:
        add_para(doc, text)
    add_heading(doc, "Observed Operating State", 3)
    obs_35 = [
        "Inspection showed samba-ad-dc active on both C2IdM1 at 172.30.65.66 and C2IdM2 at 172.30.65.67, with isc-dhcp-server active on both nodes.",
        "Checks also confirmed DHCP failover roles as Primary on C2IdM1 and Secondary on C2IdM2.",
        "DNS records on both nodes resolved c1-webserver.c1.local to 172.30.64.162 and 172.30.65.162, and c2-webserver.c2.local to 172.30.64.170 and 172.30.65.170.",
        "The hosted zones observed were c2.local, c1.local, and _msdcs.c2.local, and the principal set visible on both nodes included Administrator, admin, employee1, employee2, and c2_file_users.",
        "C2LinuxClient later consumed this identity plane with DNS servers 172.30.65.66 and 172.30.65.67 and search domains c1.local and c2.local, which matched the expected client reading.",
    ]
    section_35.extend(obs_35)
    for text in obs_35:
        add_para(doc, text)
    add_table(
        doc,
        "Table 9. Company 2 identity, DNS, and DHCP summary",
        ["Object", "Address or scope", "Observed state", "Operational reading"],
        [
            ["C2IdM1", "172.30.65.66", "Ubuntu 22.04.5 LTS, samba-ad-dc active, isc-dhcp-server active, DHCP Primary", "Primary Company 2 identity and lease node."],
            ["C2IdM2", "172.30.65.67", "Ubuntu 22.04.5 LTS, samba-ad-dc active, isc-dhcp-server active, DHCP Secondary", "Secondary Company 2 identity and lease node."],
            ["Hosted zones", "c2.local, c1.local, _msdcs.c2.local", "Visible on both nodes", "Cross-domain name visibility aligns with the documented shared forest."],
            ["A records", "c1-webserver.c1.local and c2-webserver.c2.local", "172.30.64.162 and 172.30.65.162; 172.30.64.170 and 172.30.65.170", "Dual-record publication supports routed and cross-site consumption paths."],
            ["Principal inventory", "Administrator, admin, employee1, employee2, c2_file_users", "Visible on both nodes", "Confirms directory objects required for client and file-share tests."],
            ["Resolver dependency", "C2LinuxClient", "DNS servers 172.30.65.66 and 172.30.65.67; search domains c1.local and c2.local", "Client behavior matched the identity-plane design."],
        ],
    )
    add_figure(doc, *FIGURES[11])
    add_figure(doc, *FIGURES[12])
    add_heading(doc, "Service Composition and Operational Reading", 3)
    tail_35 = [
        "Company 2 therefore provides the authoritative naming and lease layer for its own tenant while also carrying the shared-forest visibility that makes dual-hostname access intelligible. In operational terms, that means C2IdM1 and C2IdM2 are not just background directory servers; they are the systems that explain why later checks on C2FS, C2LinuxClient, and C2WebServer behaved consistently under the documented names and domains.",
        "The paired-node design also gives support staff a cleaner way to reason about faults. A hostname failure, lease anomaly, or logon problem can be tested against two equivalent identity nodes before anyone assumes that the entire tenant stack has failed. Because DNS, DHCP, and directory roles sit together on both systems, the control plane remains compact enough to troubleshoot without losing service separation.",
        "The shared-forest reading becomes stronger in that context. Cross-domain DNS visibility and client resolver behavior show that Company 2 identity is not operating in isolation, yet the tenant boundary still holds because cross-tenant reachability remains policy-bound and service-specific. That combination is what makes the Company 2 control plane feel engineered rather than accidental.",
        "Taken together, the identity pair gives the site a stable explanation for downstream behavior. File-share access on C2FS, user resolution on C2LinuxClient, and hostname-based publication on C2WebServer all depend on the same naming and authentication base. When those later sections succeed, they do so because the Company 2 identity layer is functioning as the intended control plane.",
    ]
    section_35.extend(tail_35)
    for text in tail_35:
        add_para(doc, text)
    add_figure(doc, *FIGURES[13])

    add_heading(doc, "3.6 Storage, File Services, and Isolated SAN Design", 2)
    add_para(doc, "The storage layer is where Site 2 becomes operational instead of merely descriptive. Both tenants hide block transport behind file-service hosts, which keeps tenant users on SMB paths and keeps storage troubleshooting inside a smaller, more predictable fault domain [5]-[6], [14].")
    add_heading(doc, "Company 2 File-Service State (C2FS)", 3)
    add_bullets(
        doc,
        [
            "smbd: active",
            "Mounted: /dev/sdb at /mnt/c2_public",
            "Active iSCSI session: 172.30.65.194:3260, target iqn.2024-03.org.clearroots:c2san",
            "Shares: [C2_Public] -> /mnt/c2_public/Public; [C2_Private] -> /mnt/c2_public/Private/%U",
            "Hostname-based SMB validated from C2LinuxClient: //c2fs.c2.local/C2_Public and //c2fs.c2.local/C2_Private",
            "Sync result: successful (Site 1 -> Site 2)",
        ],
    )
    add_figure(doc, *FIGURES[14])
    add_figure(doc, *FIGURES[15])
    add_heading(doc, "Company 1 File-Service State (C1FS)", 3)
    add_bullets(
        doc,
        [
            "Windows SMB service: active",
            "Dedicated data volume: F: drive labeled SharedData",
            "Named SMB shares: present on the F: SharedData volume",
            "Active iSCSI initiator session: confirmed via Get-IscsiSession and tied to Company 1 SAN",
        ],
    )
    add_heading(doc, "SAN Isolation Model", 3)
    add_para(doc, "The two SAN bridges are intentionally not ordinary routed service hosts. Company 1 storage uses C1SAN 172.30.65.186/29 with gateway 172.30.65.185. Company 2 storage uses C2SAN 172.30.65.194/29 with gateway 172.30.65.193. Neither bridge is meant to act as a tenant-facing endpoint. Their role is to deliver block transport to the file-service hosts and then disappear behind those hosts operationally.")
    add_figure(doc, *FIGURES[16])
    add_figure(doc, *FIGURES[17])
    add_heading(doc, "Share Presentation Model", 3)
    add_para(doc, "C2FS presents two share types to the Company 2 tenant: the shared C2_Public path and the per-user C2_Private path under /mnt/c2_public/Private/%U. C1FS follows the same layered idea on Windows even though the exact share names were not enumerated in the environment facts. In both cases, users interact with SMB shares while the storage bridge remains hidden underneath the file-service host.")
    add_table(
        doc,
        "Table 10. Storage and isolated SAN summary",
        ["Layer", "Company 1", "Company 2", "Operational reading"],
        [
            ["Isolated storage bridge", "C1SAN 172.30.65.186/29, gateway 172.30.65.185", "C2SAN 172.30.65.194/29, gateway 172.30.65.193", "Both SAN paths stay outside the routed tenant LAN and DMZ segments."],
            ["File-service host", "C1FS 172.30.65.4", "C2FS 172.30.65.68 with service NIC ens19 172.30.65.68/26 and storage NIC ens18 172.30.65.195/29", "Each tenant consumes block storage through a dedicated file-service host."],
            ["Mounted or dedicated volume", "F: SharedData", "/dev/sdb mounted at /mnt/c2_public", "The file-service host, not the SAN bridge, presents usable data paths."],
            ["SMB presentation", "Windows SMB shares on F: SharedData", "C2_Public and C2_Private via Samba", "Share-level administration remains above the storage layer."],
            ["Transport signal", "Active iSCSI initiator session confirmed on C1FS", "Active iSCSI session to 172.30.65.194:3260 target iqn.2024-03.org.clearroots:c2san", "The correct operational signal is the consumer session on the file-service host."],
            ["User-visible access", "Company 1 SMB shares", "//c2fs.c2.local/C2_Public and //c2fs.c2.local/C2_Private", "Users consume names and shares, not storage endpoints."],
        ],
    )
    add_para(doc, "The next section keeps the same operational reading but shifts it to the client perspective. Once identity and file services are proven, the important question is how real endpoints consume them.")

    add_heading(doc, "3.7 Client Access, Identity Validation, and Dual-Hostname Web Delivery", 2)
    add_para(doc, "Client validation closes the gap between server-side service state and user-visible service delivery. Site 2 had three usable client perspectives: C1UbuntuClient for Company 1 Linux behavior, C1WindowsClient for Company 1 Windows behavior, and C2LinuxClient for Company 2 Linux behavior.")
    add_para(doc, "Those perspectives matter because hostname-based web publication and domain-backed file access are only meaningful when they succeed from endpoints. They also help separate identity issues from web binding or storage issues when the same user reports a general service failure.")
    add_para(doc, "C1UbuntuClient was observed as admin@C1UbuntuClient with C1.LOCAL active, and both web hostnames returned HTTP 200. C2LinuxClient showed C2.LOCAL state, employee1@c2.local and employee2@c2.local through getent passwd, DNS servers 172.30.65.66 and 172.30.65.67, search domains c1.local and c2.local, and successful HTTP 200 access to both web hostnames.")
    add_figure(doc, *FIGURES[18])
    add_figure(doc, *FIGURES[19])
    add_table(
        doc,
        "Table 11. Client access and identity summary",
        ["Validation item", "C1UbuntuClient", "C1WindowsClient", "C2LinuxClient"],
        [
            ["Role", "Company 1 Linux client", "Company 1 Windows client", "Company 2 Linux client"],
            ["Identity context", "Realm C1.LOCAL active", "Domain membership confirmed in c1.local", "C2.LOCAL visible"],
            ["Inspection method", "Direct shell observation from MSPUbuntuJump", "Jump64 WMI-backed inspection; reason stated in Section 3.4", "Direct shell observation"],
            ["Resolver state", "Company 1 naming path working; both hostnames resolved", "Company 1 DNS use confirmed", "DNS servers 172.30.65.66 and 172.30.65.67; search domains c1.local and c2.local"],
            ["https://c1-webserver.c1.local", "HTTP 200", "HTTP 200", "HTTP 200"],
            ["https://c2-webserver.c2.local", "HTTP 200", "HTTP 200", "HTTP 200"],
            ["File-share perspective", "Direct client-side Company 1 SMB access was not re-tested from this endpoint in the final pass; Company 1 share availability was confirmed from C1FS instead", "Direct client-side Company 1 SMB access was not re-tested from this endpoint in the final pass; Company 1 share availability was confirmed from C1FS instead", "//c2fs.c2.local/C2_Public and //c2fs.c2.local/C2_Private validated"],
            ["Operational meaning", "Confirms Company 1 service consumption from Linux", "Confirms Company 1 service consumption from Windows", "Confirms Company 2 identity, naming, SMB, and web consumption from Linux"],
        ],
    )
    add_table(
        doc,
        "Table 12. Internal web delivery summary",
        ["Endpoint or name", "Publication path", "Observed result", "Operational reading"],
        [
            ["https://c1-webserver.c1.local", "Hostname-only IIS publication on C1WebServer", "HTTP/2 200", "Company 1 web delivery is intended to succeed by hostname."],
            ["https://172.30.65.162", "Direct IP to C1WebServer", "HTTP/2 404", "Raw IP access is intentionally not the publication contract."],
            ["https://c2-webserver.c2.local", "Hostname-based nginx publication on C2WebServer", "HTTP/1.1 200 OK", "Company 2 web delivery follows the same named-service model [12]."],
            ["https://172.30.65.170", "Direct IP to C2WebServer", "HTTP/1.1 404 Not Found", "Raw IP access is intentionally not the publication contract."],
        ],
    )
    add_para(doc, "The internal web results lead naturally into protection and recovery. Once named service delivery is proven, the next operational question is how the site protects and restores those services when the live layer fails.")

    add_heading(doc, "3.8 Backup, Recovery, and Offsite Protection", 2)
    add_para(doc, "S2Veeam at 172.30.65.180 is the protection anchor for Site 2. The host was reachable from MSPUbuntuJump on TCP 445, 9392, 5985, 10005, and 10006, and it was also reachable from Jump64 through the local administrator context. That means the backup layer was not treated as a theoretical design component; it was an observed administrative service [8], [11].")
    add_para(doc, "The protection model is intentionally layered. Backup jobs preserve Windows and Ubuntu server workloads, file-share jobs protect the Company 1 and Company 2 file estates directly, and copy jobs push data toward Site 1 over the documented OpenVPN path. Synchronization remains a separate file-continuity mechanism and should not be mistaken for backup.")
    add_table(
        doc,
        "Table 13. Backup and offsite-protection summary",
        ["Item", "Observed state", "Operational reading"],
        [
            ["S2Veeam host", "172.30.65.180 reachable from MSPUbuntuJump and Jump64", "Backup administration is part of the MSP control plane."],
            ["Active services", "VeeamBackupSvc, VeeamBrokerSvc, VeeamDeploySvc, VeeamExplorersRecoverySvc, VeeamFilesysVssSvc, VeeamMountSvc, VeeamNFSSvc", "Core backup, broker, deploy, restore, mount, and NFS functions were active."],
            ["Local repository", "Site2Veeam on Z:\\Site2AgentBackups", "Primary on-site repository for Site 2 recoveries."],
            ["Offsite SMB target", "\\\\192.168.64.20\\Site2OffsiteFromSite2", "Site 1 receives backup copies over the documented route."],
            ["Job families", "Ubuntu_Servers, Windows_Servers, C1_FileShare, C2_FileShare", "Protection covers both host-level and file-share-level recovery paths."],
            ["Copy jobs", "Present", "Offsite protection is implemented, not merely planned."],
            ["Protected workload count", "10 machines", "The backup scope reaches beyond one tenant or one platform."],
        ],
    )
    add_para(doc, "Current operational state is strongest when the repository, job families, and copy targets are read together. Local backups support fast recovery inside Site 2, while the Site 1 copy path matters when the failure boundary is larger than one host or one volume.")
    add_figure(doc, *FIGURES[20])

    add_heading(doc, "3.9 Requirement-to-Implementation Traceability", 2)
    add_para(doc, "The traceability matrix below maps the site requirements to specific implementation outcomes and evidence sources. Company 1 and Company 2 are shown at the same level of granularity so the handover record does not privilege one tenant over the other.")
    add_table(
        doc,
        "Table 14. Requirement-to-implementation traceability matrix",
        ["Requirement area", "Implemented systems and facts", "Evidence source", "Validated outcome", "Operational reading"],
        [
            ["Company 1 directory services", "C1DC1 and C1DC2; WinRM-observed from Jump64", "Jump64 service-state observation", "Company 1 domain and controller state returned successfully", "Company 1 identity is active from the approved Windows bastion."],
            ["Company 1 file services and isolated storage", "C1FS F: SharedData volume, named shares, active iSCSI session, C1SAN", "Jump64 observation plus SAN design record", "File-service and storage-consumer chain confirmed", "Company 1 file delivery rests on an isolated storage bridge that stays behind C1FS."],
            ["Company 1 web delivery", "C1WebServer, IIS, hostname-only binding on TCP 443, HTTP 404 on raw IP", "Jump64 observation and HTTP validation", "Hostname succeeds while raw IP does not", "Company 1 web publication follows a named-service contract."],
            ["Company 1 client validation", "C1WindowsClient and C1UbuntuClient", "Jump64 endpoint inspection and MSPUbuntuJump shell observation", "Windows and Linux endpoints consumed both internal web hostnames", "Company 1 services are usable from actual endpoints."],
            ["Company 2 identity, DNS, and DHCP", "C2IdM1 and C2IdM2, Samba AD, DHCP failover", "Direct node observation", "AD, DNS, and DHCP active on both nodes", "Company 2 control plane is live and paired."],
            ["Company 2 file services and isolated storage", "C2FS, iSCSI session, C2_Public and C2_Private shares, C2SAN", "Direct node observation and hostname-based share validation", "Mounted storage and share definitions confirmed", "Company 2 file delivery rests on an isolated storage bridge that stays behind C2FS."],
            ["Company 2 web delivery", "C2WebServer, nginx, hostname-only publication, HTTP 404 on raw IP", "HTTP validation and service observation", "Hostname succeeds while raw IP does not", "Company 2 web publication follows the same named-service model."],
            ["Company 2 client validation", "C2LinuxClient, resolver state, SMB access", "Direct client observation", "Employee principals, dual-web success, and hostname-based SMB access confirmed", "Company 2 services are usable from an actual tenant client."],
        ],
    )

    add_heading(doc, "3.10 Service Dependencies, Failure Domains, and Access Model", 2)
    add_para(doc, "Dependency mapping matters because Site 2 failures do not all begin in the same layer. An access-path problem, a DNS problem, a storage problem, and a backup problem all surface differently and should not be chased with the same first action. The two tables below keep service dependency separate from administrative authority so support staff can choose the correct boundary first.")
    add_table(
        doc,
        "Table 15. Service dependency and failure-domain view",
        ["Service area", "Depends on", "Primary failure sign", "First authoritative check", "Escalation path"],
        [
            ["Administrative entry", "WAN NAT, Jump64, MSPUbuntuJump", "No approved management entry", "NAT publication and bastion reachability", "MSP networking and firewall control."],
            ["Company 1 identity", "C1DC1, C1DC2, DNS path", "Logon or Company 1 name failures", "Jump64 WinRM to both controllers", "Company 1 Windows identity administration."],
            ["Company 2 identity", "C2IdM1, C2IdM2, DNS path, DHCP state", "Company 2 name resolution or lease faults", "Service state on both identity nodes", "Company 2 Linux identity administration."],
            ["Company 1 file and storage", "C1FS, C1SAN, iSCSI session, SMB service", "Shares unavailable or F: volume unavailable", "Get-IscsiSession, volume state, SMB service", "Company 1 Windows file-service administration."],
            ["Company 2 file and storage", "C2FS, C2SAN, mount state, smbd", "Shares unavailable or mount missing", "iSCSI session, /mnt/c2_public, smbd state", "Company 2 Linux file-service administration."],
            ["Internal web delivery", "DNS, hostname bindings, web-service state", "Hostname fails or raw IP behavior changes", "DNS records and host bindings", "Tenant web-service administration."],
            ["Backup and copy", "S2Veeam, repository, SITE1_VEEAM route, VEEAM_COPY_PORTS", "Job or copy failures", "Veeam services, repository path, route, and copy-job state", "MSP backup administration."],
        ],
    )
    add_table(
        doc,
        "Table 16. Authentication and authorization model",
        ["Scope", "Authority source", "Administrative access model", "Consumer access model", "Operational meaning"],
        [
            ["MSP control plane", "Local administrative control on Jump64, MSPUbuntuJump, S2Veeam, and OPNsense", "Controlled bastion-first administration", "Not tenant-consumable", "Administration begins from the management scope, not the tenant scopes."],
            ["Company 1 identity", "c1.local on C1DC1 and C1DC2", "Jump64 WinRM to domain controllers and related Windows services", "Company 1 endpoints consume directory-backed services", "Identity authority stays inside the Company 1 controller pair."],
            ["Company 2 identity", "c2.local on C2IdM1 and C2IdM2", "Linux-side administrative inspection and service management", "Company 2 client and file-service behavior consume the shared control plane", "Identity and DNS authority stay inside the Company 2 controller pair."],
            ["Company 1 file access", "Windows SMB share permissions on C1FS", "Windows server administration", "Company 1 endpoints consume named SMB shares", "Share presentation stays above the isolated storage layer."],
            ["Company 2 file access", "Samba share rules, valid users, and c2_file_users group", "Linux server administration on C2FS", "C2LinuxClient consumes //c2fs.c2.local/C2_Public and //c2fs.c2.local/C2_Private", "Authentication and authorization meet at the file-service host."],
            ["Internal web access", "DNS and host binding configuration", "Web-host administration on C1WebServer and C2WebServer", "Clients consume https://c1-webserver.c1.local and https://c2-webserver.c2.local", "Web delivery is controlled by name and binding, not by raw address."],
            ["Backup administration", "Local administrator context on S2Veeam", "Jump64 and MSP administrative access", "Not tenant-consumable", "Recovery control remains in the MSP scope."],
        ],
    )

    add_heading(doc, "3.11 Data Protection Flow", 2)
    add_para(doc, "Data protection in Site 2 must be read as three related but different mechanisms. First, isolated SAN paths deliver block storage to the file-service hosts. Second, synchronization keeps selected file content aligned from Site 1 to Site 2. Third, Veeam creates backup and copy chains that support restore activity when continuity alone is not enough.")
    add_para(doc, "Keeping backup separate from synchronization matters operationally. A sync success does not replace a restore point, and a restore point does not replace the live share path that users need during normal service. The table below preserves that separation so a receiving team can choose the right recovery mechanism for the right failure type.")
    add_table(
        doc,
        "Table 17. Storage, backup, and recovery data-flow summary",
        ["Data set or layer", "Live path", "Protection method", "Offsite element", "Recovery use"],
        [
            ["Company 1 shared data", "C1SAN -> C1FS F: SharedData -> Company 1 SMB shares", "Veeam C1_FileShare and Windows workload protection", "Copy jobs to Site 1 via S2Veeam", "File restore or host-level recovery depending on the fault boundary."],
            ["Company 2 public data", "C2SAN -> C2FS /mnt/c2_public/Public -> C2_Public", "Successful Site 1 -> Site 2 sync plus Veeam C2_FileShare protection", "Copy jobs to Site 1 via S2Veeam", "Continuity through sync, restore through backup."],
            ["Company 2 private data", "C2SAN -> C2FS /mnt/c2_public/Private/%U -> C2_Private", "Veeam C2_FileShare protection", "Copy jobs to Site 1 via S2Veeam", "Per-user data recovery without rebuilding the whole host."],
            ["Windows and Ubuntu server workloads", "Windows_Servers and Ubuntu_Servers job families", "Veeam workload protection", "Copy jobs present", "Host-level recovery when the service VM itself fails."],
            ["Backup repository", "Z:\\Site2AgentBackups", "Local repository handling on S2Veeam", "Offsite SMB target \\\\192.168.64.20\\Site2OffsiteFromSite2", "Keeps local and offsite recovery scopes distinct."],
        ],
    )

    add_heading(doc, "3.12 Maintenance and Routine Checks", 2)
    add_para(doc, "Routine maintenance should follow the control plane before it follows the tenant planes. The goal is to catch drift while the failure boundary is still small: entry paths first, identity next, storage and share presentation after that, web publication next, and recovery confidence last.")
    add_table(
        doc,
        "Table 18. Operational maintenance checks",
        ["Routine check", "System or path", "Why it matters"],
        [
            ["Confirm Jump64 and MSPUbuntuJump are reachable through the published entry points", "33464 -> 172.30.65.178:3389 and 33564 -> 172.30.65.179:22", "If the bastions fail, the approved management model fails before any tenant service can be repaired."],
            ["Confirm OPNsense service boundary", "172.30.65.177 plus route and alias review", "The gateway defines segmentation, remote access, and backup-copy transport."],
            ["Confirm Company 1 controller health", "C1DC1 and C1DC2", "Directory and DNS faults break downstream Company 1 services quickly."],
            ["Confirm Company 2 controller health", "C2IdM1 and C2IdM2", "AD, DNS, and DHCP faults break Company 2 endpoint behavior quickly."],
            ["Confirm Company 1 file and storage chain", "C1FS F: SharedData, SMB service, iSCSI session", "Share issues should be separated from storage-consumer faults."],
            ["Confirm Company 2 file and storage chain", "C2FS smbd, mount state, iSCSI session, sync result", "Share issues, mount issues, and transport issues must stay distinct."],
            ["Confirm hostname-based web delivery", "C1WebServer and C2WebServer", "A hostname success plus raw-IP 404 confirms the intended publication model remains intact."],
            ["Confirm backup and copy posture", "S2Veeam services, repository, job families, copy jobs", "Protection value depends on both local repository health and offsite-copy continuity."],
        ],
    )

    add_heading(doc, "3.13 Troubleshooting and Fast Triage Guide", 2)
    add_para(doc, "Fast triage works when the first check is authoritative. Site 2 is documented to reduce the habit of checking every host for every symptom. The table below points each common service symptom to the layer that is most likely to narrow the fault first.")
    add_table(
        doc,
        "Table 19. Troubleshooting and fast triage guide",
        ["Symptom", "First authoritative check", "Likely layer", "Immediate next action"],
        [
            ["No approved administrative entry", "WAN NAT and bastion reachability", "MSP edge or bastion layer", "Restore controlled entry before touching tenant services."],
            ["Company 1 logon or name failure", "Jump64 inspection of C1DC1 and C1DC2", "Company 1 identity", "Compare controller state and DNS availability before checking clients."],
            ["Company 2 logon, lease, or DNS failure", "Service state on C2IdM1 and C2IdM2", "Company 2 identity", "Validate AD, DHCP, and hosted-zone state before checking downstream services."],
            ["Company 1 share unavailable", "C1FS SMB state, F: SharedData, iSCSI session", "Company 1 file or storage chain", "Decide whether the fault is share, volume, or transport."],
            ["Company 2 share unavailable", "C2FS smbd, /mnt/c2_public, iSCSI session, sync state", "Company 2 file or storage chain", "Separate mount, Samba, and sync concerns."],
            ["Hostname fails but raw IP works", "DNS and host binding", "Naming or web-publication layer", "Fix named-service publication before widening network exposure."],
            ["Hostname and raw IP both fail", "Service state plus network reachability", "Host or path failure", "Determine whether the web host itself is down or unreachable."],
            ["Backup copy job fails", "S2Veeam services, repository, static route to 192.168.64.20/32, VEEAM_COPY_PORTS", "Protection transport layer", "Separate repository faults from inter-site path faults."],
        ],
    )

    add_heading(doc, "3.14 Integrated Design Summary", 2)
    add_para(doc, "The integrated reading of Site 2 is straightforward once the control plane, tenant planes, and protection plane are kept distinct. The site works because each layer does one job clearly and exposes only the paths it is supposed to expose.")
    add_table(
        doc,
        "Table 20. Integrated design summary",
        ["Design area", "Integrated outcome", "Evidence basis", "Operational meaning"],
        [
            ["MSP control plane", "Bastion-first administration with limited edge exposure", "NAT rules, OPNsense evidence, Jump64 and MSPUbuntuJump checks", "The site can be administered without exposing tenant workloads directly."],
            ["Company 1 service chain", "Directory, file, web, client, and isolated storage layers behave coherently", "Jump64 and MSPUbuntuJump observations", "Company 1 is a complete tenant stack, not a partial dependency."],
            ["Company 2 control plane", "Paired identity, DNS, DHCP, shared-forest visibility, and client behavior align", "Direct node and client observations", "Company 2 has a stable control plane that explains downstream service success."],
            ["Storage isolation", "Both tenants hide block transport behind file-service hosts", "iSCSI and volume observations plus design records", "Support teams can troubleshoot shares without treating SAN endpoints as user services."],
            ["Web publication", "Both tenants publish by hostname and reject raw-IP browsing", "HTTP validation plus host-side observations", "Web access depends on naming and bindings, not on accidental address exposure."],
            ["Protection and offsite copy", "Local repository and Site 1 copy path are both present", "S2Veeam administrative observations and route or rule evidence", "Recovery remains available when continuity alone is insufficient."],
        ],
    )

    add_heading(doc, "3.15 Limitations and Outstanding Items", 2)
    add_para(doc, "The final pass confirmed OPNsense management-plane presence through HTTP 403 on port 80 and TCP 53 reachability, but it did not include an authenticated GUI walkthrough. That is an evidence depth limit, not a service failure.")
    add_para(doc, "The final pass also confirmed live administrative access to S2Veeam from Jump64, but refreshed GUI screenshots were not captured for the repository, job families, and copy-job views. Those screenshots should be collected before final submission so the visual evidence matches the observed state.")
    add_para(doc, "C1SAN direct management access remains intentionally blocked from Jump64 and MSPUbuntuJump. The relevant confirmation is the active iSCSI consumer session observed on C1FS, which is the correct support signal for normal file-service operations.")

    add_heading(doc, "4. Conclusion", 1)
    add_para(doc, "Site 2 can be handed over as a coherent operating environment because the MSP, Company 1, and Company 2 scopes were all validated at the points that matter most for support. The MSP scope showed controlled entry through Jump64 and MSPUbuntuJump, enforceable segmentation through OPNsense, and active recovery tooling on S2Veeam.")
    add_para(doc, "Company 1 showed a complete Windows-led tenant chain: C1DC1 and C1DC2 returned active directory-service state, C1FS exposed its dedicated F: SharedData volume and active iSCSI consumer session, C1WebServer answered only on the documented hostname, and both Company 1 clients consumed the internal web services successfully. Company 2 showed an equally complete Linux-led chain: C2IdM1 and C2IdM2 carried active AD, DNS, and DHCP roles, C2FS mounted isolated storage and exposed the expected Samba shares, C2LinuxClient consumed identity and share services correctly, and C2WebServer answered only on the documented hostname.")
    add_para(doc, "Those results support the design choices recorded in this report. Entry stays bastion-first, tenant boundaries stay segmented, web publication stays name-based, storage stays isolated behind file-service hosts, and backup stays separate from synchronization. The remaining gaps are explicit and limited to evidence depth, not to service failure. On that basis, the site is ready for formal handover with a clear operational reading for day-to-day administration and first-line triage.")

    add_heading(doc, "5. Appendices", 1)
    add_heading(doc, "Appendix A: Observed Addressing, Gateways, and Endpoints", 2)
    add_table(
        doc,
        "Table A1. Observed addressing, gateways, and endpoints",
        ["Object", "Address or port", "Gateway or interface", "Notes"],
        [
            ["WAN", "172.20.64.1/16", "Upstream WAN", "External network used for NAT publication."],
            ["OPNsense MSP interface", "172.30.65.177/29", "MSP segment", "Gateway and segmentation control."],
            ["Jump64", "172.30.65.178", "MSP segment", "Published through 33464 -> 172.30.65.178:3389."],
            ["MSPUbuntuJump", "172.30.65.179", "MSP segment", "Published through 33564 -> 172.30.65.179:22."],
            ["S2Veeam", "172.30.65.180", "MSP segment", "Backup and recovery host."],
            ["C1LAN", "172.30.65.1/26", "OPNsense C1LAN", "Company 1 routed LAN."],
            ["C1DC1", "172.30.65.2", "C1LAN", "Primary Company 1 controller."],
            ["C1DC2", "172.30.65.3", "C1LAN", "Secondary Company 1 controller."],
            ["C1FS", "172.30.65.4", "C1LAN", "Company 1 file server."],
            ["C1WindowsClient", "172.30.65.11", "C1LAN", "Company 1 Windows client."],
            ["C1UbuntuClient", "172.30.65.36", "C1LAN", "Company 1 Linux client."],
            ["C1DMZ", "172.30.65.161/29", "OPNsense C1DMZ", "Company 1 web segment."],
            ["C1WebServer", "172.30.65.162", "C1DMZ", "Company 1 IIS host."],
            ["C1SAN", "172.30.65.186/29", "Gateway 172.30.65.185", "Company 1 isolated storage bridge."],
            ["C2LAN", "172.30.65.65/26", "OPNsense C2LAN", "Company 2 routed LAN."],
            ["C2IdM1", "172.30.65.66", "ens18", "Company 2 primary identity node."],
            ["C2IdM2", "172.30.65.67", "ens18", "Company 2 secondary identity node."],
            ["C2FS service NIC", "172.30.65.68/26", "ens19", "Company 2 file-service interface."],
            ["C2LinuxClient", "172.30.65.70", "ens18", "Company 2 Linux client."],
            ["C2DMZ", "172.30.65.169/29", "OPNsense C2DMZ", "Company 2 web segment."],
            ["C2WebServer", "172.30.65.170", "ens18", "Company 2 nginx host."],
            ["C2SAN", "172.30.65.194/29", "Gateway 172.30.65.193", "Company 2 isolated storage bridge."],
            ["C2FS storage NIC", "172.30.65.195/29", "ens18", "Private storage-facing NIC for C2FS."],
            ["SITE1_VEEAM", "192.168.64.20/32", "Via Site 1 OpenVPN gateway", "Static route target for offsite backup copy."],
        ],
    )

    add_heading(doc, "Appendix B: Evidence and Reference Traceability", 2)
    add_table(
        doc,
        "Table B1. Evidence and reference traceability",
        ["Report area", "Primary evidence", "Reference set", "Confidence reading"],
        [
            ["Network and edge exposure", "OPNsense interface, NAT, alias, and path evidence", "[1]-[3], [9]-[11]", "High"],
            ["Company 1 identity and Windows service stack", "Jump64 WinRM and WMI-backed inspection plus MSP reachability checks", "[7], [10]-[11], [13]-[14]", "High"],
            ["Company 2 identity stack", "Direct node-side service observations and DNS output", "[4]-[6], [10]-[11]", "High"],
            ["Storage and file services", "C1FS and C2FS service-state evidence, share definitions, iSCSI session evidence", "[5]-[6], [10]-[11], [14]", "High"],
            ["Client validation", "C1UbuntuClient, C1WindowsClient, and C2LinuxClient endpoint checks", "[10]-[11], [13]", "High"],
            ["Web delivery", "Hostname and raw-IP HTTP validation plus host-side observations", "[7], [12], [10]-[11]", "High"],
            ["Backup and copy", "S2Veeam reachability, service inventory, repository, job, and copy-path evidence", "[8]-[11]", "High"],
            ["Shared forest interpretation", "DNS visibility, client behavior, and design record", "[4], [10]-[11]", "Medium-High"],
        ],
    )

    add_heading(doc, "Appendix C: Service Verification Matrix", 2)
    add_table(
        doc,
        "Table C1. Service verification and assurance matrix",
        ["Service area", "Validation method", "Observed result", "Assurance level"],
        [
            ["Administrative entry", "WAN NAT and bastion reachability checks", "Both published jump-host paths were present and usable", "High"],
            ["OPNsense management plane", "HTTP and TCP reachability from MSPUbuntuJump", "HTTP 403 on port 80, TCP 53 reachable, TCP 443 timed out", "Medium-High"],
            ["Company 1 directory services", "Jump64 WinRM to C1DC1 and C1DC2", "Domain, forest, controller, and directory-service state returned successfully", "High"],
            ["Company 1 file service", "Jump64 observation of C1FS", "SMB active, F: SharedData present, shares visible, iSCSI session active", "High"],
            ["Company 1 web service", "Jump64 observation and HTTP checks", "Hostname binding on TCP 443; raw IP returned HTTP 404", "High"],
            ["Company 1 client consumption", "WMI-backed Windows check and direct Linux shell check", "Windows and Linux endpoints both reached the two internal web hostnames", "High"],
            ["Company 2 identity services", "Direct node checks on C2IdM1 and C2IdM2", "AD, DNS, and DHCP active; zones and principals visible", "High"],
            ["Company 2 file service", "Direct C2FS checks and client share validation", "Mounted iSCSI-backed storage, active shares, successful client SMB access", "High"],
            ["Company 2 web service", "HTTP validation and host observation", "Hostname returned 200 OK; raw IP returned 404 Not Found", "High"],
            ["Backup and offsite copy", "S2Veeam administrative observation and route or port review", "Repository, job families, copy jobs, and protected workload count aligned with design", "High"],
        ],
    )

    add_heading(doc, "Appendix D: Unresolved Items and Known Gaps", 2)
    add_para(doc, "The final revision pass left three explicit gaps in evidence depth. First, the OPNsense management plane was not walked through with an authenticated GUI session. Second, current Veeam GUI screenshots were not refreshed even though live access to S2Veeam was confirmed from Jump64. Third, C1SAN remained intentionally unmanageable from the bastions, so the storage reading continues to rely on the active C1FS consumer session rather than a direct SAN administration session.")

    add_heading(doc, "Appendix E: Sanitized SMB Configuration Excerpt (C2FS)", 2)
    add_para(doc, "The following excerpt is included verbatim because it defines the active Company 2 share model observed on C2FS.")
    for line in APPENDIX_E_BLOCK:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    add_heading(doc, "6. References", 1)
    refs = [
        '[1] OPNsense Documentation, "Firewall Rules," https://docs.opnsense.org/manual/firewall.html, accessed Mar. 24, 2026.',
        '[2] OPNsense Documentation, "Network Address Translation," https://docs.opnsense.org/manual/nat.html, accessed Mar. 24, 2026.',
        '[3] OPNsense Documentation, "Setup SSL VPN Road Warrior," https://docs.opnsense.org/manual/how-tos/sslvpn_client.html, accessed Mar. 24, 2026.',
        '[4] SambaWiki, "Setting up Samba as an Active Directory Domain Controller," https://wiki.samba.org/index.php/Setting_up_Samba_as_an_Active_Directory_Domain_Controller, accessed Mar. 24, 2026.',
        '[5] Ubuntu Server Documentation, "Set up Samba as a file server," https://documentation.ubuntu.com/server/how-to/samba/file-server/, accessed Mar. 24, 2026.',
        '[6] Ubuntu Server Documentation, "iSCSI initiator (or client)," https://documentation.ubuntu.com/server/how-to/storage/iscsi-initiator-or-client/, accessed Mar. 24, 2026.',
        '[7] Microsoft Learn, "binding Element for bindings for site [IIS Settings Schema]," https://learn.microsoft.com/en-us/previous-versions/iis/settings-schema/ms691267(v=vs.90), accessed Mar. 24, 2026.',
        '[8] Veeam Help Center, "Configuring Backup Repositories," https://helpcenter.veeam.com/docs/vbr/userguide/sch_configure_repository.html, accessed Mar. 24, 2026.',
        '[9] Site 2 gateway configuration record, Mar. 23, 2026.',
        '[10] Site 2 environment inventory, SAN addressing record, and namespace design record, Mar. 24, 2026.',
        '[11] Site 2 operating-state review record, Mar. 23-27, 2026.',
        '[12] nginx Documentation, "Server names," https://nginx.org/en/docs/http/server_names.html, accessed Mar. 27, 2026.',
        '[13] Microsoft Learn, "About WMI," https://learn.microsoft.com/en-us/windows/win32/wmisdk/about-wmi, accessed Mar. 27, 2026.',
        '[14] Microsoft Learn, "SMB features in Windows and Windows Server," https://learn.microsoft.com/en-us/windows-server/storage/file-server/smb-feature-descriptions, accessed Mar. 27, 2026.',
    ]
    for ref in refs:
        add_para(doc, ref)

    return doc, section_34, section_35


def verify_document(doc: Document, section_34: List[str], section_35: List[str]) -> dict:
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    text = "\n".join(parts)
    required_titles = [
        "3.4 Company 1 Directory Services, File Services, Web Delivery, and Client Access",
        "3.5 Company 2 Identity Services, DNS, DHCP, and Shared Forest Design",
        "3.6 Storage, File Services, and Isolated SAN Design",
        "3.9 Requirement-to-Implementation Traceability",
        "3.10 Service Dependencies, Failure Domains, and Access Model",
    ]
    missing_titles = [title for title in required_titles if title not in text]
    required_values = [
        "172.20.64.1/16", "172.30.65.177/29", "172.30.65.1/26", "172.30.65.161/29", "172.30.65.65/26", "172.30.65.169/29",
        "SITE1_OVPN", "172.30.65.186/29", "172.30.65.185", "172.30.65.194/29", "172.30.65.193", "33464", "33564",
        "172.30.65.178", "172.30.65.179", "172.30.65.180", "172.30.65.2", "172.30.65.3", "172.30.65.4", "172.30.65.11",
        "172.30.65.36", "172.30.65.162", "172.30.65.66", "172.30.65.67", "172.30.65.68", "172.30.65.70", "172.30.65.170",
        "ens18", "ens19", "172.30.65.195/29", "c1-webserver.c1.local", "c2-webserver.c2.local", "172.30.64.162", "172.30.64.170",
        "iqn.2024-03.org.clearroots:c2san", "C2_Public", "C2_Private", "Administrator", "admin", "employee1", "employee2", "c2_file_users",
        "VeeamBackupSvc", "VeeamBrokerSvc", "VeeamDeploySvc", "VeeamExplorersRecoverySvc", "VeeamFilesysVssSvc", "VeeamMountSvc", "VeeamNFSSvc", "192.168.64.20/32",
    ]
    missing_values = [value for value in required_values if value not in text]
    return {
        "missing_titles": missing_titles,
        "missing_values": missing_values,
        "wmi_phrase_count": text.count("WMI because TCP 5985 was not open on that host"),
        "figure_count": text.count("[FIGURE PLACEHOLDER - Figure "),
        "table_count": len(doc.tables),
        "s34_sentences": sentence_count(section_34),
        "s35_sentences": sentence_count(section_35),
        "parity_diff": abs(sentence_count(section_34) - sentence_count(section_35)) / max(sentence_count(section_34), sentence_count(section_35)),
    }


def main() -> None:
    OUTDIR.mkdir(parents=True, exist_ok=True)
    doc, section_34, section_35 = build_document()
    doc.save(str(DOCX_PATH))
    reopen = Document(str(DOCX_PATH))
    extracted = []
    for p in reopen.paragraphs:
        if p.text:
            extracted.append(p.text)
    for table in reopen.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    extracted.append(cell.text)
    TEXT_PATH.write_text("\n".join(extracted), encoding="utf-8")
    checks = verify_document(reopen, section_34, section_35)
    pdf_status = export_pdf_if_possible()
    print(f"docx={DOCX_PATH}")
    print(f"pdf={PDF_PATH}")
    print(f"pdf_status={pdf_status}")
    print(f"figure_placeholders={checks['figure_count']}")
    print(f"table_objects={checks['table_count']}")
    print(f"wmi_phrase_count={checks['wmi_phrase_count']}")
    print(f"s34_sentences={checks['s34_sentences']}")
    print(f"s35_sentences={checks['s35_sentences']}")
    print(f"parity_diff={checks['parity_diff']:.3f}")
    print(f"missing_titles={checks['missing_titles']}")
    print(f"missing_values={checks['missing_values']}")


if __name__ == "__main__":
    main()
