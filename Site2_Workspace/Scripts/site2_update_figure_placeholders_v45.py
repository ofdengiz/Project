from pathlib import Path

from docx import Document


ROOT = Path(r"C:\Algonquin\Winter2026\Emerging_Tech\Project")
DOCX = ROOT / "Site2_SourceBased_Technical_Report_V4.5.docx"
GUIDE = ROOT / "Site2_Figure_Placement_Guide_V4.5.md"


PLACEHOLDERS = {
    "Figure 6.": "INSERT UPDATED SCREENSHOT HERE - Figure 6 / C2IdM1 AD-DNS-DHCP evidence",
    "Figure 7.": "INSERT UPDATED SCREENSHOT HERE - Figure 7 / C2IdM2 AD-DNS-DHCP evidence",
    "Figure 8.": "INSERT UPDATED SCREENSHOT HERE - Figure 8 / Shared-forest and cross-domain DNS evidence",
    "Figure 9.": "INSERT UPDATED SCREENSHOT HERE - Figure 9 / C2FS iSCSI and mounted volume evidence",
    "Figure 10.": "INSERT UPDATED SCREENSHOT HERE - Figure 10 / C2FS SMB shares and sync evidence",
    "Figure 13.": "INSERT UPDATED SCREENSHOT HERE - Figure 13 / C1UbuntuClient dual-web evidence",
    "Figure 15.": "INSERT UPDATED SCREENSHOT HERE - Figure 15 / S2Veeam repository, jobs, and offsite-copy evidence",
}


GUIDE_TEXT = """# Site 2 Figure Placement Guide - V4.5

Use this guide with:
- `C:\\Algonquin\\Winter2026\\Emerging_Tech\\Project\\Site2_SourceBased_Technical_Report_V4.5.docx`

Figures that currently need refreshed screenshots:

1. Figure 6
   C2IdM1 Active Directory, DNS, and DHCP evidence
   Capture: active `samba-ad-dc`, active DHCP service, and DNS query results for both web hostnames.

2. Figure 7
   C2IdM2 Active Directory, DNS, and DHCP evidence
   Capture: same style as Figure 6, but from `C2IdM2`.

3. Figure 8
   Shared-forest and cross-domain DNS evidence
   Capture: one view showing both `c1.local` and `c2.local` visibility in the same naming context.

4. Figure 9
   C2FS iSCSI-backed storage and mounted volume evidence
   Capture: active iSCSI session and mounted `/mnt/c2_public` evidence.

5. Figure 10
   C2FS SMB share definitions and synchronization evidence
   Capture: `C2_Public`, `C2_Private`, and successful sync evidence.

6. Figure 13
   C1UbuntuClient Company 1 client dual-web evidence
   Capture: `admin@C1UbuntuClient` prompt, Company 1 realm context, and successful access to both web hostnames.

7. Figure 15
   S2Veeam repository, backup jobs, and offsite-copy evidence
   Capture: refreshed Veeam repository, job inventory, and copy-job views.

Figures currently still present and not marked as missing:
- Figure 11 - C1SAN isolated storage interface evidence
- Figure 12 - C2SAN isolated storage interface evidence
- Figure 14 - C2LinuxClient domain identity and dual-web evidence
"""


def main() -> None:
    doc = Document(str(DOCX))
    for i, para in enumerate(doc.paragraphs):
        for key, replacement in PLACEHOLDERS.items():
            if para.text.startswith(key):
                if i + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[i + 1]
                    if "Screenshot evidence available in supporting evidence set; image pending final formatting." in next_para.text:
                        next_para.text = replacement
                break
    doc.save(str(DOCX))
    GUIDE.write_text(GUIDE_TEXT, encoding="utf-8")


if __name__ == "__main__":
    main()
