from pathlib import Path
import shutil

from docx import Document


ROOT = Path(r"C:\Algonquin\Winter2026\Emerging_Tech\Project")
SRC = ROOT / "Site2_SourceBased_Technical_Report_V4.3.docx"
DST = ROOT / "Site2_SourceBased_Technical_Report_V4.4.docx"


def set_text(paragraph, text: str) -> None:
    paragraph.text = text


def main() -> None:
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    # Title page version bump for the new finalized revision.
    if len(doc.paragraphs) > 8 and doc.paragraphs[8].text.strip() == "4.3":
        set_text(doc.paragraphs[8], "4.4")

    # Clean punctuation artifacts that remained from prior encoding passes.
    replacements = {
        " ? not an accident ? ": ", not an accident, ",
        " ? the documented hostnames": " - the documented hostnames",
        " operational credibility ? it distinguishes ": " operational credibility - it distinguishes ",
        " sections ? network, identity, storage, backup ? describe ": " sections - network, identity, storage, backup - describe ",
    }
    for para in doc.paragraphs:
        text = para.text
        new_text = text
        for old, new in replacements.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            set_text(para, new_text)

    # Remove the stray leftover line in the backup section.
    for para in doc.paragraphs:
        if para.text.strip() == "TCP 10006":
            set_text(para, "")

    # Fix Appendix D ordering and keep all content under the real heading.
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == "Appendix D. Unresolved Items and Known Gaps":
            base = i
            break
    else:
        raise RuntimeError("Appendix D heading not found")

    expected = {
        base + 1: "This appendix distinguishes items that were resolved during the March 27 Jump64 inspection pass from items that remain at partial validation.",
        base + 2: "Resolved since V4.0:",
        base + 3: "C1WindowsClient was revalidated during the March 27 inspection pass via WMI and controlled remote process execution from Jump64. Domain membership, Company 1 DNS usage, and successful resolution of and access to both internal web hostnames were all confirmed. WinRM (TCP 5985) was not open on this host, which is why WMI-backed inspection was used instead. The management method difference is noted in Section 3.4 but does not represent an unresolved gap.",
        base + 4: "C1FS was actively inspected from Jump64 over WinRM. A dedicated F: SharedData volume, named SMB shares, and an active iSCSI consumer session were all observed directly. This closes the earlier gap around Company 1 file and storage chain observability.",
        base + 5: "Remaining items:",
        base + 6: "OPNsense management reachability was confirmed to the point of an HTTP 403 response on port 80 and successful TCP 53 access from MSPUbuntuJump. A fully authenticated GUI walkthrough was not performed in this revision pass. This is an evidence depth limit, not a service failure.",
        base + 7: "C1SAN direct management access remains intentionally blocked. MSPUbuntuJump and Jump64 do not receive routed access to the isolated storage address. The relevant confirmation is the active iSCSI consumer session observed on C1FS, which shows the storage chain is operating as designed. No management session into C1SAN itself is expected or required for normal operations.",
        base + 8: "The Veeam GUI screenshot evidence was not refreshed in this revision pass despite live administrative access being confirmed from Jump64. Updated screenshots showing current repository state, job inventory, and copy-job status should be captured and attached as a figure update before the final submission.",
    }
    for idx, text in expected.items():
        set_text(doc.paragraphs[idx], text)

    # Keep paragraph styling stable around Appendix D labels.
    doc.paragraphs[base + 2].style = "Normal"
    doc.paragraphs[base + 5].style = "Normal"

    doc.save(str(DST))

    # Refresh TOC / LOF / LOT through Word if available.
    try:
        import win32com.client  # type: ignore

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        wdoc = word.Documents.Open(str(DST))
        wdoc.Fields.Update()
        for toc in wdoc.TablesOfContents:
            toc.Update()
        for tof in wdoc.TablesOfFigures:
            tof.Update()
        wdoc.Save()
        wdoc.Close(False)
        word.Quit()
    except Exception:
        pass


if __name__ == "__main__":
    main()
