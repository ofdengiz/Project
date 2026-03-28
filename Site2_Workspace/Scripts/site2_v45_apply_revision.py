from copy import deepcopy
from pathlib import Path
import shutil

from typing import Optional

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


ROOT = Path(r"C:\Algonquin\Winter2026\Emerging_Tech\Project")
SRC = ROOT / "Site2_SourceBased_Technical_Report_V4.4.docx"
DST = ROOT / "Site2_SourceBased_Technical_Report_V4.5.docx"


def set_text(paragraph, text: str) -> None:
    paragraph.text = text


def find_paragraph(doc: Document, needle: str):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    raise ValueError(f"Paragraph not found: {needle}")


def para_index(doc: Document, paragraph) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p._p is paragraph._p:
            return i
    raise ValueError("Paragraph object not found")


def insert_paragraph_after(paragraph, text: str = "", style: Optional[str] = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    return new_para


def insert_row_before(table, before_idx: int):
    template = deepcopy(table.rows[before_idx]._tr)
    table.rows[before_idx]._tr.addprevious(template)
    return table.rows[before_idx]


def main() -> None:
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    # Title version bump.
    if len(doc.paragraphs) > 8 and doc.paragraphs[8].text.strip() == "4.4":
        set_text(doc.paragraphs[8], "4.5")

    # Section 3.7 client list and hostname list cleanup.
    p = find_paragraph(doc, "Three different client perspectives were available:")
    idx = para_index(doc, p)
    set_text(doc.paragraphs[idx + 1], "C1UbuntuClient, identified in the environment evidence as the Company 1 Linux client role")
    doc.paragraphs[idx + 1].style = "List Bullet"
    set_text(doc.paragraphs[idx + 2], "C1WindowsClient, the Company 1 Windows client, observed from Jump64 through WMI and controlled remote process execution")
    doc.paragraphs[idx + 2].style = "List Bullet"
    set_text(doc.paragraphs[idx + 3], "C2LinuxClient, the Company 2 Linux client")
    doc.paragraphs[idx + 3].style = "List Bullet"
    set_text(
        doc.paragraphs[idx + 4],
        "All three client perspectives were able to resolve and consume the required internal web hostnames through name-based access. C1UbuntuClient and C2LinuxClient did so directly from their own shells. C1WindowsClient did so through a Jump64-managed client-side probe because WinRM on the client itself was not available:",
    )
    doc.paragraphs[idx + 4].style = "Normal"
    set_text(doc.paragraphs[idx + 5], "https://c1-webserver.c1.local")
    doc.paragraphs[idx + 5].style = "List Bullet"
    set_text(doc.paragraphs[idx + 6], "https://c2-webserver.c2.local")
    doc.paragraphs[idx + 6].style = "List Bullet"

    # Section 3.6 File-Service State: add C1FS equivalent list after the C2 paragraph.
    anchor = find_paragraph(doc, "Keeping those layers separate is what makes file problems diagnosable")
    p1 = insert_paragraph_after(anchor, "Inspection of C1FS from Jump64 showed:", "Normal")
    p2 = insert_paragraph_after(p1, "Windows SMB service active", "List Bullet")
    p3 = insert_paragraph_after(p2, "F: drive labeled SharedData present as a dedicated data volume", "List Bullet")
    p4 = insert_paragraph_after(p3, "Named SMB shares visible on the F: SharedData volume", "List Bullet")
    p5 = insert_paragraph_after(p4, "Active iSCSI initiator session tied to the Company 1 file-service stack", "List Bullet")
    insert_paragraph_after(
        p5,
        "The same layered reading applies: C1SAN delivers the block device, C1FS mounts it as the F: drive, Windows SMB presents named shares above it, and Company 1 clients consume those shares by name. The file-service architecture is equivalent across both tenants even though the platform stack differs.",
        "Body Text",
    )

    # Section 3.6 Share Presentation: add Company 1 share validation paragraph.
    share_anchor = find_paragraph(doc, "That is valuable because it closes the loop between storage architecture and actual user-visible SMB behavior.")
    insert_paragraph_after(
        share_anchor,
        "The Company 1 file service follows the same separation principle on a Windows platform. C1FS presents its shares from a dedicated F: SharedData volume rather than from a system drive, which keeps user data on a volume that can be managed, resized, or backed up independently of the operating system. The active iSCSI session confirmed on C1FS during the March 27 inspection shows that block storage is arriving from C1SAN through the same isolated-bridge model used on the Company 2 side. The difference is platform: where C2FS exposes Samba shares to Linux and Windows clients over c2fs.c2.local, C1FS exposes Windows SMB shares to Company 1 clients. Both designs keep storage transport invisible to users and keep the troubleshooting path clear: a share problem stays in SMB configuration, a mounted-volume problem stays at the file-service host layer, and a block-transport problem stays in the iSCSI session.",
        "Body Text",
    )

    # Strengthen section prose quality for 3.10 - 3.13.
    set_text(
        doc.paragraphs[410],
        "Dependency tables are only useful when they help an operator think in the same order that the environment fails. In Site 2, access path, identity, storage, and recovery do not break with the same symptoms, and they should not be investigated with the same reflexes. The point of laying out the access model beside the dependency view is to show which identity is supposed to touch which system, which bastion is expected to reach it first, and where a support engineer can trust the first answer they get.",
    )
    doc.paragraphs[410].style = "Normal"
    set_text(
        doc.paragraphs[411],
        "That difference becomes practical very quickly. If a web hostname fails from a client, the right first questions are not the same as they would be for a file-share outage or a Veeam copy problem. One issue begins with DNS, bindings, and client reachability; another begins with the mounted volume or SMB presentation; another begins with repository state and copy paths. The document is stronger when it names those boundaries explicitly, because real support work is mostly the act of choosing the right first boundary to inspect.",
    )
    doc.paragraphs[411].style = "Body Text"
    set_text(
        doc.paragraphs[417],
        "Synchronization and backup serve different purposes, and both need to appear in the handover record. In day-to-day operations they can look deceptively similar because each one touches the same file estate, but they answer different questions. Synchronization keeps working content aligned between sites. Backup exists for the moment when ordinary continuity is no longer enough and a known-good restore point matters more than keeping two locations similar.",
    )
    doc.paragraphs[417].style = "Normal"
    set_text(
        doc.paragraphs[418],
        "That distinction changes operator behavior during an incident. A missing file, a corrupted share, a lost mount, and a broken repository all feel like 'data problems' to the user who reports them, yet they require different first checks and different recovery choices. Writing that difference plainly matters because support teams rarely fail by lacking commands; they fail by using the wrong recovery mechanism too early or by troubleshooting a sync path when the real fault sits in backup, storage, or permissions.",
    )
    doc.paragraphs[418].style = "Body Text"
    set_text(
        doc.paragraphs[423],
        "These checks are presented in control-plane order rather than as a loose administrator checklist because that is how the site has to be defended in practice. When an operator starts the day, the first meaningful question is whether the bastions still give controlled entry. The next questions are whether routing and policy still carry the intended paths, whether names still resolve from the authoritative systems, whether storage still presents the data that users expect, and finally whether recovery remains trustworthy if the live service layer has already failed.",
    )
    doc.paragraphs[423].style = "Body Text"
    set_text(
        doc.paragraphs[424],
        "That sequence reflects lived support experience more than documentation style. Healthy environments are usually not lost all at once; they drift. A resolver path drops out, an iSCSI session disappears, a web binding is changed, or a copy job falls behind. Routine maintenance has value precisely because it catches that drift while the fix is still small. Framing the checks this way turns the table from a generic duty list into a practical reading order for the whole site.",
    )
    doc.paragraphs[424].style = "Body Text"
    set_text(
        doc.paragraphs[427],
        "Table 19 summarizes the primary fault domains and first-line checks for common service symptoms. Its real value is speed: when someone reports a problem under demo or support pressure, the table reduces the temptation to check everything at once and instead points toward the first authoritative system that can narrow the fault.",
    )
    doc.paragraphs[427].style = "Normal"
    set_text(
        doc.paragraphs[429],
        "Good triage is less about having a long command list and more about preserving confidence while uncertainty is highest. A support engineer who can tell the difference between an entry-path failure, a naming failure, a storage-presentation failure, and a recovery-layer issue will usually solve the problem faster than one who collects more output from the wrong systems. The triage table is meant to encode that judgment, so the document reads like an operating guide rather than a static inventory.",
    )
    doc.paragraphs[429].style = "Body Text"

    # Table 14: include C1WindowsClient in client-side service consumption row.
    table14 = doc.tables[14]
    table14.rows[9].cells[1].text = "C1UbuntuClient, C1WindowsClient, and C2LinuxClient resolving and reaching both required hostnames"
    table14.rows[9].cells[2].text = "Direct Linux client-side web and identity checks plus Jump64-managed Windows client probe"
    table14.rows[9].cells[4].text = "Client observations confirm that named services are consumable from both company contexts across Linux and Windows endpoints"

    # Table 15: rename C2 row and add Company 1 storage row.
    table15 = doc.tables[15]
    table15.rows[4].cells[0].text = "Company 2 file and storage"
    row = insert_row_before(table15, 5)
    values = [
        "Company 1 file and storage",
        "C1FS, C1SAN, iSCSI session, SMB share configuration",
        "Company 1 share access and file-service availability",
        "Shares unavailable or C1FS iSCSI session gone",
        "Jump64 WinRM to C1FS: Get-SmbShare, Get-IscsiSession, Get-Volume",
    ]
    for i, val in enumerate(values):
        row.cells[i].text = val

    # Table 17: add Company 1 SMB rows.
    table17 = doc.tables[17]
    row = insert_row_before(table17, 5)
    values = [
        "Company 1 public shares",
        "Local file operations on C1FS",
        "SMB over C1LAN",
        "Named share on F: SharedData volume",
        "Veeam C1_FileShare backup job",
        "File-level restore from Veeam without requiring full-host recovery",
    ]
    for i, val in enumerate(values):
        row.cells[i].text = val
    row = insert_row_before(table17, 6)
    values = [
        "Company 1 Windows file shares",
        "C1SAN iSCSI transport",
        "Mounted to C1FS F: SharedData",
        "SMB share presented to C1LAN clients",
        "Veeam file-share backup plus VM backup",
        "Both the share layer and the underlying host are protected independently",
    ]
    for i, val in enumerate(values):
        row.cells[i].text = val

    # Table 18: add Company 1 maintenance checks after OPNsense row.
    table18 = doc.tables[18]
    new_rows = [
        [
            "Confirm C1DC1 and C1DC2 are reachable from Jump64 on WinRM (TCP 5985) and that NTDS, DNS, KDC, and Netlogon services are running",
            "Company 1 domain controller health is the foundation of all Company 1 identity, authentication, and name resolution; if either DC is unhealthy, client logon and hostname resolution will degrade",
        ],
        [
            "Confirm C1FS F: SharedData volume is mounted and SMB shares are accessible; confirm active iSCSI session to C1SAN",
            "Storage transport issues on the Company 1 file service will appear as share unavailability; separating volume, share, and iSCSI checks makes the fault domain immediately visible",
        ],
        [
            "Confirm C1WebServer IIS service is running and c1-webserver.c1.local returns HTTP 200 while the raw IP returns 404",
            "Confirms that Company 1 web delivery continues to follow the hostname-only publication model and that IIS has not been reconfigured to answer raw IP requests",
        ],
    ]
    insert_at = 3
    for vals in new_rows:
        row = insert_row_before(table18, insert_at)
        row.cells[0].text = vals[0]
        row.cells[1].text = vals[1]
        insert_at += 1

    # Table 20 systems administration row.
    table20 = doc.tables[20]
    table20.rows[3].cells[1].text = "Identity, client access, file services, storage transport, Windows-side inspection, and backup paths are aligned with the documented operating model across both tenant scopes"
    table20.rows[3].cells[2].text = "Samba AD and DHCP state, DNS records, three-client web access validation, C2FS and C1FS storage checks, Jump64 WinRM sessions confirming C1DC1/C1DC2/C1FS/C1WebServer/C1WindowsClient service state, and S2Veeam administrative access from Jump64"

    # Table B1 rows for C1-inclusive evidence basis.
    table22 = doc.tables[22]
    table22.rows[4].cells[1].text = "Ubuntu Server Samba documentation plus C2FS service-state evidence; Windows Server SMB documentation plus C1FS share and volume evidence from Jump64 WinRM inspection"
    table22.rows[5].cells[1].text = "Ubuntu Server iSCSI documentation, SAN addressing evidence, and C2FS session state; Windows iSCSI initiator evidence from C1FS Get-IscsiSession output observed via Jump64"

    # Appendix C body / Table C1 rows updated to V4.4 state.
    table23 = doc.tables[23]
    rows = [
        ["Administrative entry and bastion access", "Controlled remote-access verification and jump-host reachability checks", "Both jump systems were reachable and remained the intended administrative entry points", "High"],
        ["Segmented networking and limited edge exposure", "OPNsense configuration review plus management-path checks", "Interfaces, aliases, NAT publication limits, and service-specific rules were all evidenced", "High"],
        ["Company 1 services", "MSP-bastion checks, Jump64 remoting or WMI inspection, hostname review, DNS visibility, and client access checks", "Company 1 directory, file, web, client, and SAN roles were all revalidated, with direct Windows-side observation from Jump64 now covering the previously thin areas", "High"],
        ["Company 2 identity, DNS, and DHCP", "Service-state checks and DNS queries on C2IdM1 and C2IdM2", "Samba AD, DHCP, and required hostname records were all present and consistent", "High"],
        ["Dual-hostname web delivery", "Client-side and bastion-side HTTPS checks by hostname and by raw IP", "Both required hostnames returned successful responses while raw IP access reflected hardened behavior", "High"],
        ["File services and share isolation", "C2FS service checks, mount review, testparm, sync-log review, and hostname-based SMB validation; C1FS inspected via Jump64 WinRM", "SMB service state, mounted storage, shares, synchronization, and named share access confirmed for C2FS; F: SharedData volume, named shares, and active iSCSI session confirmed for C1FS", "High"],
        ["Isolated SAN transport", "SAN addressing evidence plus iSCSI session review on C2FS and C1FS", "Storage transport confirmed as isolated from tenant LAN segments and correctly consumed by both file-service layers", "High"],
        ["Backup and offsite protection", "Veeam host reachability, Jump64 administrative session, repository evidence, job inventory, copy-job inventory, route or rule review, and backup-design evidence", "Protection architecture, local backup handling, file-share backup scope, offsite-copy design, and Windows administrative access to S2Veeam were all consistent with the documented operating model", "High"],
        ["Shared-forest interpretation", "DNS visibility, client behavior, and confirmed namespace relationship", "Cross-domain naming behavior is consistent with the documented forest relationship", "Medium-High"],
        ["Administrative Linux SSH access", "Administrative-path testing within the documented operating scope", "Managed SSH access paths aligned with the documented operating scope", "Medium"],
    ]
    for r_idx, vals in enumerate(rows, start=1):
        while r_idx >= len(table23.rows):
            table23.add_row()
        for c_idx, val in enumerate(vals):
            table23.rows[r_idx].cells[c_idx].text = val

    doc.save(str(DST))

    # Refresh fields and lists through Word if present.
    try:
        import win32com.client  # type: ignore

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        wdoc = word.Documents.Open(str(DST))
        wdoc.Fields.Update()
        for toc in wdoc.TablesOfContents:
            toc.Update()
        for tof in wdoc.TablesOfFigures:
            tof.Update()
        wdoc.Save()
        wdoc.Close(False)
        word.Quit()
    except Exception:
        pass


if __name__ == "__main__":
    main()
