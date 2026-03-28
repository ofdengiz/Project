from pathlib import Path
import shutil
from typing import Optional

from docx import Document
from docx.shared import Inches


ROOT = Path(r"C:\Algonquin\Winter2026\Emerging_Tech\Project")
SRC = ROOT / "Site2_SourceBased_Technical_Report_V4.6.docx"
DST = ROOT / "Site2_SourceBased_Technical_Report_V4.7.docx"
IMGDIR = ROOT / "Site2_CLI_Evidence_2026-03-27_193749"


FIGURES = [
    (
        "Figure 5A. Jump64 Windows inspection baseline.\nDescription: This figure shows the Jump64 Windows bastion baseline, including platform state and its internal Site 2 management address.",
        IMGDIR / "Jump64_Baseline.png",
    ),
    (
        "Figure 5B. C1DC1 service-state evidence.\nDescription: This figure shows C1DC1 service-state output from Jump64, confirming the active Company 1 directory stack on the Windows management path.",
        IMGDIR / "Jump64_C1DC1_ServiceState.png",
    ),
    (
        "Figure 5C. C1DC2 service-state evidence.\nDescription: This figure shows C1DC2 service-state output from Jump64, reinforcing the dual-controller operating model on the Company 1 side.",
        IMGDIR / "Jump64_C1DC2_ServiceState.png",
    ),
    (
        "Figure 5D. C1FS storage, share, and iSCSI evidence.\nDescription: This figure shows the C1FS F: SharedData volume, named SMB shares, and active iSCSI session as observed from Jump64.",
        IMGDIR / "Jump64_C1FS_StorageShares.png",
    ),
    (
        "Figure 5E. C1WebServer IIS binding evidence.\nDescription: This figure shows the workgroup-hosted C1WebServer state and the IIS binding restricted to c1-webserver.c1.local on TCP 443.",
        IMGDIR / "Jump64_C1WebServer_IIS.png",
    ),
    (
        "Figure 5F. C1WindowsClient endpoint and dual-web evidence.\nDescription: This figure shows the Company 1 Windows client validated from Jump64, including successful resolution and consumption of both required internal web hostnames.",
        IMGDIR / "Jump64_C1WindowsClient_WebProbe.png",
    ),
]

def main() -> None:
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    target = None
    for para in doc.paragraphs:
        if para.text.strip() == "3.5 Company 2 Identity, Shared Forest Context, DNS, and DHCP":
            target = para
            break
    if target is None:
        raise RuntimeError("Target heading for 3.5 not found")

    # Insert in reverse order so final document preserves 5A -> 5F.
    for caption, img_path in reversed(FIGURES):
        img_para = target.insert_paragraph_before("")
        try:
            img_para.style = "Body Text"
        except Exception:
            pass
        if img_path.exists():
            run = img_para.add_run()
            run.add_picture(str(img_path), width=Inches(6.5))
        cap_para = target.insert_paragraph_before(caption)
        try:
            cap_para.style = "Body Text"
        except Exception:
            pass

    doc.save(str(DST))


if __name__ == "__main__":
    main()
