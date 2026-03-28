from copy import copy, deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.shared import Inches

ROOT = Path(r"C:\Algonquin\Winter2026\Emerging_Tech\Project")
SRC = ROOT / "Site2_SourceBased_Technical_Report_V4.2.docx"
DST = ROOT / "Site2_SourceBased_Technical_Report_V4.3.docx"


def find_paragraph(doc: Document, needle: str):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    raise ValueError(f"Paragraph not found: {needle}")


def find_heading(doc: Document, text: str, style_name: str):
    for p in doc.paragraphs:
        if p.text == text and p.style and p.style.name == style_name:
            return p
    raise ValueError(f"Heading not found: {text}")


def para_index(doc: Document, paragraph) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p._p is paragraph._p:
            return i
    raise ValueError("Paragraph object not found")


def set_plain_text(paragraph, text: str):
    paragraph.text = text


def set_bold_label(paragraph, text: str):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = True


def set_style(paragraph, style_name: str):
    try:
        paragraph.style = style_name
    except Exception:
        pass


def insert_row_before(table, before_idx: int):
    template = deepcopy(table.rows[before_idx]._tr)
    table.rows[before_idx]._tr.addprevious(template)
    return table.rows[before_idx]


def main():
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    # Fix 1: title page version and bold labels
    set_plain_text(doc.paragraphs[8], "4.3")
    for idx in [3, 5, 7, 9, 11, 13, 15, 17, 19]:
        set_bold_label(doc.paragraphs[idx], doc.paragraphs[idx].text)

    # Fix 2: remove {.mark} artifacts if any remain
    replacements = {
        "3.4 [Company 1 Services]{.mark}": "3.4 Company 1 Services",
        "3.[5 Company 2 Identity, Shared Forest Context, DNS, and DHCP]{.mark}": "3.5 Company 2 Identity, Shared Forest Context, DNS, and DHCP",
        "## 3.4 [Company 1 Services]{.mark}": "## 3.4 Company 1 Services",
        "[C1WindowsClient at 172.30.65.11]{.mark}": "C1WindowsClient at 172.30.65.11",
        "[C1UbuntuClient at 172.30.65.36, aligned to the active Company 1 Linux client role]{.mark}": "C1UbuntuClient at 172.30.65.36, aligned to the active Company 1 Linux client role",
        "services[, internal web services]{.mark}, and isolated storage": "services, internal web services, and isolated storage",
        "## 3.[5 Company 2 Identity, Shared Forest Context, DNS, and DHCP]{.mark}": "## 3.5 Company 2 Identity, Shared Forest Context, DNS, and DHCP",
        "{.mark}": "",
    }
    for p in doc.paragraphs:
        text = p.text
        new_text = text
        for old, new in replacements.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            p.text = new_text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                new_text = text
                for old, new in replacements.items():
                    new_text = new_text.replace(old, new)
                if new_text != text:
                    cell.text = new_text

    # Fix 3: relocate C1UbuntuClient hostname paragraph
    old_para = find_paragraph(doc, "One useful example is the Company 1 Linux client role.")
    moved_text = old_para.text
    set_plain_text(old_para, "Table 5 maps every Site 2 system to its documented service role across all three scopes.")
    set_style(old_para, "Body Text")
    target_para = find_paragraph(doc, "The final Site 2 state also reflects a small number of deliberate service refinements")
    inserted = target_para.insert_paragraph_before(moved_text)
    set_style(inserted, "Body Text")

    # Fix 4: Table 5A rename and expansion
    p_195 = find_paragraph(doc, "Beyond role mapping, the current platform baseline is useful because it shows that the Linux service nodes were sized")
    p_196 = doc.paragraphs[para_index(doc, p_195) + 1]
    p_197 = doc.paragraphs[para_index(doc, p_195) + 2]
    set_plain_text(
        p_195,
        "Beyond the full site inventory, the Linux node baseline across all three service scopes shows that the platforms were shaped to match their roles rather than built to a generic template. MSP contributes one Linux bastion with no storage role. Company 1 contributes one Linux client endpoint validated for domain context and web access. Company 2 uses the densest Linux stack: two identity controllers, one dual-interface file server, one client, and one focused web application host. Those differences are worth documenting because they explain later choices about where to perform identity checks, why storage complexity concentrates on C2FS, and why the client validation was done from a full desktop endpoint rather than a headless node."
    )
    set_plain_text(
        p_196,
        "Platform shape matters in a handover document because it explains decisions that would otherwise look arbitrary. A single bastion host in the MSP scope with no storage role is the right choice when the purpose is controlled inspection. A dual-controller identity pair with matching CPU and memory in the Company 2 scope is easier to reason about for failover than a single node. A file-service VM with a dedicated data disk is easier to defend and troubleshoot than one that stores user data in its root filesystem. And a full Ubuntu desktop client is more representative for end-to-end validation than a headless server node."
    )
    set_plain_text(p_197, "Table 5A. Observed Linux node platform baseline - all service scopes")

    table5a = doc.tables[5]
    row = insert_row_before(table5a, 1)
    values = [
        "MSPUbuntuJump",
        "Ubuntu Linux (exact version not queried in final pass)",
        "Not directly observed",
        "Not directly observed",
        "Not directly observed",
        "172.30.65.179/29 on MSP segment",
        "MSP Linux bastion; primary CLI inspection path into Company 1 and Company 2 Linux systems; no storage or tenant service role",
    ]
    for i, v in enumerate(values):
        row.cells[i].text = v
    row = insert_row_before(table5a, 2)
    values = [
        "C1UbuntuClient",
        "Ubuntu 25.04",
        "4",
        "7.3 GiB",
        "32 GB root disk plus 3.8 GiB swap",
        "172.30.65.36/26 on C1LAN",
        "Company 1 Linux client endpoint; validated Company 1 realm membership, resolver state, and dual-hostname web access; hardware baseline captured during the March inspection pass",
    ]
    for i, v in enumerate(values):
        row.cells[i].text = v

    # Fix 5: Table 10 company 1 storage rows
    table10 = doc.tables[10]
    arch_idx = len(table10.rows) - 1
    rows_to_add = [
        ("C1FS service interface", "172.30.65.4 on C1LAN"),
        ("C1FS dedicated data volume", "F: drive labeled SharedData, observed from Jump64 WinRM session"),
        ("C1FS SMB shares", "Named shares present on the F: SharedData volume, observed from Jump64"),
        ("Active Company 1 iSCSI session", "Active iSCSI consumer session confirmed on C1FS, tied to the Company 1 file-service stack, observed from Jump64"),
        ("Company 1 share access model", "SMB shares backed by isolated SAN storage via C1SAN; block transport separated from C1LAN segment"),
    ]
    for item, evidence in rows_to_add:
        row = insert_row_before(table10, arch_idx)
        row.cells[0].text = item
        row.cells[1].text = evidence
        arch_idx += 1
    table10.rows[-1].cells[1].text = (
        "Two isolated storage bridges support both tenant file-service domains. C2FS presents Samba shares backed by C2SAN iSCSI transport. "
        "C1FS presents Windows SMB shares backed by C1SAN iSCSI transport. Both operate independently without exposing block traffic to tenant LAN or DMZ segments."
    )

    # Fix 6: add C1WindowsClient column to Table 11
    table11 = doc.tables[11]
    table11.add_column(Inches(2.2))
    new_col_values = [
        "C1WindowsClient",
        "Company 1 Windows client role",
        "c1.local / C1.LOCAL",
        "Domain membership confirmed; Company 1 DNS use confirmed via WMI and remote process execution from Jump64",
        "Company 1 DNS resolvers active; c1.local and c2.local hostnames resolved successfully",
        "Resolved and returned HTTP 200 via Jump64-managed client-side probe",
        "Resolved and returned HTTP 200 via Jump64-managed client-side probe",
        "Not directly validated from C1WindowsClient; Company 1 share access confirmed via C1FS inspection",
        "Shows Company 1 named-service contract holds on a native Windows endpoint, not only from the Linux client",
    ]
    # move current C2LinuxClient column to last and place new col in between
    for r, row in enumerate(table11.rows):
        row.cells[3].text = row.cells[2].text
        row.cells[2].text = new_col_values[r]

    # Fix 7: Appendix D replacement
    appendix_d = find_heading(doc, "Appendix D. Unresolved Items and Known Gaps", "Heading 2")
    idx_d = para_index(doc, appendix_d)
    p_intro = doc.paragraphs[idx_d + 1]
    p1 = doc.paragraphs[idx_d + 2]
    p2 = doc.paragraphs[idx_d + 3]
    p3 = doc.paragraphs[idx_d + 4]
    set_plain_text(p_intro, "This appendix distinguishes items that were resolved during the March 27 Jump64 inspection pass from items that remain at partial validation.")
    set_style(p_intro, "Body Text")
    set_plain_text(p1, "Resolved since V4.0:")
    set_style(p1, "Normal")
    set_plain_text(p2, "C1WindowsClient was revalidated during the March 27 inspection pass via WMI and controlled remote process execution from Jump64. Domain membership, Company 1 DNS usage, and successful resolution of and access to both internal web hostnames were all confirmed. WinRM (TCP 5985) was not open on this host, which is why WMI-backed inspection was used instead. The management method difference is noted in Section 3.4 but does not represent an unresolved gap.")
    set_style(p2, "Body Text")
    set_plain_text(p3, "C1FS was actively inspected from Jump64 over WinRM. A dedicated F: SharedData volume, named SMB shares, and an active iSCSI consumer session were all observed directly. This closes the earlier gap around Company 1 file and storage chain observability.")
    set_style(p3, "Body Text")
    appendix_e = find_heading(doc, "Appendix E. Sanitized SMB Configuration Excerpt", "Heading 2")
    q4 = appendix_e.insert_paragraph_before("The Veeam GUI screenshot evidence was not refreshed in this revision pass despite live administrative access being confirmed from Jump64. Updated screenshots showing current repository state, job inventory, and copy-job status should be captured and attached as a figure update before the final submission.")
    q3 = appendix_e.insert_paragraph_before("C1SAN direct management access remains intentionally blocked. MSPUbuntuJump and Jump64 do not receive routed access to the isolated storage address. The relevant confirmation is the active iSCSI consumer session observed on C1FS, which shows the storage chain is operating as designed. No management session into C1SAN itself is expected or required for normal operations.")
    q2 = appendix_e.insert_paragraph_before("OPNsense management reachability was confirmed to the point of an HTTP 403 response on port 80 and successful TCP 53 access from MSPUbuntuJump. A fully authenticated GUI walkthrough was not performed in this revision pass. This is an evidence depth limit, not a service failure.")
    q1 = appendix_e.insert_paragraph_before("Remaining items:")
    for para in [q1, q2, q3, q4]:
        set_style(para, "Body Text" if para is not q1 else "Normal")

    # Fix 8: Section 3.7 bullet formatting and hostname list
    p335 = find_paragraph(doc, "Three different client perspectives were available:")
    idx = para_index(doc, p335)
    set_style(doc.paragraphs[idx + 1], "List Bullet")
    set_plain_text(doc.paragraphs[idx + 1], "C1UbuntuClient, identified in the environment evidence as the Company 1 Linux client role")
    set_style(doc.paragraphs[idx + 2], "List Bullet")
    set_plain_text(doc.paragraphs[idx + 2], "C1WindowsClient, the Company 1 Windows client, observed from Jump64 through WMI and controlled remote process execution")
    set_style(doc.paragraphs[idx + 3], "List Bullet")
    set_plain_text(doc.paragraphs[idx + 3], "C2LinuxClient, the Company 2 Linux client")
    set_style(doc.paragraphs[idx + 4], "Normal")
    set_plain_text(doc.paragraphs[idx + 4], "All three client perspectives were able to resolve and consume the required internal web hostnames through name-based access. C1UbuntuClient and C2LinuxClient did so directly from their own shells. C1WindowsClient did so through a Jump64-managed client-side probe because WinRM on the client itself was not available.")
    new_url = doc.paragraphs[idx + 5].insert_paragraph_before("https://c1-webserver.c1.local")
    set_style(new_url, "List Bullet")
    set_plain_text(doc.paragraphs[idx + 5], "https://c2-webserver.c2.local")
    set_style(doc.paragraphs[idx + 5], "List Bullet")

    # Fix 9: Table 12 actual client rows
    table12 = doc.tables[12]
    table12.rows[5].cells[0].text = "C1UbuntuClient to both hostnames (direct curl from shell)"
    table12.rows[5].cells[1].text = "Success"
    table12.rows[6].cells[0].text = "C1WindowsClient to both hostnames (Jump64-managed probe)"
    table12.rows[6].cells[1].text = "Success"
    row = table12.add_row()
    row.cells[0].text = "C2LinuxClient to both hostnames (direct curl from shell)"
    row.cells[1].text = "Success"

    # Fix 10: mark caption-only figures explicitly
    figure_notes = {
        "Figure 6.": "Screenshot evidence available in supporting evidence set; image pending final formatting.",
        "Figure 7.": "Screenshot evidence available in supporting evidence set; image pending final formatting.",
        "Figure 8.": "Screenshot evidence available in supporting evidence set; image pending final formatting.",
        "Figure 9.": "Screenshot evidence available in supporting evidence set; image pending final formatting.",
        "Figure 10.": "Screenshot evidence available in supporting evidence set; image pending final formatting.",
        "Figure 13.": "Screenshot evidence available in supporting evidence set; image pending final formatting.",
        "Figure 15.": "Screenshot evidence available in supporting evidence set; image pending final formatting.",
    }
    for p in doc.paragraphs:
        for key, note in figure_notes.items():
            if p.text.startswith(key) and p.style and p.style.name == "Body Text":
                idxp = para_index(doc, p)
                # use immediate following blank paragraph if present
                if idxp + 1 < len(doc.paragraphs) and doc.paragraphs[idxp + 1].text == "":
                    set_plain_text(doc.paragraphs[idxp + 1], note)
                    set_style(doc.paragraphs[idxp + 1], "Body Text")
                else:
                    q = p.insert_paragraph_before(note)
                    set_style(q, "Body Text")
                break

    # Update title text in list of tables/table captions if static
    for p in doc.paragraphs:
        if p.text == "Table 5A. Observed Linux VM platform baseline":
            p.text = "Table 5A. Observed Linux node platform baseline - all service scopes"
        if p.text == "Table 10. Storage and isolated SAN summary":
            p.text = "Table 10. Storage and isolated SAN summary"

    doc.save(str(DST))

    # Refresh Word fields if available
    try:
        import win32com.client  # type: ignore

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        wdoc = word.Documents.Open(str(DST))
        wdoc.Fields.Update()
        for toc in wdoc.TablesOfContents:
            toc.Update()
        for tof in wdoc.TablesOfFigures:
            tof.Update()
        wdoc.Save()
        wdoc.Close()
        word.Quit()
    except Exception:
        pass


if __name__ == "__main__":
    main()
