from dataclasses import dataclass
from pathlib import Path
import importlib.util
import re
import shutil
from typing import Optional

from pypdf import PdfReader
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"C:\Algonquin\Winter2026\Emerging_Tech\Project")
SOURCE_SCRIPT = ROOT / "Site2_Workspace" / "Scripts" / "build_site2_report_v60.py"
PDF_PATH = ROOT / "output" / "pdf" / "Site2_Technical_Report_V6.0.pdf"
OUT_DIR = ROOT / "output" / "doc"
FINAL_DIR = ROOT / "Site2_Deliverables" / "Final"
OUT_DOCX = OUT_DIR / "Site2_Technical_Report_V6.0.docx"
FINAL_DOCX = FINAL_DIR / "Site2_Technical_Report_V6.0.docx"
EXEC_SUMMARY_LEAD = "Site 2 is the second operating location in our two-site managed service environment."


@dataclass
class CatalogEntry:
    text: str
    level: int = 0
    page: Optional[int] = None


def load_source_module():
    spec = importlib.util.spec_from_file_location("site2v60", SOURCE_SCRIPT)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def ensure_style(doc, name, style_type, base_style=None):
    if name in doc.styles:
        style = doc.styles[name]
    else:
        style = doc.styles.add_style(name, style_type)
    if base_style:
        style.base_style = doc.styles[base_style]
    return style


def set_style_font(style, name, size, bold=False, italic=False, color=None):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    if color is not None:
        style.font.color.rgb = color



def add_field(paragraph, instruction, placeholder=""):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    if placeholder:
        run._r.append(OxmlElement("w:t"))
        run._r[-1].text = placeholder
    run._r.append(end)



def set_run_font(run, name="Times New Roman", size=11, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color



def parse_markup_runs(paragraph, raw_text, default_size=11):
    parts = re.split(r"(<b>.*?</b>)", raw_text)
    for part in parts:
        if not part:
            continue
        if part.startswith("<b>") and part.endswith("</b>"):
            run = paragraph.add_run(part[3:-4])
            set_run_font(run, size=default_size, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=default_size)



def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    set_style_font(normal, "Times New Roman", 11)

    set_style_font(doc.styles["Title"], "Times New Roman", 24, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    set_style_font(doc.styles["Heading 1"], "Times New Roman", 16, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    set_style_font(doc.styles["Heading 2"], "Times New Roman", 13, bold=True, color=RGBColor(0x2F, 0x55, 0x97))
    set_style_font(doc.styles["Heading 3"], "Times New Roman", 11.5, bold=True, color=RGBColor(0x2F, 0x55, 0x97))

    cover_sub = ensure_style(doc, "CoverSubtitle", 1, "Normal")
    set_style_font(cover_sub, "Times New Roman", 14, italic=True, color=RGBColor(0x44, 0x44, 0x44))

    front_heading = ensure_style(doc, "FrontHeadingDoc", 1, "Normal")
    set_style_font(front_heading, "Times New Roman", 16, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))

    caption = doc.styles["Caption"]
    set_style_font(caption, "Times New Roman", 10.5, bold=True)

    code = ensure_style(doc, "CodeBlockDoc", 1, "Normal")
    set_style_font(code, "Courier New", 9)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Raspberry Pioneers | Site 2 Technical Report V6.0 | Page ")
    add_field(footer, "PAGE")



def plain_text(flowable):
    if hasattr(flowable, "getPlainText"):
        return flowable.getPlainText()
    return getattr(flowable, "text", "")



def walk_flowables(sequence):
    for item in sequence:
        name = type(item).__name__
        if name == "KeepTogether":
            yield from walk_flowables(item._content)
        else:
            yield item



def collect_catalog_from_source():
    source = SOURCE_SCRIPT.read_text(encoding="utf-8")

    headings = [
        CatalogEntry(text=match.group(2), level=int(match.group(1)))
        for match in re.finditer(r'add_heading\(story, s, (\d), "([^"]+)"\)', source)
    ]
    figures = [
        CatalogEntry(text=f"{match.group(1)}. {match.group(2)}")
        for match in re.finditer(r'add_figure\(story, s, "([^"]+)", "([^"]+)",', source)
    ]
    tables = [
        CatalogEntry(text=match.group(1))
        for match in re.finditer(r'add_table\(story, s, "([^"]+)",', source)
    ]
    return headings, figures, tables



def normalize(text):
    text = text.replace("\xa0", " ")
    text = text.replace("\u2011", "-")
    text = text.replace("\u2013", "-")
    text = text.replace("\u2014", "-")
    text = re.sub(r"\s+", " ", text)
    return text.strip()



def extract_pdf_page_texts(pdf_path):
    reader = PdfReader(str(pdf_path))
    return [normalize(page.extract_text() or "") for page in reader.pages]



def find_body_start(page_texts):
    lead = normalize(EXEC_SUMMARY_LEAD)
    for idx, text in enumerate(page_texts):
        if lead in text:
            return idx
    return 0



def candidate_needles(text):
    normalized = normalize(text)
    needles = [normalized]
    if ". " in normalized:
        label, remainder = normalized.split(". ", 1)
        needles.append(label)
        needles.append(f"{label}. {remainder[:100]}")
    needles.append(normalized[:120])
    seen = []
    for needle in needles:
        needle = needle.strip()
        if needle and needle not in seen:
            seen.append(needle)
    return seen



def assign_pages(entries, page_texts, start_index):
    assigned = []
    cursor = start_index
    for entry in entries:
        found = None
        for needle in candidate_needles(entry.text):
            for page_index in range(cursor, len(page_texts)):
                if needle in page_texts[page_index]:
                    found = page_index
                    break
            if found is not None:
                break
        if found is None:
            found = cursor
        cursor = max(cursor, found)
        assigned.append(CatalogEntry(text=entry.text, level=entry.level, page=found + 1))
    return assigned



def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)



def convert_table(doc, flowable):
    rows = len(flowable._cellvalues)
    cols = max(len(row) for row in flowable._cellvalues)
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER if rows == 1 and cols == 1 else WD_TABLE_ALIGNMENT.LEFT

    placeholder = rows == 1 and cols == 1 and "[FIGURE PLACEHOLDER -" in plain_text(flowable._cellvalues[0][0])

    for r_idx, row in enumerate(flowable._cellvalues):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if placeholder else WD_ALIGN_PARAGRAPH.LEFT
            text = plain_text(value)
            run = p.add_run(text)
            if r_idx == 0 and not placeholder:
                set_run_font(run, size=8.75, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
                shade_cell(cell, "1F4E79")
            else:
                set_run_font(run, size=10 if placeholder else 8.5, italic=placeholder)
                if placeholder:
                    shade_cell(cell, "F3F3F3")
    doc.add_paragraph()



def convert_preformatted(doc, flowable):
    paragraph = doc.add_paragraph(style="CodeBlockDoc")
    for idx, line in enumerate(flowable.lines):
        if idx:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, name="Courier New", size=9)


def convert_image(doc, flowable):
    image_path = getattr(flowable, "filename", None) or getattr(flowable, "_file", None)
    if not image_path:
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    width_inches = max(1.0, float(getattr(flowable, "drawWidth", 432)) / 72.0)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def derived_toc_level(text):
    if text == "Executive Summary":
        return 0
    if text.startswith("Appendix "):
        return 1
    if re.match(r"^\d+\.\d+\.\d+\s", text):
        return 2
    if re.match(r"^\d+\.\d+\s", text):
        return 1
    if re.match(r"^\d+\.\s", text):
        return 0
    return 0


def normalize_toc_paragraphs(doc):
    toc_start = None
    lof_start = None
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text == "Table of Contents":
            toc_start = idx
        elif text == "List of Figures":
            lof_start = idx
            break
    if toc_start is None or lof_start is None:
        return
    for paragraph in doc.paragraphs[toc_start + 1:lof_start]:
        raw = paragraph.text.strip()
        if not raw:
            continue
        title = raw.split("	", 1)[0].strip()
        paragraph.paragraph_format.left_indent = Inches(0.22 * derived_toc_level(title))
        paragraph.paragraph_format.first_line_indent = Inches(0)


class Converter:
    def __init__(self, doc, toc_entries, figure_entries, table_entries):
        self.doc = doc
        self.toc_entries = toc_entries
        self.figure_entries = figure_entries
        self.table_entries = table_entries
        self.front_heading = None

    def add_front_heading(self, text):
        p = self.doc.add_paragraph(style="FrontHeadingDoc")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run_font(run, size=16, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
        self.front_heading = text

    def add_body_paragraph(self, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        set_run_font(run)

    def add_reference_paragraph(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        run = p.add_run(text)
        set_run_font(run, size=10.5)

    def add_index_line(self, text, page, level=0, base_size=10.5):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.22 * level)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        title_run = p.add_run(text)
        set_run_font(title_run, size=max(9.5, base_size - (0.5 * min(level, 2))))
        page_run = p.add_run(f"\t{page}")
        set_run_font(page_run, size=max(9.5, base_size - (0.5 * min(level, 2))))

    def toc_level(self, text):
        if text == "Executive Summary":
            return 0
        if text.startswith("Appendix "):
            return 1
        if re.match(r"^\d+\.\d+\.\d+\s", text):
            return 2
        if re.match(r"^\d+\.\d+\s", text):
            return 1
        if re.match(r"^\d+\.\s", text):
            return 0
        return 0

    def add_toc(self):
        for entry in self.toc_entries:
            self.add_index_line(entry.text, entry.page or "?", level=self.toc_level(entry.text), base_size=10.5)

    def add_caption_list(self, items):
        for entry in items:
            self.add_index_line(entry.text, entry.page or "?", level=0, base_size=10)

    def convert_paragraph(self, flowable):
        style_name = getattr(flowable.style, "name", "")
        text = plain_text(flowable)
        raw_text = getattr(flowable, "text", text)

        if style_name == "CoverTitle":
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_run_font(run, size=24, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
        elif style_name == "CoverSub":
            p = self.doc.add_paragraph(style="CoverSubtitle")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_run_font(run, size=14, italic=True, color=RGBColor(0x44, 0x44, 0x44))
        elif style_name == "CoverDetail":
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            parse_markup_runs(p, raw_text, default_size=11.5)
        elif style_name == "FrontHeading":
            self.add_front_heading(text)
        elif style_name == "H1":
            self.front_heading = None
            p = self.doc.add_heading(level=1)
            run = p.add_run(text)
            set_run_font(run, size=16, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
        elif style_name == "H2":
            self.front_heading = None
            p = self.doc.add_heading(level=2)
            run = p.add_run(text)
            set_run_font(run, size=13, bold=True, color=RGBColor(0x2F, 0x55, 0x97))
        elif style_name == "H3":
            self.front_heading = None
            p = self.doc.add_heading(level=3)
            run = p.add_run(text)
            set_run_font(run, size=11.5, bold=True, color=RGBColor(0x2F, 0x55, 0x97))
        elif style_name in {"FigCap", "TabCap"}:
            p = self.doc.add_paragraph(style="Caption")
            run = p.add_run(text)
            set_run_font(run, size=10.5, bold=True)
        elif style_name == "FigDesc":
            p = self.doc.add_paragraph()
            run = p.add_run(text)
            set_run_font(run, size=9.5)
        elif style_name == "Ref":
            self.add_reference_paragraph(text)
        else:
            self.add_body_paragraph(text)

    def convert_flowable(self, flowable):
        name = type(flowable).__name__
        if name == "Spacer":
            height = getattr(flowable, "height", 0)
            if height >= 18:
                self.doc.add_paragraph()
        elif name == "PageBreak":
            self.doc.add_page_break()
            self.front_heading = None
        elif name == "Paragraph":
            self.convert_paragraph(flowable)
        elif name == "TableOfContents":
            if self.front_heading == "Table of Contents":
                self.add_toc()
            elif self.front_heading == "List of Figures":
                self.add_caption_list(self.figure_entries)
            elif self.front_heading == "List of Tables":
                self.add_caption_list(self.table_entries)
        elif name in {"LongTable", "Table"}:
            convert_table(self.doc, flowable)
        elif name == "Image":
            convert_image(self.doc, flowable)
        elif name == "KeepTogether":
            for item in flowable._content:
                self.convert_flowable(item)
        elif name == "Preformatted":
            convert_preformatted(self.doc, flowable)



def build_docx():
    if not PDF_PATH.exists():
        raise FileNotFoundError(f"PDF page-map source not found: {PDF_PATH}")

    module = load_source_module()
    flowables = module.build_story()
    toc_entries, figure_entries, table_entries = collect_catalog_from_source()

    page_texts = extract_pdf_page_texts(PDF_PATH)
    body_start = find_body_start(page_texts)
    toc_entries = assign_pages(toc_entries, page_texts, body_start)
    figure_entries = assign_pages(figure_entries, page_texts, body_start)
    table_entries = assign_pages(table_entries, page_texts, body_start)

    doc = Document()
    configure_document(doc)
    converter = Converter(doc, toc_entries, figure_entries, table_entries)

    for flowable in flowables:
        converter.convert_flowable(flowable)

    normalize_toc_paragraphs(doc)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FINAL_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    shutil.copyfile(OUT_DOCX, FINAL_DOCX)

    print(OUT_DOCX)
    print(FINAL_DOCX)
    print(f"TOC_ENTRIES={len(toc_entries)} FIGURES={len(figure_entries)} TABLE_CAPTIONS={len(table_entries)}")
    print(f"BODY_START_PAGE={body_start + 1}")
    print(f"PAGE_MAP_SOURCE={PDF_PATH}")


if __name__ == "__main__":
    build_docx()

